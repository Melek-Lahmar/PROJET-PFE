from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

def add_chapter_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_fig(doc, num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Insérer ici la figure]')
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Figure ' + str(num) + ' : ' + caption)
    run2.bold = True
    run2.font.size = Pt(10)
    doc.add_paragraph()

def add_tab_label(doc, num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Tableau ' + str(num) + ' : ' + caption)
    run.bold = True
    run.font.size = Pt(10)

def add_uc_table(doc, tab_num, rows_data):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    labels = ['Nom du cas d\'utilisation', 'Acteur principal', 'Acteurs secondaires',
              'Preconditions', 'Postconditions', 'Scenario nominal',
              'Scenarios alternatifs', 'Scenarios d\'exception']
    for i, (label, val) in enumerate(zip(labels, rows_data)):
        row = table.rows[i]
        row.cells[0].text = label
        run = row.cells[0].paragraphs[0].runs
        if run:
            run[0].bold = True
        row.cells[1].text = val
    doc.add_paragraph()

# ============================================================
# CHAPITRE 5
# ============================================================

doc.add_page_break()
add_chapter_title(doc, 'CHAPITRE 5')
add_chapter_title(doc, 'Sprint 3 : Traitement des commandes, bons de livraison et gestion des reclamations')
doc.add_paragraph()

add_h1(doc, '1. Introduction')
add_body(doc, "Ce chapitre presente le troisieme sprint de notre solution e-commerce et logistique. Il porte sur l'ensemble des fonctionnalites dediees au role de confirmateur : traitement des bons de commande (BC), transformation en bons de livraison (BL), gestion des reclamations clients, suivi des demandes de correction, echanges de messages en temps reel et pilotage des indicateurs d'activite. Ce sprint constitue le coeur du workflow commercial, car il assure la transition entre l'acte d'achat et la preparation de la livraison, tout en gerant les incidents et les demandes formulees par les clients.")
add_body(doc, "Sur le plan technique, le sprint s'appuie sur la Web API ASP.NET Core qui expose les endpoints du controleur confirmateur, sur l'application web React qui fournit les interfaces de traitement, et sur l'application mobile Flutter qui permet au confirmateur de gerer les reclamations et de consulter les commandes depuis un appareil mobile. La communication en temps reel est assuree par SignalR, qui diffuse les evenements metier (nouvelle reclamation, changement de statut, reassignation de cas) a tous les confirmateurs connectes.")

add_h1(doc, '2. Backlog du sprint 3')
add_tab_label(doc, 14, 'Backlog du Sprint 3')
doc.add_paragraph()

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(['ID', 'User Story', 'Taches principales', 'Responsable', 'Estimation']):
    hdr[i].text = h
    if hdr[i].paragraphs[0].runs:
        hdr[i].paragraphs[0].runs[0].bold = True

us_rows = [
    ('US1', 'En tant que confirmateur, je peux consulter les bons de commande a traiter.',
     'Developper la liste BC avec filtres de statut ; consommer /api/confirmateur/commandes ; afficher les informations principales.',
     'Melek Lahmar', '3 j'),
    ('US2', 'En tant que confirmateur, je peux consulter le detail d\'un bon de commande.',
     'Creer l\'interface detail BC ; afficher les lignes, informations client, paiement et livraison.',
     'Melek Lahmar', '3 j'),
    ('US3', 'En tant que confirmateur, je peux modifier le statut d\'un bon de commande.',
     'Integrer l\'action de changement de statut (0 a 3) ; appeler /api/confirmateur/commandes/{piece}/status ; gerer les erreurs.',
     'Tawfik Siala', '3 j'),
    ('US4', 'En tant que confirmateur, je peux transformer un BC en bon de livraison.',
     'Declencher /api/confirmateur/commandes/{piece}/transform-to-bl ; gerer le doublon BL ; afficher la reference generee.',
     'Tawfik Siala', '5 j'),
    ('US5', 'En tant que confirmateur, je peux consulter les bons de livraison generes.',
     'Creer la liste BL (/api/confirmateur/bl) ; integrer le detail BL ; afficher les lignes et informations de livraison.',
     'Melek Lahmar', '3 j'),
    ('US6', 'En tant que confirmateur, je peux gerer les reclamations des clients.',
     'Lister /api/confirmateur/reclamations ; consulter le detail ; changer le statut ; prendre en charge un cas ; creer un echange.',
     'Melek Lahmar / Tawfik Siala', '5 j'),
    ('US7', 'En tant que confirmateur, je peux envoyer une demande de correction a un client.',
     'Appeler /api/confirmateur/reclamations/{id}/correction ; saisir la nature de la demande ; notifier le client.',
     'Tawfik Siala', '3 j'),
    ('US8', 'En tant que confirmateur, je peux consulter et repondre aux messages d\'une reclamation.',
     'Afficher le fil de messages ; envoyer un nouveau message ; integrer les mises a jour en temps reel via SignalR.',
     'Melek Lahmar', '3 j'),
    ('US9', 'En tant que confirmateur, je peux consulter mes indicateurs de performance.',
     'Consommer /api/confirmateur/status/me/stats ; afficher cas actifs, clotures, taux ; gerer la pause activite.',
     'Tawfik Siala', '2 j'),
    ('US10', 'En tant qu\'administrateur, je peux consulter les tableaux de bord de suivi des commandes.',
     'Integrer les endpoints statistiques ; afficher les indicateurs globaux des bons de commande et de livraison.',
     'Melek Lahmar', '3 j'),
]

for us in us_rows:
    row = table.add_row().cells
    for i, val in enumerate(us):
        row[i].text = val

doc.add_paragraph()

add_h1(doc, '3. Les activites du sprint 3')
add_body(doc, "Les activites du Sprint 3 couvrent l'ensemble du cycle de traitement des commandes et des reclamations. Elles consistent a implementer les workflows de validation documentaire, a mettre en place les mecanismes de communication entre le confirmateur et le client, et a assurer la coherence des donnees en temps reel.")
for item in [
    "Analyse du workflow de confirmation des bons de commande et de la transition BC vers BL ;",
    "Conception des interactions entre le confirmateur, l'interface React, l'application mobile Flutter et la Web API ;",
    "Developpement des interfaces de liste, detail et changement de statut des bons de commande ;",
    "Mise en place de la transformation documentaire BC vers BL avec gestion des doublons et des erreurs metier ;",
    "Implementation du module de gestion des reclamations : liste, detail, assignation, changement de statut, creation d'echange ;",
    "Developpement du systeme de demandes de correction adressees aux clients ;",
    "Integration du chat en temps reel entre confirmateur et client sur chaque reclamation via SignalR ;",
    "Mise en place du suivi de session confirmateur (pause, reprise, statistiques personnelles) ;",
    "Validation des regles metier liees aux etats des documents, aux droits d'acces et aux transitions de statut.",
]:
    add_bullet(doc, item)

add_h1(doc, '4. Les activites d\'analyse et de conception')
add_body(doc, "L'analyse et la conception du Sprint 3 formalisent les regles de passage d'un bon de commande vers un bon de livraison, ainsi que les flux de traitement des reclamations. Cette etape permet d'eviter les incoherences entre les documents commerciaux, les informations client et les donnees de la base SQL Server. Elle precise egalement les roles de chaque acteur dans le circuit de validation et les regles d'assignation des cas de reclamation entre confirmateurs.")

add_h1(doc, '5. Diagramme de cas d\'utilisation du sprint 3')
add_body(doc, "Le diagramme de cas d'utilisation du Sprint 3 presente les actions realisees par le confirmateur dans le traitement des commandes, des bons de livraison et des reclamations. L'administrateur peut suivre les indicateurs associes a ce processus.")
add_fig(doc, 74, 'Diagramme de cas d\'utilisation du Sprint 3')
add_body(doc, "Ce diagramme illustre les trois grands domaines d'action du confirmateur : le traitement documentaire (BC/BL), la gestion des reclamations (assignation, traitement, cloture, creation d'echange) et le pilotage de son activite (session, pause, statistiques). L'administrateur dispose d'une vue globale sur les indicateurs de performance.")

add_h1(doc, '6. Description textuelle des cas d\'utilisation')

uc_data = [
    ('Consulter les bons de commande', 'Confirmateur',
     'Web API ASP.NET Core, base SQL Server',
     'Le confirmateur doit etre authentifie avec le role CONFIRMATEUR.',
     'La liste des bons de commande correspondant aux criteres est affichee.',
     'Le confirmateur ouvre l\'interface des commandes. Le systeme interroge /api/confirmateur/commandes. Les BC sont charges depuis la base et affiches avec leur statut (EN_ATTENTE, CONFIRME, REFUSE, TENTATIVE).',
     'Le confirmateur peut filtrer la liste selon le statut ou rechercher une commande par reference.',
     'Si l\'utilisateur n\'a pas le role requis, l\'acces est refuse (401). En cas d\'erreur serveur, un message est affiche.'),
    ('Modifier le statut d\'un bon de commande', 'Confirmateur',
     'Web API ASP.NET Core, base SQL Server',
     'Le bon de commande doit exister et etre accessible au confirmateur.',
     'Le statut de la commande est mis a jour dans la base de donnees.',
     'Le confirmateur consulte la commande, choisit le nouveau statut (0=Annule, 1=Confirme, 2=En livraison, 3=Livre) et valide. La Web API verifie la commande, met a jour l\'etat et retourne une confirmation.',
     'Le confirmateur peut annuler l\'action avant validation.',
     'Si le statut demande est invalide ou si la commande est introuvable, la mise a jour est refusee avec un message explicite.'),
    ('Transformer un BC en bon de livraison', 'Confirmateur',
     'Web API ASP.NET Core, service metier, base SQL Server',
     'Le bon de commande doit exister et etre dans un etat compatible avec la transformation.',
     'Un bon de livraison est genere et associe aux lignes correspondantes.',
     'Le confirmateur demande la transformation via POST /api/confirmateur/commandes/{piece}/transform-to-bl. Le service verifie l\'existence du BC, controle l\'absence de doublon BL, cree le bon de livraison, met a jour les donnees et retourne la reference du BL genere.',
     'Si un BL existe deja pour la commande, le systeme retourne la reference du document existant afin d\'eviter les doublons.',
     'Si le BC est introuvable, si le stock est insuffisant ou si une erreur de transaction survient, la transformation est annulee et un message d\'erreur est retourne.'),
    ('Gerer les reclamations clients', 'Confirmateur',
     'Web API ASP.NET Core, base SQL Server, SignalR Hub',
     'Le confirmateur doit etre authentifie. Des reclamations doivent etre presentes dans la base.',
     'La reclamation est traitee : statut mis a jour, echange cree, correction envoyee ou cas reassigne.',
     'Le confirmateur consulte la liste via GET /api/confirmateur/reclamations, selectionne une reclamation, consulte son detail et les tentatives de livraison associees. Il peut prendre en charge le cas (POST /reprendre), modifier le statut (PUT /status), envoyer une demande de correction au client (PUT /correction), creer un bon d\'echange (POST /echange) ou reassigner le cas a un autre confirmateur (PUT /assign).',
     'Le confirmateur peut filtrer les reclamations par statut, motif, source, type de cas ou periode. Il peut annoter la reclamation avec une note interne.',
     'Si la reclamation est deja assignee a un autre confirmateur, la prise en charge est refusee. Si le stock est insuffisant pour un echange, le systeme retourne une erreur metier explicite.'),
    ('Echanger des messages en temps reel avec le client', 'Confirmateur',
     'Web API ASP.NET Core, SignalR Hub, base SQL Server',
     'La reclamation doit exister. Le confirmateur et le client doivent etre connectes au hub SignalR.',
     'Le message est enregistre et diffuse en temps reel au destinataire.',
     'Le confirmateur ouvre le fil de discussion associe a une reclamation. Il saisit son message et l\'envoie. Le hub SignalR diffuse le message au client concerne. Les nouvelles reclamations et reassignations sont egalement notifiees en temps reel (NouveauCas, CasReattribue, StatutCasChange).',
     'Si le client n\'est pas connecte, le message est conserve en base et sera visible a sa prochaine connexion.',
     'En cas de deconnexion SignalR, le systeme applique une periode de grace de 5 secondes avant de liberer les cas assignes, afin d\'eviter les interruptions lors des changements de reseau.'),
    ('Consulter ses indicateurs de performance', 'Confirmateur',
     'Web API ASP.NET Core, base SQL Server',
     'Le confirmateur doit etre authentifie.',
     'Les statistiques de performance sont affichees (cas actifs, clotures aujourd\'hui, cette semaine, ce mois).',
     'Le confirmateur accede a son profil ou tableau de bord. Le systeme interroge GET /api/confirmateur/status/me/stats et retourne les indicateurs : nombre de cas actifs, cas clotures par periode, taux de resolution. Le confirmateur peut aussi activer ou desactiver sa disponibilite (pause/reprise).',
     'L\'administrateur dispose d\'un tableau de bord global avec les statistiques de tous les confirmateurs.',
     'Si aucune donnee n\'est disponible pour la periode, les indicateurs affichent zero.'),
]

tab_nums = [15, 16, 17, 18, 19, 20]
for i, uc in enumerate(uc_data):
    sub = str(i + 1)
    add_h2(doc, '6.' + sub + '. ' + uc[0])
    add_uc_table(doc, tab_nums[i], list(uc))
    add_tab_label(doc, tab_nums[i], 'Description textuelle du cas d\'utilisation « ' + uc[0] + ' »')
    doc.add_paragraph()

add_h1(doc, '7. Diagramme de classes du sprint 3')
add_body(doc, "Le diagramme de classes du Sprint 3 represente les entites manipulees lors du traitement des commandes et des reclamations. Il met en evidence les relations entre les documents commerciaux, les reclamations et les sessions confirmateur.")
add_fig(doc, 75, 'Diagramme de classes du Sprint 3')
add_body(doc, "Le modele repose sur l'entite F_DOCENTETE, qui porte l'en-tete du document (BC ou BL), et F_DOCLIGNE, qui detaille les articles. La transformation BC vers BL s'appuie sur ces entites et sur les controles de disponibilite. L'entite F_RECLAMATION centralise les informations de la reclamation (statut, confirmateur assigne, date de cloture, notes internes, tentatives de livraison). F_RECLAMATION_TENTATIVE enregistre chaque tentative de livraison avec ses coordonnees GPS, son motif et l'eventuelle photo d'incident. F_CONFIRMATRICE_SESSION trace les connexions et deconnexions des confirmateurs a des fins d'audit.")

add_h1(doc, '8. Diagrammes de sequence')

seq3 = [
    (76, '8.1. Modification du statut d\'un bon de commande',
     "Ce diagramme montre les echanges entre l'interface React, la Web API et la base SQL Server lors de la modification du statut d'un bon de commande par le confirmateur. Il illustre egalement la gestion des droits d'acces et des erreurs de validation."),
    (77, '8.2. Transformation BC en bon de livraison',
     "Ce diagramme decrit la transformation d'un bon de commande en bon de livraison. Il presente les controles metier (existence du BC, absence de doublon BL, verification du stock), la creation du document de livraison et la gestion des cas d'erreur."),
    (78, '8.3. Prise en charge et traitement d\'une reclamation',
     "Ce diagramme illustre le cycle complet de traitement d'une reclamation : assignation au confirmateur, consultation des tentatives de livraison, changement de statut et notification SignalR au client. Il montre egalement la creation d'un bon d'echange lorsque la reclamation le justifie."),
    (79, '8.4. Communication en temps reel sur une reclamation',
     "Ce diagramme presente les echanges de messages en temps reel entre le confirmateur et le client via le hub SignalR. Il montre la diffusion des notifications de nouvelles reclamations, de reassignation de cas et de changements de statut."),
]

for fig_num, heading, desc in seq3:
    add_h2(doc, heading)
    add_body(doc, desc)
    add_fig(doc, fig_num, heading.split('. ', 1)[1])

add_h1(doc, '9. Maquettes')
maquettes3 = [
    (80, 'Maquette de l\'interface de liste des bons de commande (confirmateur)'),
    (81, 'Maquette de l\'interface de detail d\'un bon de commande'),
    (82, 'Maquette de l\'interface de transformation BC en BL'),
    (83, 'Maquette de l\'interface de liste des bons de livraison'),
    (84, 'Maquette de l\'interface de liste des reclamations'),
    (85, 'Maquette de l\'interface de detail d\'une reclamation avec fil de messages'),
    (86, 'Maquette de l\'interface des indicateurs de performance confirmateur'),
]
for i, (fig_num, caption) in enumerate(maquettes3):
    idx = str(i + 1)
    add_h2(doc, '9.' + idx + '. ' + caption)
    add_fig(doc, fig_num, caption)

add_h1(doc, '10. Realisation')
real3 = [
    (87, '10.1. Interface de liste des bons de commande (confirmateur)',
     "L'interface React permet au confirmateur de consulter l'ensemble des bons de commande a traiter. Elle consomme GET /api/confirmateur/commandes et propose des filtres par statut. Les cartes presentent la reference, le client, la date et le statut courant de chaque document."),
    (88, '10.2. Interface de detail d\'un bon de commande',
     "Cette interface affiche les lignes du bon de commande, les informations client (adresse, telephone, mode de livraison), le mode de paiement, la ventilation de la facture et le statut du document. Elle permet de declencher le changement de statut ou la transformation en BL."),
    (89, '10.3. Interface de transformation en bon de livraison',
     "Cette interface permet au confirmateur de declencher la transformation du bon de commande en bon de livraison. Le systeme verifie l'absence de doublon, genere le BL et affiche sa reference. En cas d'erreur metier (stock insuffisant, BC introuvable), un message explicite est retourne."),
    (90, '10.4. Interface de liste des bons de livraison',
     "Cette interface affiche les bons de livraison generes, avec leur statut (CONFIRME, EN_LIVRAISON, LIVRE, RETOUR, REPORTE, DEPOT), les informations de livraison et les lignes associees."),
    (91, '10.5. Interface de gestion des reclamations',
     "L'interface de reclamations permet au confirmateur de consulter les cas ouverts, de prendre en charge un cas, de modifier son statut, d'envoyer une demande de correction au client, de creer un bon d'echange ou de reassigner le cas. Les nouvelles reclamations arrivent en temps reel via SignalR sans rechargement de la page."),
    (92, '10.6. Interface de chat reclamation (temps reel)',
     "Cette interface affiche le fil de messages associe a une reclamation. Les messages echanges entre le confirmateur et le client sont affiches en temps reel via SignalR. Le confirmateur peut consulter les photos d'incident jointes a la reclamation et les tentatives de livraison avec leurs coordonnees GPS."),
    (93, '10.7. Interface des indicateurs de performance confirmateur',
     "Cette interface presente les statistiques de performance du confirmateur : cas actifs, cas clotures aujourd'hui, cette semaine et ce mois, taux de resolution. Elle permet egalement d'activer ou de desactiver la disponibilite (pause/reprise) depuis l'application mobile Flutter."),
]
for fig_num, heading, desc in real3:
    add_h2(doc, heading)
    add_body(doc, desc)
    add_fig(doc, fig_num, heading.split('. ', 1)[1])

add_h1(doc, '11. Tests et validation')
add_body(doc, "Les tests du Sprint 3 visent a verifier la conformite du traitement des commandes, des bons de livraison et des reclamations. Les validations portent sur les droits d'acces, l'exactitude des donnees affichees, les transitions de statut, la transformation documentaire et le comportement en temps reel.")
add_tab_label(doc, 21, 'Tests fonctionnels du Sprint 3')
doc.add_paragraph()

test3_table = doc.add_table(rows=1, cols=4)
test3_table.style = 'Table Grid'
for i, h in enumerate(['Test', 'Objectif', 'Resultat attendu', 'Statut']):
    test3_table.rows[0].cells[i].text = h

tests3 = [
    ('Acces confirmateur', 'Verifier que seul le role CONFIRMATEUR accede aux commandes et aux reclamations.', 'L\'acces est autorise uniquement au role prevu. Un acces non autorise retourne 401 ou 403.', 'A valider par capture'),
    ('Consultation BC', 'Verifier l\'affichage du detail et des lignes d\'un bon de commande.', 'Les informations du BC sont coherentes avec la base SQL Server.', 'A valider par capture'),
    ('Mise a jour statut BC', 'Verifier le changement de statut 0, 1, 2, 3.', 'Le statut est mis a jour et visible dans la liste.', 'A valider par capture'),
    ('Transformation BC vers BL', 'Verifier la generation d\'un BL a partir d\'un BC.', 'Un bon de livraison est cree sans doublon. La reference du BL est retournee.', 'A valider par capture'),
    ('Gestion des doublons BL', 'Verifier le comportement lorsqu\'un BL existe deja.', 'Le systeme retourne la reference existante sans creer de doublon.', 'A valider par capture'),
    ('Prise en charge reclamation', 'Verifier l\'assignation d\'une reclamation a un confirmateur.', 'La reclamation est associee au confirmateur et retiree de la liste commune.', 'A valider par capture'),
    ('Changement statut reclamation', 'Verifier la cloture ou le refus d\'une reclamation.', 'Le statut est mis a jour et une notification est envoyee au client.', 'A valider par capture'),
    ('Creation bon d\'echange', 'Verifier la creation d\'un BL d\'echange depuis une reclamation.', 'Un bon d\'echange est genere avec les lignes appropriees.', 'A valider par capture'),
    ('Chat en temps reel', 'Verifier la reception des messages SignalR sans rechargement.', 'Les messages arrivent instantanement dans l\'interface du destinataire.', 'A valider par capture'),
    ('Statistiques confirmateur', 'Verifier le calcul des indicateurs de performance.', 'Les cas actifs, clotures et le taux de resolution sont coherents avec la base.', 'A valider par capture'),
]
for t in tests3:
    row = test3_table.add_row().cells
    for i, v in enumerate(t):
        row[i].text = v

doc.add_paragraph()

add_h1(doc, '12. Conclusion du sprint 3')
add_body(doc, "Le Sprint 3 a permis de mettre en place le traitement interne des commandes et la gestion complete des reclamations. Il a couvert la consultation et la modification des bons de commande, la transformation BC vers BL avec gestion des doublons, la consultation des bons de livraison generes, ainsi que le circuit complet de traitement des reclamations : assignation, communication en temps reel avec le client, creation de bons d'echange et suivi des indicateurs de performance. Ce sprint consolide la chaine commerciale et prepare le Sprint 4, consacre au suivi mobile des livraisons par le livreur, a l'optimisation des tournees et au suivi en temps reel de la position du livreur par le client.")

# ============================================================
# CHAPITRE 6
# ============================================================

doc.add_page_break()
add_chapter_title(doc, 'CHAPITRE 6')
add_chapter_title(doc, 'Sprint 4 : Suivi mobile des livraisons, optimisation des tournees et transit inter-depots')
doc.add_paragraph()

add_h1(doc, '1. Introduction')
add_body(doc, "Ce chapitre presente le quatrieme et dernier sprint de notre projet. Il est consacre au suivi mobile des livraisons et a l'ensemble des fonctionnalites operationnelles accessibles depuis l'application mobile Flutter. Ce sprint couvre les besoins du livreur standard (consultation des livraisons, affectation, mise a jour du statut, localisation GPS, optimisation de tournee et caisse), du livreur de transit (missions inter-depots avec scan de codes-barres), ainsi que les fonctionnalites dediees au client mobile (suivi de commande en temps reel, reclamations, chat).")
add_body(doc, "Sur le plan technique, le sprint met en relation l'application mobile Flutter, la Web API ASP.NET Core et la base SQL Server. La localisation GPS est geree par le plugin Geolocator, la communication en temps reel s'appuie sur SignalR (diffusion de la position du livreur en direct au client), et l'optimisation des tournees utilise l'algorithme du voisin le plus proche avec calcul de distance Haversine combine au service OSRM pour les ETA de livraison.")

add_h1(doc, '2. Backlog du sprint 4')
add_tab_label(doc, 22, 'Backlog du Sprint 4')
doc.add_paragraph()

table4 = doc.add_table(rows=1, cols=5)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
for i, h in enumerate(['ID', 'User Story', 'Taches principales', 'Responsable', 'Estimation']):
    hdr4[i].text = h

us6_rows = [
    ('US1', 'En tant que livreur, je peux consulter les livraisons disponibles dans ma zone.',
     'Developper l\'ecran mobile des nouvelles livraisons ; consommer /api/livreur/pool/disponibles ; filtrer par gouvernorat et delegation du livreur.',
     'Tawfik Siala', '4 j'),
    ('US2', 'En tant que livreur, je peux prendre en charge une livraison.',
     'Implementer l\'action d\'affectation POST /api/livreur/pool/{doPiece}/prendre ; gerer le conflit si la livraison est deja prise ; mettre a jour la liste.',
     'Melek Lahmar', '3 j'),
    ('US3', 'En tant que livreur, je peux consulter mes livraisons affectees.',
     'Developper l\'ecran mobile de mes commandes ; charger via /api/livreur/pool/mes-livraisons ; afficher les informations de contact, adresse, paiement et statut.',
     'Tawfik Siala', '3 j'),
    ('US4', 'En tant que livreur, je peux mettre a jour le statut d\'une livraison.',
     'Preparer les statuts (EN_LIVRAISON, LIVRE, REPORTE, RETOUR, DEPOT, TENTATIVE) ; envoyer via PUT /api/livreur/orders/{piece}/status ; gerer la date de replanification.',
     'Melek Lahmar', '4 j'),
    ('US5', 'En tant que livreur, je peux demarrer une livraison active avec diffusion GPS.',
     'Implementer POST /start-heading ; activer le service de localisation (Geolocator) ; envoyer la position toutes les 10 secondes via POST /api/livreur/location/ping ; diffuser la position au client via SignalR.',
     'Tawfik Siala', '5 j'),
    ('US6', 'En tant que livreur, je peux optimiser mon itineraire de tournee.',
     'Appeler /api/livreur/tournee/optimize ; afficher l\'ordre optimal des livraisons (algorithme voisin le plus proche + Haversine) ; integrer les ETA via OSRM.',
     'Melek Lahmar', '4 j'),
    ('US7', 'En tant que livreur de transit, je peux gerer mes missions inter-depots.',
     'Developper l\'ecran transit_home_screen ; lister les missions /api/transit/my-missions ; scanner les codes-barres avec mobile_scanner pour valider la reception des articles.',
     'Tawfik Siala', '5 j'),
    ('US8', 'En tant que livreur, je peux consulter mes statistiques de performance.',
     'Consommer /api/livreur/stats ; afficher livraisons du jour, taux de reussite, montant encaisse, tendance 7 jours ; gerer la caisse COD.',
     'Melek Lahmar', '3 j'),
    ('US9', 'En tant que client, je peux suivre ma commande en temps reel depuis le mobile.',
     'Afficher l\'etat adaptatif (AT_DEPOT, IN_DELIVERY_QUEUE, HEADING_TO_YOU, TERMINAL) ; afficher la carte live avec la position du livreur quand HEADING_TO_YOU ; consommer /api/tracking-state.',
     'Tawfik Siala', '4 j'),
    ('US10', 'En tant que client mobile, je peux deposer et suivre une reclamation.',
     'Creer une reclamation POST /api/reclamations ; consulter la liste et le detail ; echanger des messages avec le confirmateur via SignalR ; joindre des photos (support HEIC/HEIF).',
     'Melek Lahmar', '4 j'),
]

for us in us6_rows:
    row = table4.add_row().cells
    for i, val in enumerate(us):
        row[i].text = val

doc.add_paragraph()

add_h1(doc, '3. Les activites du sprint 4')
add_body(doc, "Les activites du Sprint 4 sont orientees vers l'usage mobile et le suivi operationnel en conditions reelles. Elles consistent a adapter les donnees de livraison a un contexte terrain, a securiser les actions par role, a assurer la continuite de la communication entre le livreur et le client, et a maintenir la coherence entre le document commercial et le suivi de livraison.")
for item in [
    "Analyse du parcours livreur et des statuts de livraison (EN_LIVRAISON, LIVRE, REPORTE, RETOUR, DEPOT, TENTATIVE) ;",
    "Conception des echanges entre Flutter, la Web API ASP.NET Core et la base SQL Server ;",
    "Developpement des ecrans mobiles de livraisons disponibles et de livraisons affectees ;",
    "Implementation de la prise en charge d'une livraison avec gestion des conflits d'affectation ;",
    "Integration du service de localisation GPS (Geolocator) avec diffusion en temps reel via SignalR ;",
    "Developpement de l'optimisation de tournee (voisin le plus proche + Haversine + ETA via OSRM) ;",
    "Implementation du module transit inter-depots avec scan de codes-barres (mobile_scanner) ;",
    "Developpement du tableau de bord livreur avec statistiques quotidiennes, hebdomadaires et mensuelles ;",
    "Mise en place du suivi client adaptatif (TrackingStateCard : AT_DEPOT, IN_DELIVERY_QUEUE, HEADING_TO_YOU, TERMINAL) ;",
    "Integration du module de reclamations client mobile avec chat en temps reel et support photo HEIC/HEIF.",
]:
    add_bullet(doc, item)

add_h1(doc, '4. Les activites d\'analyse et de conception')
add_body(doc, "L'analyse et la conception du Sprint 4 definissent la maniere dont les bons de livraison sont exposes aux livreurs selon leur zone geographique (gouvernorat/delegation). Elles precisent les regles d'affectation, les statuts autorises et leurs transitions, ainsi que les mecanismes de diffusion de la position GPS du livreur au client concerne. Elles formalisent egalement le workflow de transit inter-depots, qui implique la validation de la reception physique des articles par scan de code-barres.")

add_h1(doc, '5. Diagramme de cas d\'utilisation du sprint 4')
add_body(doc, "Le diagramme de cas d'utilisation du Sprint 4 presente les fonctionnalites mobiles liees aux roles livreur, livreur de transit et client. Il met en evidence les services externes utilises : SignalR pour la diffusion de position, OSRM pour les ETA de livraison et Geolocator pour la capture GPS.")
add_fig(doc, 88, 'Diagramme de cas d\'utilisation du Sprint 4')
add_body(doc, "Ce diagramme montre que le livreur standard peut consulter les livraisons disponibles, en prendre en charge, consulter ses livraisons affectees, mettre a jour leur statut, demarrer une livraison active avec diffusion GPS et optimiser sa tournee. Le livreur de transit gere ses missions inter-depots par scan de codes-barres. Le client peut suivre sa commande en temps reel et interagir avec le confirmateur via le module de reclamations.")

add_h1(doc, '6. Description textuelle des cas d\'utilisation')

uc6_data = [
    ('Consulter les livraisons disponibles', 'Livreur',
     'Application Flutter, Web API ASP.NET Core, base SQL Server',
     'Le livreur doit etre authentifie avec le role LIVREUR et disposer d\'un profil avec gouvernorat defini.',
     'La liste des livraisons disponibles filtrees par zone est affichee.',
     'Le livreur ouvre l\'ecran des nouvelles livraisons. L\'application interroge GET /api/livreur/pool/disponibles. Le backend identifie le gouvernorat/delegation du livreur, filtre les BL non affectes et retourne la liste avec les informations de livraison.',
     'Si aucune livraison n\'est disponible dans la zone, l\'application affiche un etat vide avec un message informatif.',
     'Si le profil livreur est introuvable ou si le token JWT est invalide ou expire, l\'acces est refuse.'),
    ('Prendre en charge une livraison', 'Livreur',
     'Application Flutter, Web API ASP.NET Core, base SQL Server',
     'La livraison doit exister, etre disponible et ne pas etre deja affectee a un autre livreur.',
     'La livraison est associee au livreur dans la base. Une entree est creee dans F_LIVRAISON.',
     'Le livreur selectionne une livraison et confirme la prise en charge. L\'application appelle POST /api/livreur/pool/{doPiece}/prendre. La Web API verifie l\'absence d\'affectation, cree l\'entree F_LIVRAISON associee au livreur et retourne une confirmation.',
     'Le livreur peut annuler l\'action avant confirmation. Il peut egalement abandonner une livraison deja prise via POST /abandon avec une note justificative.',
     'Si la livraison est deja prise par un autre livreur, le systeme retourne un conflit 409 et l\'application informe le livreur avec un message clair.'),
    ('Mettre a jour le statut d\'une livraison', 'Livreur',
     'Application Flutter, Web API ASP.NET Core, base SQL Server',
     'La livraison doit etre affectee au livreur connecte.',
     'Le statut, les notes et les dates utiles de la livraison sont mis a jour dans F_LIVRAISON.',
     'Le livreur selectionne une livraison, choisit le nouveau statut et valide. L\'API verifie l\'affectation et normalise le statut via PUT /api/livreur/orders/{piece}/status. Pour les statuts LIVRE et REPORTE, les dates correspondantes sont automatiquement enregistrees. Le montant est encaisse automatiquement pour les paiements en especes.',
     'Si le statut est REPORTE, une date de replanification doit etre fournie. La mise a jour par lot (batch) est possible via PUT /api/livreur/orders/batch-status pour traiter plusieurs commandes simultanement.',
     'Si le statut est invalide, si la livraison est introuvable ou si la date de report est absente pour un report, l\'operation est refusee avec un message explicite.'),
    ('Demarrer une livraison active avec diffusion GPS', 'Livreur',
     'Application Flutter, Geolocator, Web API ASP.NET Core, SignalR Hub',
     'Le livreur doit avoir une livraison affectee. Les permissions de localisation doivent etre accordees sur l\'appareil.',
     'La position du livreur est diffusee en temps reel au client concerne via SignalR.',
     'Le livreur demarre la livraison active via POST /api/livreur/orders/{piece}/start-heading. Le service LivreurLocationService capture la position GPS toutes les 10 secondes et l\'envoie via POST /api/livreur/location/ping. La Web API met a jour F_LIVREUR_POSITION et diffuse les coordonnees au client via SignalR (evenement LocationUpdate). Les positions hors connexion sont stockees dans une file d\'attente locale et envoyees en lot via POST /ping-batch a la reconnexion.',
     'Si la connexion est perdue, la file d\'attente locale accumule les positions et les envoie par lot a la reconnexion. Le livreur peut arreter la diffusion via POST /stop-heading.',
     'Si les permissions de localisation sont refusees, la diffusion GPS est desactivee. Un message informe le livreur que le client ne pourra pas suivre sa position en direct.'),
    ('Optimiser la tournee de livraison', 'Livreur',
     'Application Flutter, Web API ASP.NET Core (algorithme Haversine + OSRM)',
     'Le livreur doit avoir plusieurs livraisons affectees avec coordonnees GPS disponibles.',
     'L\'ordre optimal des livraisons est calcule et affiche avec les distances et ETA estimes.',
     'Le livreur demande l\'optimisation de sa tournee. L\'application appelle GET /api/livreur/tournee/optimize. Le backend applique l\'algorithme du voisin le plus proche avec la distance Haversine pour ordonner les arrets, puis consulte le service OSRM pour calculer les ETA de livraison. La liste ordonnee est retournee avec les distances et temps estimes par arret.',
     'Si des coordonnees sont manquantes pour certaines livraisons, elles sont integrees en fin de tournee.',
     'Si le service OSRM est indisponible, les ETA sont estimes a partir de la distance Haversine seule.'),
    ('Gerer une mission de transit inter-depots', 'Livreur de transit',
     'Application Flutter (mobile_scanner), Web API ASP.NET Core, base SQL Server',
     'Le livreur doit avoir le role LIVREUR et une mission de transit assignee dans F_TRANSFERTS.',
     'Les articles scannes sont marques comme recus dans la base. La mission est mise a jour.',
     'Le livreur de transit ouvre l\'ecran transit_home_screen et consulte ses missions. Il selectionne une mission et accede a transit_mission_details_screen. Il scanne les codes-barres des articles recus via transit_barcode_scanner_screen (plugin mobile_scanner avec retour sonore). Chaque scan valide la reception d\'un article et met a jour F_TRANSFERTS.',
     'Le livreur peut scanner plusieurs articles en sequence. Un retour sonore confirme chaque scan reussi.',
     'Si le code-barres ne correspond a aucun article de la mission, un message d\'erreur est affiche. Si la camera est indisponible, le livreur peut saisir le code manuellement.'),
    ('Suivre une commande en temps reel (client mobile)', 'Client',
     'Application Flutter, Web API ASP.NET Core, SignalR Hub',
     'Le client doit etre authentifie. Une commande doit lui etre associee.',
     'L\'etat adaptatif de la commande est affiche avec les informations de suivi correspondantes.',
     'Le client ouvre l\'ecran de suivi client_order_tracking_screen. La carte TrackingStateCard interroge /api/tracking-state et affiche l\'etat adaptatif : AT_DEPOT (commande en depot), IN_DELIVERY_QUEUE (en cours de livraison), HEADING_TO_YOU (livreur en route avec carte live et ETA), ou TERMINAL (livre/retourne). En etat HEADING_TO_YOU, la carte live LiveDeliveryMapSheet affiche la position du livreur mise a jour toutes les 15 secondes.',
     'Le client peut appeler ou envoyer un SMS au livreur directement depuis l\'ecran de suivi.',
     'Si les coordonnees GPS du livreur sont indisponibles, le bouton de carte live est masque. Si la connexion SignalR est perdue, le suivi bascule sur un mode de rafraichissement toutes les 15 secondes.'),
]

tab6_nums = [23, 24, 25, 26, 27, 28, 29]
for i, uc in enumerate(uc6_data):
    sub = str(i + 1)
    add_h2(doc, '6.' + sub + '. ' + uc[0])
    add_uc_table(doc, tab6_nums[i], list(uc))
    add_tab_label(doc, tab6_nums[i], 'Description textuelle du cas d\'utilisation « ' + uc[0] + ' »')
    doc.add_paragraph()

add_h1(doc, '7. Diagramme de classes du sprint 4')
add_body(doc, "Le diagramme de classes du Sprint 4 represente les entites necessaires au suivi des livraisons, a la localisation GPS et aux missions de transit. Il relie le bon de livraison, le profil livreur, l'entite de suivi et les tables de position.")
add_fig(doc, 89, 'Diagramme de classes du Sprint 4')
add_body(doc, "L'entite F_LIVRAISON memorise l'affectation d'une livraison a un livreur, son statut, ses dates cles (creation, livraison, replanification), le montant encaisse et la note de suivi. F_LIVREUR_POSITION stocke la position GPS courante du livreur (latitude, longitude, timestamp, etat de diffusion). F_LIVREUR_POSITION_HISTORY conserve l'historique des positions pour audit. F_LIVREUR_ABANDON_LOG trace les abandons avec leur motif. F_TRANSFERTS represente les missions de transit inter-depots avec les articles a valider par scan.")

add_h1(doc, '8. Diagrammes de sequence')

seq4 = [
    (90, '8.1. Consultation des livraisons disponibles',
     "Ce diagramme montre comment l'application Flutter recupere les bons de livraison disponibles pour le livreur connecte, filtres selon son gouvernorat et sa delegation."),
    (91, '8.2. Prise en charge d\'une livraison',
     "Ce diagramme decrit l'affectation d'un bon de livraison a un livreur. Il presente egalement la gestion du conflit lorsqu'une livraison a deja ete prise par un autre livreur."),
    (92, '8.3. Mise a jour du statut de livraison',
     "Ce diagramme presente la modification du statut d'une livraison depuis l'application mobile, avec les controles d'affectation, la normalisation du statut, l'enregistrement des dates et la gestion automatique de l'encaissement."),
    (93, '8.4. Diffusion GPS en temps reel',
     "Ce diagramme illustre la sequence complete de demarrage d'une livraison active, depuis l'activation du service Geolocator jusqu'a la diffusion de la position au client via le hub SignalR, en passant par l'UPSERT de F_LIVREUR_POSITION."),
    (94, '8.5. Suivi client en temps reel (TrackingStateCard)',
     "Ce diagramme decrit la logique adaptative de l'ecran de suivi client : interrogation de /api/tracking-state, selection de l'etat (AT_DEPOT, IN_DELIVERY_QUEUE, HEADING_TO_YOU, TERMINAL) et affichage de la carte live en etat HEADING_TO_YOU."),
    (95, '8.6. Scan de codes-barres pour le transit',
     "Ce diagramme presente le workflow de validation d'un article de transit : ouverture de la camera mobile_scanner, scan du code-barres, envoi au backend, mise a jour de F_TRANSFERTS et retour sonore de confirmation."),
]

for fig_num, heading, desc in seq4:
    add_h2(doc, heading)
    add_body(doc, desc)
    add_fig(doc, fig_num, heading.split('. ', 1)[1])

add_h1(doc, '9. Maquettes')
maquettes4 = [
    (96, 'Maquette de l\'ecran mobile de connexion livreur'),
    (97, 'Maquette de l\'ecran des livraisons disponibles'),
    (98, 'Maquette de l\'ecran de mes livraisons affectees'),
    (99, 'Maquette de l\'ecran detail livraison avec actions de suivi'),
    (100, 'Maquette de l\'ecran de carte et itineraire optimise'),
    (101, 'Maquette de l\'ecran de statistiques livreur'),
    (102, 'Maquette de l\'ecran de missions de transit inter-depots'),
    (103, 'Maquette de l\'ecran de scan de codes-barres'),
    (104, 'Maquette de l\'ecran de suivi commande client (TrackingStateCard)'),
    (105, 'Maquette de l\'ecran de reclamation client mobile avec chat'),
]
for i, (fig_num, caption) in enumerate(maquettes4):
    idx = str(i + 1)
    add_h2(doc, '9.' + idx + '. ' + caption)
    add_fig(doc, fig_num, caption)

add_h1(doc, '10. Realisation')
real4 = [
    (106, '10.1. Interface mobile des livraisons disponibles',
     "L'ecran Flutter new_orders_screen permet au livreur de consulter les livraisons disponibles dans sa zone geographique. Il consomme GET /api/livreur/pool/disponibles et affiche les cartes de livraison avec l'adresse, le montant, le mode de paiement et les informations de contact du destinataire."),
    (107, '10.2. Interface mobile de mes livraisons affectees',
     "L'ecran my_orders_screen affiche les livraisons affectees au livreur connecte, organisees par statut. Il permet de naviguer vers le detail de chaque livraison pour declencher une mise a jour de statut ou demarrer une livraison active."),
    (108, '10.3. Interface mobile de mise a jour du statut',
     "L'ecran delivery_details_screen permet de changer le statut d'une livraison (EN_LIVRAISON, LIVRE, REPORTE, RETOUR, DEPOT, TENTATIVE). Pour le statut REPORTE, un selecteur de date de replanification est presente. Le montant encaisse est automatiquement enregistre lors du passage au statut LIVRE pour les paiements en especes."),
    (109, '10.4. Interface mobile de carte et itineraire optimise',
     "La carte interactive integre l'optimisation de tournee (algorithme du voisin le plus proche + Haversine) et les ETA calcules via le service OSRM. Le livreur peut visualiser l'ordre optimal de ses livraisons sur la carte et demarrer la navigation vers chaque arret."),
    (110, '10.5. Interface mobile de statistiques livreur',
     "L'ecran livreur_stats_screen consomme GET /api/livreur/stats et affiche les livraisons du jour, le taux de reussite, le montant total encaisse, la tendance des 7 derniers jours (sparkline) et la repartition par type de statut. Le livreur peut aussi soumettre sa caisse au depot via POST /api/livreur/cashbox/remettre."),
    (111, '10.6. Interface mobile de missions de transit (livreur de transit)',
     "L'ecran transit_home_screen liste les missions de transit inter-depots assignees au livreur. L'ecran transit_mission_details_screen presente les articles attendus avec leur etat de reception. L'ecran transit_barcode_scanner_screen active la camera via le plugin mobile_scanner pour scanner les codes-barres des articles recus, avec un retour sonore a chaque scan reussi."),
    (112, '10.7. Interface mobile de suivi commande client (TrackingStateCard)',
     "La carte adaptative TrackingStateCard affiche l'etat de la commande : AT_DEPOT (commande en depot), IN_DELIVERY_QUEUE (en cours, bouton d'appel livreur), HEADING_TO_YOU (livreur en route avec carte live LiveDeliveryMapSheet, ETA en minutes, fraicheur de la position GPS), ou TERMINAL (livre ou retourne). En etat HEADING_TO_YOU, le client peut appeler ou envoyer un SMS au livreur directement depuis l'ecran."),
    (113, '10.8. Interface mobile de reclamation client',
     "Le module de reclamations client mobile permet de creer une reclamation (POST /api/reclamations), de consulter la liste et le detail, et d'echanger des messages en temps reel avec le confirmateur via SignalR. Les photos d'incident peuvent etre jointes a la reclamation avec support des formats HEIC et HEIF (iPhone)."),
]
for fig_num, heading, desc in real4:
    add_h2(doc, heading)
    add_body(doc, desc)
    add_fig(doc, fig_num, heading.split('. ', 1)[1])

add_h1(doc, '11. Tests et validation')
add_body(doc, "Les tests du Sprint 4 valident le bon fonctionnement de l'application mobile dans les differents contextes d'usage : livreur standard, livreur de transit et client. Ils portent sur la coherence des donnees de livraison, les transitions de statut, la precision de la localisation GPS, l'optimisation de tournee et le comportement en temps reel.")
add_tab_label(doc, 30, 'Tests fonctionnels du Sprint 4')
doc.add_paragraph()

test4_table = doc.add_table(rows=1, cols=4)
test4_table.style = 'Table Grid'
for i, h in enumerate(['Test', 'Objectif', 'Resultat attendu', 'Statut']):
    test4_table.rows[0].cells[i].text = h

tests4 = [
    ('Authentification livreur', 'Verifier l\'acces aux ecrans proteges.', 'Seul un compte livreur authentifie accede au module. Le role est correctement detecte.', 'A valider par capture'),
    ('Livraisons disponibles', 'Verifier le filtrage par zone geographique.', 'Seules les livraisons du gouvernorat/delegation du livreur apparaissent.', 'A valider par capture'),
    ('Affectation livraison', 'Verifier la prise en charge d\'une livraison.', 'Une entree F_LIVRAISON est creee pour le livreur. La livraison disparait de la liste des disponibles.', 'A valider par capture'),
    ('Conflit d\'affectation', 'Verifier le cas d\'une livraison deja prise.', 'Le systeme retourne un message clair. La livraison n\'est pas doublee.', 'A valider par capture'),
    ('Mise a jour statut', 'Verifier les statuts LIVRE, REPORTE, RETOUR, DEPOT, TENTATIVE.', 'Le statut est sauvegarde avec les bonnes dates. Le montant est encaisse pour les paiements COD.', 'A valider par capture'),
    ('Diffusion GPS temps reel', 'Verifier la reception de la position livreur par le client.', 'La carte live affiche la position du livreur avec une fraicheur inferieure a 30 secondes.', 'A valider par capture'),
    ('File d\'attente GPS hors ligne', 'Verifier l\'envoi differe des positions en cas de perte de reseau.', 'Les positions sont envoyees en lot via /ping-batch a la reconnexion.', 'A valider par capture'),
    ('Optimisation de tournee', 'Verifier l\'ordre optimal calcule par l\'algorithme.', 'L\'ordre affiche minimise la distance totale. Les ETA OSRM sont coherents.', 'A valider par capture'),
    ('Scan code-barres transit', 'Verifier la validation par scan d\'un article de transit.', 'L\'article est marque comme recu dans F_TRANSFERTS. Un retour sonore est emis.', 'A valider par capture'),
    ('TrackingStateCard etats', 'Verifier l\'affichage adaptatif des 4 etats de suivi client.', 'Chaque etat affiche les bonnes informations et les bons boutons d\'action.', 'A valider par capture'),
    ('Reclamation mobile avec photo', 'Verifier l\'envoi d\'une reclamation avec photo HEIC.', 'La photo est acceptee et stockee. La reclamation apparait cote confirmateur.', 'A valider par capture'),
    ('Chat reclamation temps reel', 'Verifier la reception instantanee des messages.', 'Les messages arrivent sans rechargement via SignalR.', 'A valider par capture'),
    ('Responsive mobile', 'Verifier l\'affichage sur differentes tailles d\'ecran.', 'Les cartes, boutons et listes restent lisibles et fonctionnels.', 'A valider par capture'),
]
for t in tests4:
    row = test4_table.add_row().cells
    for i, v in enumerate(t):
        row[i].text = v

doc.add_paragraph()

add_h1(doc, '12. Conclusion du sprint 4')
add_body(doc, "Le Sprint 4 complete le workflow de la solution e-commerce et logistique par un suivi mobile complet des livraisons. Il apporte au livreur standard les fonctionnalites necessaires pour consulter les livraisons disponibles dans sa zone, prendre en charge une livraison, mettre a jour son statut et optimiser sa tournee grace a l'algorithme du voisin le plus proche et au service OSRM. La diffusion GPS en temps reel via SignalR permet au client de suivre la progression de son livreur sur une carte interactive depuis l'application mobile Flutter.")
add_body(doc, "Le livreur de transit dispose d'un espace dedie pour gerer les missions inter-depots et valider la reception des articles par scan de codes-barres. Le client mobile beneficie d'un suivi adaptatif (TrackingStateCard), d'un module de reclamations avec chat en temps reel et du support des photos HEIC/HEIF. Ce sprint renforce la continuite entre la Web API ASP.NET Core, la base SQL Server et l'application mobile Flutter, et conclut le perimetre fonctionnel du projet.")

out_path = '/home/user/PROJET-PFE/Rapport/Chapitres_5_6_Corriges.docx'
doc.save(out_path)
print('Document saved: ' + out_path)
