#!/usr/bin/env python3
"""
Build script: Generate PlantUML diagrams for Sprints 2-5 and insert into Word document.
"""

import os
import subprocess
import sys
from pathlib import Path

# ─── CONFIGURATION ────────────────────────────────────────────────────────────
DIAGRAMS_DIR = Path("/tmp/diagrams")
PLANTUML_JAR = "/tmp/plantuml.jar"
DOC_PATH = Path("/home/user/PROJET-PFE/Rapport/Rapport_PFE_Final_Modifie.docx")

SKINPARAM = """
skinparam backgroundColor white
skinparam handwritten false
skinparam sequenceMessageAlign center
skinparam responseMessageBelowArrow true
skinparam ParticipantPadding 20
skinparam BoxPadding 10
skinparam roundCorner 0
skinparam shadowing false
skinparam DefaultFontName Arial
skinparam DefaultFontSize 12
skinparam TitleFontStyle bold
skinparam TitleFontSize 14
"""

# ─── PUML DEFINITIONS ─────────────────────────────────────────────────────────
PUML_FILES = {}

PUML_FILES["s2_usecase.puml"] = f"""@startuml
title Diagramme de cas d'utilisation - Sprint 2
{SKINPARAM}
left to right direction

actor "Visiteur" as V
actor "Client B2C" as CB2C
actor "Client B2B" as CB2B

rectangle "Système e-commerce" {{
  usecase "Consulter le catalogue" as UC1
  usecase "Rechercher un produit" as UC2
  usecase "Filtrer les produits" as UC3
  usecase "Consulter le détail produit" as UC4
  usecase "Comparer des produits" as UC5
  usecase "Gérer ses favoris" as UC6
  usecase "Gérer le panier" as UC7
  usecase "Gérer ses adresses" as UC8
  usecase "Passer une commande" as UC9
  usecase "Payer en ligne (Konnect)" as UC10
  usecase "Checkout invité" as UC11
  usecase "Suivre une commande" as UC12
  usecase "Demander un devis B2B" as UC13
}}

V --> UC1
V --> UC2
V --> UC3
V --> UC4
V --> UC5
V --> UC11

CB2C --|> V
CB2C --> UC6
CB2C --> UC7
CB2C --> UC8
CB2C --> UC9
CB2C --> UC10
CB2C --> UC12

CB2B --|> CB2C
CB2B --> UC13

UC2 .> UC1 : <<include>>
UC3 .> UC1 : <<include>>
UC10 .> UC9 : <<include>>
@enduml
"""

PUML_FILES["s2_seq_catalogue.puml"] = f"""@startuml
title Diagramme de séquence - Consultation du catalogue
{SKINPARAM}
participant "Visiteur" as V
participant "Interface React" as R
participant "ArticlesController" as C
participant "ArticlesService" as S
participant "SQL Server" as DB

V -> R : Accéder à la page catalogue
R -> C : GET /api/articles?filters
C -> S : GetArticlesFiltres(filtres)
S -> DB : SELECT articles + prix + stock
DB --> S : Liste des articles
S --> C : ArticleDto[]
C --> R : 200 OK + articles paginés
R --> V : Afficher les articles

note over R : React Query met en cache\\nles résultats 5 minutes

@enduml
"""

PUML_FILES["s2_seq_panier.puml"] = f"""@startuml
title Diagramme de séquence - Gestion du panier
{SKINPARAM}
participant "Client" as C
participant "Interface React" as R
participant "CartController" as CC
participant "CartService" as CS
participant "SQL Server" as DB

C -> R : Cliquer "Ajouter au panier"
R -> CC : POST /api/cart/items\\n{{articleId, quantite}}
CC -> CS : AjouterArticle(userId, articleId, qte)
CS -> DB : SELECT stock disponible
DB --> CS : Quantite en stock
CS -> DB : UPSERT F_PANIER_LIGNE
DB --> CS : Ligne sauvegardée
CS --> CC : PanierDto mis à jour
CC --> R : 200 OK + panier
R --> C : Mettre à jour le panier (icône + count)

note over R : Zustand met à jour\\nle state global du panier

@enduml
"""

PUML_FILES["s2_seq_commande.puml"] = f"""@startuml
title Diagramme de séquence - Passer une commande et paiement Konnect
{SKINPARAM}
participant "Client" as C
participant "Interface React" as R
participant "OrdersController" as OC
participant "BonCommandeService" as BCS
participant "Konnect API" as K
participant "SQL Server" as DB

C -> R : Valider le panier (checkout)
R -> OC : POST /api/orders\\n{{adresse, modePaiement, lignes}}
OC -> BCS : CreerBonCommande(dto)
BCS -> DB : INSERT F_DOCENTETE (DO_Type=0)
BCS -> DB : INSERT F_DOCLIGNE (lignes)
DB --> BCS : BC créé (piece)

alt Paiement Konnect
  BCS -> K : POST /api/payments\\n{{montant, reference}}
  K --> BCS : paymentUrl
  BCS --> OC : BC + paymentUrl
  OC --> R : 201 + redirectUrl
  R --> C : Rediriger vers Konnect
  C -> K : Payer en ligne
  K -> OC : Webhook /api/payments/callback\\n{{status: paid}}
  OC -> DB : UPDATE DO_Statut = CONFIRME
  OC --> R : Confirmation paiement
  R --> C : Afficher confirmation
else Paiement à la livraison
  BCS --> OC : BC créé
  OC --> R : 201 Created
  R --> C : Afficher confirmation commande
end

@enduml
"""

PUML_FILES["s2_seq_devis.puml"] = f"""@startuml
title Diagramme de séquence - Demande de devis B2B
{SKINPARAM}
participant "Client B2B" as C
participant "Interface React" as R
participant "DevisController" as DC
participant "DevisService" as DS
participant "SQL Server" as DB

C -> R : Cliquer "Demander un devis" depuis le panier
R -> DC : POST /api/devis\\n{{lignes, clientId, note}}
DC -> DS : CreerDevis(dto)
DS -> DB : INSERT F_DEVIS_ENTETE\\n(statut=EN_ATTENTE)
DS -> DB : INSERT F_DEVIS_LIGNE (articles)
DB --> DS : Devis créé
DS --> DC : DevisDto
DC --> R : 201 Created + devisId
R --> C : "Votre devis a été soumis"

note over DB : Notifie l'équipe commerciale\\nvia SignalR ou email

@enduml
"""

PUML_FILES["s3_usecase.puml"] = f"""@startuml
title Diagramme de cas d'utilisation - Sprint 3
{SKINPARAM}
left to right direction

actor "Administrateur" as A
actor "Vendeur" as V
actor "Client B2B" as CB2B
actor "Sage X3" as SX3

rectangle "Gestion commerciale" {{
  usecase "Gérer les produits" as UC1
  usecase "Gérer les images (Cloudinary)" as UC2
  usecase "Synchroniser Sage X3" as UC3
  usecase "Gérer les stocks" as UC4
  usecase "Gérer les dépôts" as UC5
  usecase "Gérer les clients B2B" as UC6
  usecase "Gérer les remises B2B" as UC7
  usecase "Configurer la homepage" as UC8
  usecase "Gérer les paramètres" as UC9
  usecase "Consulter les indicateurs" as UC10
  usecase "Créer une commande client" as UC11
  usecase "Consulter le catalogue vendeur" as UC12
  usecase "Soumettre un devis" as UC13
  usecase "Valider / Négocier un devis" as UC14
  usecase "Convertir devis en commande" as UC15
}}

A --> UC1
A --> UC2
A --> UC3
A --> UC4
A --> UC5
A --> UC6
A --> UC7
A --> UC8
A --> UC9
A --> UC10
A --> UC14

V --> UC11
V --> UC12

CB2B --> UC13
CB2B --> UC15

SX3 --> UC3 : <<provide>>

UC2 .> UC1 : <<include>>
UC4 .> UC3 : <<include>>
UC15 .> UC13 : <<extend>>
@enduml
"""

PUML_FILES["s3_seq_sage.puml"] = f"""@startuml
title Diagramme de séquence - Synchronisation Sage X3
{SKINPARAM}
participant "Administrateur" as A
participant "Interface React" as R
participant "SyncController" as SC
participant "SageX3Service" as SS
participant "Sage X3 API" as SX
participant "SQL Server" as DB

A -> R : Cliquer "Synchroniser"
R -> SC : POST /api/articles/sync-sage
SC -> SS : SynchroniserArticles()
SS -> SX : GET /api/articles (Sage X3)
SX --> SS : Articles + prix + stocks
SS -> DB : UPSERT F_ARTICLE (code, désignation, prix)
SS -> DB : UPSERT F_TARIF (prix par catalogue)
SS -> DB : UPSERT F_STOCK (quantités par dépôt)
DB --> SS : Rapport (créés/mis à jour/erreurs)
SS --> SC : SyncResultDto
SC --> R : 200 OK + {{crées: n, mis à jour: m}}
R --> A : Afficher rapport de synchronisation

note over SS : Hangfire peut planifier\\nla synchronisation automatique

@enduml
"""

PUML_FILES["s3_seq_devis_workflow.puml"] = f"""@startuml
title Diagramme de séquence - Workflow devis B2B
{SKINPARAM}
participant "Client B2B" as C
participant "Interface React" as R
participant "DevisController" as DC
participant "DevisService" as DS
participant "Confirmateur" as CF
participant "SQL Server" as DB

C -> R : Soumettre devis depuis panier
R -> DC : POST /api/devis
DC -> DS : CreerDevis(dto)
DS -> DB : INSERT F_DEVIS_ENTETE (EN_ATTENTE)
DS -> DB : INSERT F_DEVIS_LIGNE
DB --> DS : Devis créé
DS --> DC : DevisDto
DC --> R : 201 Created

note over CF : Notifié par email / SignalR

CF -> DC : PUT /api/admin/devis/{{id}}/negocier\\n{{nouveauPrix, commentaire}}
DC -> DS : NegocierDevis(id, dto)
DS -> DB : UPDATE statut=EN_NEGOCIATION\\nINSERT F_DEVIS_EVENT
DB --> DS : OK
DS --> DC : DevisDto négocié
DC --> CF : 200 OK

C -> DC : PUT /api/devis/{{id}}/accepter-negociation
DC -> DS : AccepterNegociation(id)
DS -> DB : UPDATE statut=ACCEPTE
DB --> DS : OK
DS --> DC : DevisDto
DC --> C : 200 OK

C -> DC : POST /api/devis/{{id}}/convertir
DC -> DS : ConvertirEnCommande(id)
DS -> DB : INSERT F_DOCENTETE (BC)\\nINSERT F_DOCLIGNE\\nUPDATE statut=CONVERTI
DB --> DS : BC créé
DS --> DC : BonCommandeDto
DC --> C : 201 + BC reference

@enduml
"""

PUML_FILES["s3_seq_vendeur.puml"] = f"""@startuml
title Diagramme de séquence - Création de commande vendeur
{SKINPARAM}
participant "Vendeur" as V
participant "Interface React" as R
participant "VendeurController" as VC
participant "VendeurOrdersService" as VS
participant "RemisesService" as RS
participant "SQL Server" as DB

V -> R : Sélectionner un client B2B
R -> VC : GET /api/vendeur/clients?search=...
VC --> R : Liste clients

V -> R : Ajouter articles au panier vendeur
R -> RS : GET /api/admin/b2b/remises/{{clientId}}
RS -> DB : SELECT remises par client
DB --> RS : RemiseDto[]

V -> R : Valider la commande
R -> VC : POST /api/vendeur/commandes\\n{{clientId, lignes, remises}}
VC -> VS : CreerCommandeVendeur(dto)
VS -> DB : INSERT F_DOCENTETE\\nINSERT F_DOCLIGNE\\n(avec remises B2B appliquées)
DB --> VS : BC créé
VS --> VC : BonCommandeDto
VC --> R : 201 Created + BC reference
R --> V : Afficher confirmation commande

@enduml
"""

PUML_FILES["s4_usecase.puml"] = f"""@startuml
title Diagramme de cas d'utilisation - Sprint 4
{SKINPARAM}
left to right direction

actor "Confirmatrice" as CF
actor "Livreur" as LV
actor "Livreur Transit" as LT
actor "Client" as CL
actor "Superviseur" as SV

rectangle "Logistique et réclamations" {{
  usecase "Consulter commandes en attente" as UC1
  usecase "Modifier statut BC" as UC2
  usecase "Transformer BC en BL" as UC3
  usecase "Exporter BL vers Sage X3" as UC4
  usecase "Gérer les réclamations" as UC5
  usecase "Consulter pool livraisons" as UC6
  usecase "Prendre en charge livraison" as UC7
  usecase "Démarrer diffusion GPS" as UC8
  usecase "Mettre à jour statut livraison" as UC9
  usecase "Gérer la caisse COD" as UC10
  usecase "Optimiser la tournée" as UC11
  usecase "Demander correction adresse" as UC12
  usecase "Gérer missions transit" as UC13
  usecase "Scanner codes-barres" as UC14
  usecase "Créer une réclamation" as UC15
  usecase "Suivre commande (TrackingCard)" as UC16
  usecase "Superviser livreurs actifs" as UC17
}}

CF --> UC1
CF --> UC2
CF --> UC3
CF --> UC4
CF --> UC5

LV --> UC6
LV --> UC7
LV --> UC8
LV --> UC9
LV --> UC10
LV --> UC11
LV --> UC12

LT --> UC13
LT --> UC14

CL --> UC15
CL --> UC16

SV --> UC17

UC3 .> UC1 : <<include>>
UC8 .> UC7 : <<include>>
UC14 .> UC13 : <<include>>
@enduml
"""

PUML_FILES["s4_seq_bctbl.puml"] = f"""@startuml
title Diagramme de séquence - Transformation BC en BL
{SKINPARAM}
participant "Confirmatrice" as CF
participant "Interface React" as R
participant "ConfirmateurController" as CC
participant "BonCommandeService" as BCS
participant "SQL Server" as DB

CF -> R : Cliquer "Transformer en BL"
R -> CC : POST /api/confirmateur/commandes/{{piece}}/transform-to-bl
CC -> BCS : TransformerEnBL(piece)
BCS -> DB : SELECT F_DOCENTETE WHERE DO_Piece=piece

alt BC introuvable
  BCS --> CC : Erreur 404
  CC --> R : 404 Not Found
  R --> CF : "Commande introuvable"
else BL existe déjà
  BCS -> DB : SELECT BL existant pour le BC
  DB --> BCS : BL reference
  BCS --> CC : BL existant (référence)
  CC --> R : 200 + blReference
  R --> CF : "BL déjà créé : " + référence
else Création BL
  BCS -> DB : BEGIN TRANSACTION
  BCS -> DB : INSERT F_DOCENTETE (DO_Type=1, BL)
  BCS -> DB : INSERT F_DOCLIGNE (lignes du BC)
  BCS -> DB : UPDATE F_DOCENTETE SET BL_Ref (lien BC->BL)
  DB --> BCS : COMMIT
  BCS --> CC : BlPiece
  CC --> R : 201 Created + blPiece
  R --> CF : "BL créé : " + blPiece
end

@enduml
"""

PUML_FILES["s4_seq_gps.puml"] = f"""@startuml
title Diagramme de séquence - Livraison COD et diffusion GPS
{SKINPARAM}
participant "Livreur" as L
participant "App Flutter" as F
participant "LivreurController" as LC
participant "LocationService" as LS
participant "SignalR Hub" as SH
participant "Client" as CL
participant "SQL Server" as DB

L -> F : Démarrer la livraison
F -> LC : POST /api/livreur/orders/{{piece}}/start-heading
LC -> DB : UPDATE F_LIVRAISON statut=EN_COURS
LC -> DB : INSERT F_LIVREUR_POSITION (initial)
LC --> F : 200 OK

loop Toutes les 10 secondes
  F -> F : Geolocator.getPosition()
  F -> LC : POST /api/livreur/location/ping\\n{{lat, lng, timestamp}}
  LC -> DB : UPSERT F_LIVREUR_POSITION
  LC -> SH : LocationUpdate(userId, lat, lng)
  SH -> CL : LocationUpdate event (WebSocket)
  CL -> CL : Mettre à jour carte live
end

note over F : Si réseau indisponible :\\nstockage local puis\\nPOST /ping-batch

L -> F : Livraison effectuée
F -> LC : PUT /api/livreur/orders/{{piece}}/status\\n{{statut: LIVRE, montantCOD}}
LC -> DB : UPDATE F_LIVRAISON\\nstatut=LIVRE, Encaisse=true\\nMontantEncaisse=montant
LC -> F : POST /stop-heading
LC --> F : 200 OK
F --> L : Confirmation livraison

@enduml
"""

PUML_FILES["s4_seq_reclamation.puml"] = f"""@startuml
title Diagramme de séquence - Réclamation client et SignalR
{SKINPARAM}
participant "Client" as C
participant "App Flutter" as F
participant "ReclamationController" as RC
participant "ReclamationHub (SignalR)" as SH
participant "Confirmatrice" as CF
participant "SQL Server" as DB

C -> F : Créer une réclamation (motif + photos)
F -> RC : POST /api/reclamations\\n{{motif, commandeId, photos}}
RC -> DB : INSERT F_RECLAMATION (EN_ENVOI)
RC -> SH : Diffuser NouveauCas\\nau groupe "confirmateurs"
SH -> CF : Événement NouveauCas reçu
CF -> CF : Réclamation apparaît\\nsans rechargement

CF -> RC : POST /api/confirmateur/reclamations/{{id}}/reprendre
RC -> DB : UPDATE F_RECLAMATION SET confirmateurId
RC -> SH : Diffuser CasAssigne
SH -> C : Événement CasAssigne reçu

CF -> RC : POST /api/confirmateur/reclamations/{{id}}/messages\\n{{texte}}
RC -> DB : INSERT F_RECLAMATION_MESSAGE
RC -> SH : Diffuser NouveauMessage\\nau groupe "client-{{userId}}"
SH -> C : NouveauMessage reçu en temps réel

C -> RC : POST /api/reclamations/{{id}}/messages\\n{{réponse}}
RC -> DB : INSERT F_RECLAMATION_MESSAGE
RC -> SH : NouveauMessage au confirmateur

CF -> RC : PUT /api/confirmateur/reclamations/{{id}}/status\\n{{statut: CLOTUREE}}
RC -> DB : UPDATE statut=CLOTUREE
RC -> SH : StatutCasChange -> client
SH -> C : Notification clôture

@enduml
"""

PUML_FILES["s4_seq_transit.puml"] = f"""@startuml
title Diagramme de séquence - Scan codes-barres transit inter-dépôts
{SKINPARAM}
participant "Livreur Transit" as LT
participant "App Flutter" as F
participant "TransitController" as TC
participant "TransitService" as TS
participant "SQL Server" as DB

LT -> F : Ouvrir la mission de transit
F -> TC : GET /api/transit/my-missions/{{id}}
TC -> DB : SELECT F_TRANSFERTS lignes
DB --> TC : Articles attendus
TC --> F : MissionDto (articles + états)
F --> LT : Afficher liste des articles

LT -> F : Activer le scanner
F -> F : Démarrer mobile_scanner

loop Pour chaque article
  LT -> F : Scanner le code-barres
  F -> TC : PUT /api/transit/my-missions/{{id}}/scan\\n{{barcode}}
  TC -> TS : ValiderArticle(missionId, barcode)
  TS -> DB : SELECT F_TRANSFERTS WHERE barcode

  alt Article appartient à la mission
    TS -> DB : UPDATE F_TRANSFERTS SET recu=true
    DB --> TS : OK
    TS --> TC : ArticleValidéDto
    TC --> F : 200 OK
    F -> F : Vibration + son de confirmation
    F --> LT : Article marqué reçu (vert)
  else Article inconnu ou autre mission
    TS --> TC : Erreur 400
    TC --> F : 400 Bad Request
    F --> LT : Son d'erreur + message
  end
end

F -> TC : PUT /api/transit/my-missions/{{id}}/terminer
TC -> TS : TerminerMission(id)
TS -> DB : UPDATE F_TRANSFERT_ENTETE statut=TERMINE
TC --> F : 200 OK
F --> LT : Mission clôturée

@enduml
"""

PUML_FILES["s5_usecase.puml"] = f"""@startuml
title Diagramme de cas d'utilisation - Sprint 5
{SKINPARAM}
left to right direction

actor "Administrateur" as A
actor "n8n Workflow" as N8N

rectangle "Pilotage et aide à la décision" {{
  usecase "Consulter tableau de bord global" as UC1
  usecase "Analyser les ventes" as UC2
  usecase "Analyser les commandes" as UC3
  usecase "Analyser les produits" as UC4
  usecase "Analyser les stocks" as UC5
  usecase "Analyser la logistique" as UC6
  usecase "Analyser les livreurs" as UC7
  usecase "Analyser les clients" as UC8
  usecase "Suivre les réclamations" as UC9
  usecase "Suivre les confirmatrices" as UC10
  usecase "Suivre la sync Sage X3" as UC11
  usecase "Consulter les insights" as UC12
  usecase "Générer export PDF" as UC13
  usecase "Générer export Excel" as UC14
  usecase "Interagir avec le chatbot" as UC15
  usecase "Consulter conversations chatbot" as UC16
}}

A --> UC1
A --> UC2
A --> UC3
A --> UC4
A --> UC5
A --> UC6
A --> UC7
A --> UC8
A --> UC9
A --> UC10
A --> UC11
A --> UC12
A --> UC13
A --> UC14
A --> UC15
A --> UC16

N8N --> UC15 : <<provide>>

UC2 .> UC1 : <<include>>
UC3 .> UC1 : <<include>>
UC16 .> UC15 : <<include>>
@enduml
"""

PUML_FILES["s5_seq_chatbot.puml"] = f"""@startuml
title Diagramme de séquence - Chatbot administrateur (n8n)
{SKINPARAM}
participant "Administrateur" as A
participant "Interface React" as R
participant "ChatbotController" as CC
participant "n8n Workflow" as N8N
participant "Web API (endpoints)" as API
participant "SQL Server" as DB

A -> R : Saisir une question\\n"Quels livreurs ont le plus de retours ?"
R -> CC : POST /api/chatbot/ask\\n{{question, conversationId}}
CC -> N8N : POST webhook n8n\\n{{question, contexte, historique}}

N8N -> API : GET /api/dashboard/livreurs?filter=retours
API -> DB : SELECT stats livreurs
DB --> API : Données livreurs
API --> N8N : JSON résultats

N8N -> N8N : Traitement LLM (Groq)\\nFormater réponse en français

N8N --> CC : {{réponse, données, insights}}
CC -> DB : INSERT F_CHATBOT_MESSAGE\\n(question + réponse)
CC --> R : 200 OK + réponse formatée
R --> A : Afficher réponse conversationnelle

note over N8N : L'historique de la conversation\\nest injecté dans chaque requête\\npour le contexte

@enduml
"""

PUML_FILES["s5_seq_export.puml"] = f"""@startuml
title Diagramme de séquence - Export Excel et PDF
{SKINPARAM}
participant "Administrateur" as A
participant "Interface React" as R
participant "ExportController" as EC
participant "ExportService" as ES
participant "SQL Server" as DB

A -> R : Sélectionner type export\\n(Commandes, Livraisons, Réclamations)
A -> R : Choisir période et filtres
R -> EC : GET /api/admin/exports/excel\\n?type=commandes&debut=...&fin=...

EC -> ES : GenererExcel(type, filtres)
ES -> DB : SELECT données selon type et période
DB --> ES : Données brutes
ES -> ES : Construire classeur Excel\\n(colonnes, formatage, totaux)
ES --> EC : Stream Excel
EC --> R : 200 OK + fichier .xlsx (Content-Disposition)
R --> A : Téléchargement automatique

A -> R : Générer PDF
R -> EC : GET /api/admin/exports/pdf\\n?type=livraisons&...
EC -> ES : GenererPDF(type, filtres)
ES -> DB : SELECT données
DB --> ES : Données
ES -> ES : Construire PDF\\n(en-tête, tableaux, graphiques)
ES --> EC : Stream PDF
EC --> R : 200 OK + fichier .pdf
R --> A : Téléchargement automatique

@enduml
"""

# ─── STEP 1: Write all .puml files ────────────────────────────────────────────
def step1_write_puml():
    DIAGRAMS_DIR.mkdir(parents=True, exist_ok=True)
    for filename, content in PUML_FILES.items():
        path = DIAGRAMS_DIR / filename
        path.write_text(content, encoding="utf-8")
        print(f"  Written: {path}")
    print(f"  Total: {len(PUML_FILES)} .puml files")

# ─── STEP 2: Generate PNGs ────────────────────────────────────────────────────
def step2_generate_pngs():
    puml_files = sorted(DIAGRAMS_DIR.glob("*.puml"))
    errors = []
    for pf in puml_files:
        result = subprocess.run(
            ["java", "-jar", PLANTUML_JAR, "-tpng", str(pf)],
            capture_output=True, text=True
        )
        if result.returncode != 0:
            errors.append(f"ERROR generating {pf.name}: {result.stderr}")
        else:
            png = pf.with_suffix(".png")
            size = png.stat().st_size if png.exists() else 0
            print(f"  Generated: {png.name} ({size} bytes)")
    if errors:
        for e in errors:
            print(e)
    png_count = len(list(DIAGRAMS_DIR.glob("*.png")))
    print(f"  Total PNG generated: {png_count}")
    return png_count

# ─── STEP 3: Modify the Word document ─────────────────────────────────────────

# Data for "Modifications" column
MODIFICATIONS = {
    # Sprint 2
    "US2.1":  "API: GET /api/homepage\nReact: HomePage.tsx\nFlutter: —",
    "US2.2":  "API: GET /api/articles (filtres, pagination)\nReact: ArticlesPage.tsx, ArticleCard.tsx\nFlutter: —",
    "US2.3":  "API: GET /api/articles?search=\nReact: ArticlesFilterPanel.tsx\nFlutter: —",
    "US2.4":  "API: GET /api/articles/{id}\nReact: ArticleDetailsPage.tsx\nFlutter: —",
    "US2.5":  "API: GET /api/articles/compare\nReact: ComparePage.tsx\nFlutter: —",
    "US2.6":  "API: GET/POST/DELETE /api/favorites\nReact: FavoritesPage.tsx, FavoriteToggleButton.tsx\nFlutter: —",
    "US2.7":  "API: GET/POST/PUT/DELETE /api/cart\nReact: CartPage.tsx (Zustand)\nFlutter: —",
    "US2.8":  "API: GET/POST/PUT/DELETE /api/addresses\nReact: AddressesPage.tsx\nFlutter: —",
    "US2.9":  "API: POST /api/orders (BC)\nReact: CheckoutPage.tsx\nFlutter: —",
    "US2.10": "API: POST /api/payments/konnect, webhook\nReact: CheckoutPage.tsx\nFlutter: —",
    "US2.11": "API: POST /api/orders (guest)\nReact: GuestCheckoutLocationSection.tsx\nFlutter: —",
    "US2.12": "API: GET /api/orders/{id}/tracking\nReact: OrderDetailsPage.tsx\nFlutter: client_order_tracking_screen.dart",
    "US2.13": "API: GET /api/pages/{slug}\nReact: StaticPage.tsx\nFlutter: —",
    # Sprint 3
    "US3.1":  "API: CRUD /api/admin/articles\nReact: AdminProductsPage.tsx\nFlutter: —",
    "US3.2":  "API: POST /api/admin/articles/{id}/images (Cloudinary)\nReact: AdminProductImagesPage.tsx\nFlutter: —",
    "US3.3":  "API: POST /api/articles/sync-sage\nReact: SyncSagePage.tsx\nFlutter: —",
    "US3.4":  "API: POST /api/articles/sync-catalogues\nReact: SyncSagePage.tsx\nFlutter: —",
    "US3.5":  "API: POST /api/admin/depots/sync-sage\nReact: DepotsPage.tsx\nFlutter: —",
    "US3.6":  "API: POST /api/admin/stocks/sync-sage\nReact: StocksPage.tsx\nFlutter: —",
    "US3.7":  "API: CRUD /api/admin/clients\nReact: AdminB2BClientsPage.tsx\nFlutter: —",
    "US3.8":  "API: CRUD /api/admin/b2b/remises\nReact: AdminB2BClientsPage.tsx\nFlutter: —",
    "US3.9":  "API: GET /api/vendeur/catalogue\nReact: VendeurCataloguePage.tsx\nFlutter: —",
    "US3.10": "API: POST /api/vendeur/commandes\nReact: VendeurOrderCreatePage.tsx\nFlutter: —",
    "US3.11": "API: POST /api/devis\nReact: B2BQuotesPage.tsx\nFlutter: —",
    "US3.12": "API: GET/PUT /api/admin/devis, /api/confirmateur/devis\nReact: ConfirmateurDevisPage.tsx\nFlutter: ConfirmateurDevisDetailsPage (Flutter Web)",
    "US3.13": "API: CRUD /api/admin/depots, /api/zones\nReact: ZonesPage.tsx\nFlutter: —",
    "US3.14": "API: PUT /api/admin/homepage\nReact: HomepageAdminPage.tsx\nFlutter: —",
    "US3.15": "API: PUT /api/admin/settings\nReact: SettingsPage.tsx\nFlutter: —",
    "US3.16": "API: GET /api/admin/dashboard\nReact: DashboardPage.tsx\nFlutter: —",
    # Sprint 4
    "US4.1":  "API: GET /api/confirmateur/commandes\nReact: ConfirmateurOrdersPage.tsx\nFlutter: confirmatrice_orders_screen.dart",
    "US4.2":  "API: GET /api/confirmateur/commandes/{piece}\nReact: ConfirmateurOrderDetailsPage.tsx\nFlutter: confirmatrice_orders_screen.dart",
    "US4.3":  "API: PUT /api/confirmateur/commandes/{piece}/status\nReact: ConfirmateurOrderDetailsPage.tsx\nFlutter: confirmatrice_orders_screen.dart",
    "US4.4":  "API: POST /api/confirmateur/commandes/{piece}/transform-to-bl\nReact: ConfirmateurOrderDetailsPage.tsx\nFlutter: confirmatrice_orders_screen.dart",
    "US4.5":  "API: POST /api/confirmateur/bl/{piece}/export-sage\nReact: ConfirmateurBlPage.tsx\nFlutter: —",
    "US4.6":  "API: GET /api/confirmateur/devis\nReact: ConfirmateurDevisPage.tsx\nFlutter: —",
    "US4.7":  "API: GET /api/livreur/pool/disponibles\nReact: —\nFlutter: new_orders_screen.dart",
    "US4.8":  "API: POST /api/livreur/pool/{piece}/prendre\nReact: —\nFlutter: new_orders_screen.dart",
    "US4.9":  "API: POST /start-heading, /stop-heading, /ping, /ping-batch\nReact: —\nFlutter: LivreurLocationService.dart",
    "US4.10": "API: PUT /api/livreur/orders/{piece}/status\nReact: —\nFlutter: delivery_details_screen.dart",
    "US4.11": "API: POST /api/livreur/cashbox/remettre\nReact: —\nFlutter: livreur_stats_screen.dart",
    "US4.12": "API: POST /api/reclamations\nReact: ClientClaimsPage.tsx\nFlutter: client_create_claim_screen.dart",
    "US4.13": "API: ReclamationHub SignalR, /api/confirmateur/reclamations/*\nReact: ConfirmateurClaimsPage.tsx\nFlutter: —",
    "US4.14": "API: POST /api/livreur/requests\nReact: —\nFlutter: delivery_details_screen.dart",
    "US4.15": "API: GET /api/superviseur/livreurs, /alertes\nReact: SuperviseurPage.tsx\nFlutter: —",
    "US4.16": "API: /api/transit/my-missions/*\nReact: —\nFlutter: transit_home_screen.dart, transit_barcode_scanner_screen.dart",
    "US4.17": "API: POST /api/livreur/tournee/optimize (OSRM)\nReact: —\nFlutter: delivery_map_screen.dart",
    # Sprint 5
    "US5.1":  "API: GET /api/admin/dashboard\nReact: DashboardGlobalPage.tsx\nFlutter: —",
    "US5.2":  "API: GET /api/dashboard/ventes\nReact: DashboardVentesPage.tsx\nFlutter: —",
    "US5.3":  "API: GET /api/dashboard/commandes\nReact: DashboardCommandesPage.tsx\nFlutter: —",
    "US5.4":  "API: GET /api/dashboard/produits\nReact: DashboardProduitsPage.tsx\nFlutter: —",
    "US5.5":  "API: GET /api/dashboard/stocks\nReact: DashboardStocksPage.tsx\nFlutter: —",
    "US5.6":  "API: GET /api/dashboard/logistique\nReact: DashboardLogistiquePage.tsx\nFlutter: —",
    "US5.7":  "API: GET /api/dashboard/livreurs\nReact: DashboardLivreursPage.tsx\nFlutter: livreur_stats_screen.dart",
    "US5.8":  "API: GET /api/dashboard/clients\nReact: DashboardClientsPage.tsx\nFlutter: —",
    "US5.9":  "API: GET /api/dashboard/reclamations\nReact: DashboardReclamationsPage.tsx\nFlutter: —",
    "US5.10": "API: GET /api/dashboard/confirmateurs\nReact: DashboardConfirmateursPage.tsx\nFlutter: —",
    "US5.11": "API: GET /api/dashboard/sage-sync\nReact: DashboardSagePage.tsx\nFlutter: —",
    "US5.12": "API: GET /api/dashboard/insights\nReact: InsightsPage.tsx\nFlutter: —",
    "US5.13": "API: GET /api/admin/exports/pdf\nReact: ExportPage.tsx\nFlutter: —",
    "US5.14": "API: POST /api/chatbot/ask (n8n)\nReact: ChatbotPanel.tsx\nFlutter: —",
    "US5.15": "API: GET /api/chatbot/conversations\nReact: ChatbotHistoryPage.tsx\nFlutter: —",
    "US5.16": "API: README, Swagger\nReact: —\nFlutter: —",
}

# Placeholder → PNG mapping
PLACEHOLDER_MAP = [
    ("diagramme de cas d'utilisation du Sprint 2",      "s2_usecase.png"),
    ("diagramme de cas d'utilisation sprint 2",         "s2_usecase.png"),
    ("Consultation du catalogue",                        "s2_seq_catalogue.png"),
    ("consultation du catalogue",                        "s2_seq_catalogue.png"),
    ("Gestion du panier",                                "s2_seq_panier.png"),
    ("gestion du panier",                                "s2_seq_panier.png"),
    ("Création d'une commande client",                   "s2_seq_commande.png"),
    ("création d'une commande client",                   "s2_seq_commande.png"),
    ("commande et paiement",                             "s2_seq_commande.png"),
    ("Demande de devis B2B",                             "s2_seq_devis.png"),
    ("demande de devis B2B",                             "s2_seq_devis.png"),
    ("devis B2B du Sprint 2",                            "s2_seq_devis.png"),
    ("diagramme de cas d'utilisation du Sprint 3",      "s3_usecase.png"),
    ("diagramme de cas d'utilisation sprint 3",         "s3_usecase.png"),
    ("Synchronisation Sage X3",                          "s3_seq_sage.png"),
    ("synchronisation Sage X3",                          "s3_seq_sage.png"),
    ("synchronisation sage",                             "s3_seq_sage.png"),
    ("Workflow devis B2B",                               "s3_seq_devis_workflow.png"),
    ("workflow devis",                                   "s3_seq_devis_workflow.png"),
    ("Devis B2B",                                        "s3_seq_devis_workflow.png"),
    ("Commande vendeur",                                 "s3_seq_vendeur.png"),
    ("commande vendeur",                                 "s3_seq_vendeur.png"),
    ("diagramme de cas d'utilisation du Sprint 4",      "s4_usecase.png"),
    ("diagramme de cas d'utilisation sprint 4",         "s4_usecase.png"),
    ("Transformation BC en BL",                          "s4_seq_bctbl.png"),
    ("transformation BC en BL",                         "s4_seq_bctbl.png"),
    ("Livraison COD",                                    "s4_seq_gps.png"),
    ("livraison COD",                                    "s4_seq_gps.png"),
    ("diffusion GPS",                                    "s4_seq_gps.png"),
    ("Réclamation",                                      "s4_seq_reclamation.png"),
    ("réclamation",                                      "s4_seq_reclamation.png"),
    ("SignalR",                                          "s4_seq_reclamation.png"),
    ("Scan",                                             "s4_seq_transit.png"),
    ("scan codes",                                       "s4_seq_transit.png"),
    ("transit",                                          "s4_seq_transit.png"),
    ("diagramme de cas d'utilisation du Sprint 5",      "s5_usecase.png"),
    ("diagramme de cas d'utilisation sprint 5",         "s5_usecase.png"),
    ("Chatbot",                                          "s5_seq_chatbot.png"),
    ("chatbot",                                          "s5_seq_chatbot.png"),
    ("Export",                                           "s5_seq_export.png"),
    ("export",                                           "s5_seq_export.png"),
]

def step3_modify_document():
    from docx import Document
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from lxml import etree
    import copy

    print("\n  Opening document...")
    doc = Document(str(DOC_PATH))

    images_inserted = 0
    texts_replaced = 0

    # ── 3a. Replace name placeholders in all text ──────────────────────────
    print("  Replacing name placeholders in paragraphs...")
    for para in doc.paragraphs:
        for run in para.runs:
            if "Melek" in run.text and "Lahmar" not in run.text:
                run.text = run.text.replace("Melek", "Melek Lahmar")
                texts_replaced += 1
            if "Tawfik" in run.text and "Siala" not in run.text:
                run.text = run.text.replace("Tawfik", "Tawfik Siala")
                texts_replaced += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "Melek" in run.text and "Lahmar" not in run.text:
                            run.text = run.text.replace("Melek", "Melek Lahmar")
                            texts_replaced += 1
                        if "Tawfik" in run.text and "Siala" not in run.text:
                            run.text = run.text.replace("Tawfik", "Tawfik Siala")
                            texts_replaced += 1

    print(f"  Name replacements done: {texts_replaced}")

    # ── 3b. Insert images at placeholders ─────────────────────────────────
    print("  Scanning for image placeholders...")

    def get_para_text(para):
        return "".join(r.text for r in para.runs).strip()

    def para_matches_placeholder(text):
        text_lower = text.lower()
        for keyword, png_name in PLACEHOLDER_MAP:
            if keyword.lower() in text_lower and "[" in text:
                return png_name
        return None

    def replace_para_with_image(para, img_path):
        """Replace paragraph content with centered image."""
        # Clear all runs from the paragraph
        for child in list(para._p):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("r", "hyperlink", "bookmarkStart", "bookmarkEnd", "ins", "del"):
                para._p.remove(child)

        # Set paragraph alignment to center
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            para._p.insert(0, pPr)
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "center")

        # Add image via a temporary paragraph then move the run
        tmp_para = doc.add_paragraph()
        run = tmp_para.add_run()
        run.add_picture(str(img_path), width=Cm(15))

        # Get the run XML element and move it
        run_elem = tmp_para._p.find(qn("w:r"))
        if run_elem is not None:
            para._p.append(copy.deepcopy(run_elem))

        # Remove temporary paragraph
        tmp_para._p.getparent().remove(tmp_para._p)

    # Collect all body paragraphs (not in tables)
    for para in doc.paragraphs:
        text = get_para_text(para)
        if not text:
            continue
        png_name = para_matches_placeholder(text)
        if png_name:
            img_path = DIAGRAMS_DIR / png_name
            if img_path.exists():
                print(f"    Inserting {png_name} at: '{text[:60]}...'")
                replace_para_with_image(para, img_path)
                images_inserted += 1
            else:
                print(f"    WARNING: PNG not found: {img_path}")

    print(f"  Images inserted: {images_inserted}")

    # ── 3c. Update backlog tables ──────────────────────────────────────────
    print("  Updating backlog tables...")
    tables_updated = 0

    def get_cell_text(cell):
        return " ".join(p.text.strip() for p in cell.paragraphs).strip()

    def is_backlog_table(table):
        """Check if table has 'User Story' in its header row."""
        if len(table.rows) < 1:
            return False
        header_text = " ".join(get_cell_text(c) for c in table.rows[0].cells).lower()
        return "user story" in header_text

    def find_resp_col_index(table):
        """Find the column index of 'Resp.' in header row."""
        if len(table.rows) < 1:
            return -1
        for i, cell in enumerate(table.rows[0].cells):
            ct = get_cell_text(cell).lower()
            if "resp" in ct:
                return i
        return -1

    def already_has_modifications_col(table):
        """Check if 'Modifications' column already exists."""
        if len(table.rows) < 1:
            return False
        header_text = " ".join(get_cell_text(c) for c in table.rows[0].cells).lower()
        return "modification" in header_text

    def find_us_id_in_row(row):
        """Extract US ID like US2.1 from row cells."""
        for cell in row.cells:
            ct = get_cell_text(cell).strip()
            import re
            m = re.match(r"(US\d+\.\d+)", ct)
            if m:
                return m.group(1)
        return None

    def clone_cell_xml(source_cell):
        """Create a new tc element based on a source cell's formatting."""
        new_tc = OxmlElement("w:tc")

        # Copy tcPr (borders, margins, etc.) from source
        src_tcPr = source_cell._tc.find(qn("w:tcPr"))
        if src_tcPr is not None:
            new_tc.append(copy.deepcopy(src_tcPr))
        else:
            tcPr = OxmlElement("w:tcPr")
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), "2000")
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            new_tc.append(tcPr)

        return new_tc

    def add_text_to_tc(tc, text, font_size=18, bold=False):
        """Add a paragraph with text to a tc element."""
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "left")
        pPr.append(jc)
        p.append(pPr)

        lines = text.split("\n")
        for i, line in enumerate(lines):
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(font_size))
            rPr.append(sz)
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), str(font_size))
            rPr.append(szCs)
            if bold:
                b_elem = OxmlElement("w:b")
                rPr.append(b_elem)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = line
            r.append(t)
            p.append(r)
            # Add line break between lines (except last)
            if i < len(lines) - 1:
                r_br = OxmlElement("w:r")
                br = OxmlElement("w:br")
                r_br.append(br)
                p.append(r_br)
        tc.append(p)

    for table_idx, table in enumerate(doc.tables):
        if not is_backlog_table(table):
            continue
        if already_has_modifications_col(table):
            print(f"    Table {table_idx}: already has Modifications column, skipping.")
            continue

        resp_col_idx = find_resp_col_index(table)
        if resp_col_idx < 0:
            print(f"    Table {table_idx}: no Resp. column found, skipping.")
            continue

        print(f"    Updating table {table_idx} (resp col at index {resp_col_idx})...")
        tables_updated += 1

        # Iterate rows and insert new cell before resp_col_idx
        for row_idx, row in enumerate(table.rows):
            tr = row._tr
            cells = list(tr.findall(qn("w:tc")))

            if row_idx == 0:
                # Header row: insert "Modifications" header cell
                if resp_col_idx < len(cells):
                    ref_cell_xml = cells[resp_col_idx]
                    new_tc = clone_cell_xml(row.cells[resp_col_idx])
                    add_text_to_tc(new_tc, "Modifications", font_size=18, bold=True)
                    ref_cell_xml.addprevious(new_tc)
            else:
                # Data row: find US ID and insert modification text
                us_id = find_us_id_in_row(row)
                mod_text = MODIFICATIONS.get(us_id, "") if us_id else ""

                if resp_col_idx < len(cells):
                    ref_cell_xml = cells[resp_col_idx]
                    new_tc = clone_cell_xml(row.cells[min(resp_col_idx, len(row.cells)-1)])
                    if mod_text:
                        add_text_to_tc(new_tc, mod_text, font_size=16)
                    else:
                        add_text_to_tc(new_tc, "—", font_size=16)
                    ref_cell_xml.addprevious(new_tc)

    print(f"  Tables updated: {tables_updated}")

    # ── Save ──────────────────────────────────────────────────────────────
    print(f"\n  Saving document to {DOC_PATH}...")
    doc.save(str(DOC_PATH))
    print("  Document saved successfully.")

    return images_inserted, tables_updated

# ─── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    print("=" * 60)
    print("STEP 1: Writing .puml files")
    print("=" * 60)
    step1_write_puml()

    print("\n" + "=" * 60)
    print("STEP 2: Generating PNG diagrams")
    print("=" * 60)
    png_count = step2_generate_pngs()

    print("\n" + "=" * 60)
    print("STEP 3: Modifying Word document")
    print("=" * 60)
    images_inserted, tables_updated = step3_modify_document()

    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"  .puml files created : {len(PUML_FILES)}")
    print(f"  PNG files generated : {png_count}")
    print(f"  Images inserted     : {images_inserted}")
    print(f"  Tables updated      : {tables_updated}")
    doc_size = DOC_PATH.stat().st_size if DOC_PATH.exists() else 0
    print(f"  Output document     : {DOC_PATH} ({doc_size:,} bytes)")
    print("=" * 60)
    print("DONE")

if __name__ == "__main__":
    main()
