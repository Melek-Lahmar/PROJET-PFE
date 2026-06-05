#!/usr/bin/env python3
"""
Fusion finale : remplacement des chapitres 7 et 8 du Rapport_PFE_Copie0.docx
==============================================================================
- Conserve les styles existants (Titre1/Titre2/Titre3/Titre4, CaptionFigure,
  CaptionTableau)
- Conserve la structure des sauts de section (header/footer par chapitre)
- Légendes via champs SEQ Figure / SEQ Tableau (numérotation dynamique Word)
- Démarre à Figure 77 et Tableau 9 (continuation des 1ère-6ème chapitres)
- Active updateFields=true → Word met à jour TOC + Liste figures/tableaux
  à l'ouverture du fichier
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

# ──────────────────────────────────────────────────────────────────────
# Configuration
# ──────────────────────────────────────────────────────────────────────
COPIE0 = "/tmp/Copie0.docx"
DIAG = Path("/tmp/diagrams_staruml")
OUTPUT = Path("/home/user/PROJET-PFE/Rapport/Rapport_PFE_Final.docx")

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Mapping des styles du Copie 0
STYLES = {
    "h1":    "Titre1",         # Chapter title
    "h2":    "Titre2",         # 1., 2., 3.
    "h3":    "Titre3",         # 1.1, 4.1
    "h4":    "Titre4",         # 4.3.1
    "normal": None,             # plain Normal (no pStyle)
    "bullet": "Paragraphedeliste",
    "cap_fig": "CaptionFigure",
    "cap_tab": "CaptionTableau",
}

# ──────────────────────────────────────────────────────────────────────
# Helpers XML
# ──────────────────────────────────────────────────────────────────────
def wel(tag, **attrs):
    el = OxmlElement(f"w:{tag}")
    for k, v in attrs.items():
        el.set(qn(f"w:{k}"), str(v))
    return el

def add_text_run(p_el, text, *, bold=False, italic=False, size_pt=None, preserve_space=True):
    """Ajoute un run avec texte dans un élément paragraphe."""
    r = wel("r")
    rPr = wel("rPr")
    if bold:
        rPr.append(wel("b"))
        rPr.append(wel("bCs"))
    if italic:
        rPr.append(wel("i"))
        rPr.append(wel("iCs"))
    if size_pt:
        rPr.append(wel("sz", val=str(size_pt * 2)))
        rPr.append(wel("szCs", val=str(size_pt * 2)))
    if len(rPr):
        r.append(rPr)
    t = wel("t")
    if preserve_space:
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p_el.append(r)

def make_paragraph(style_id=None, text=None, align=None, page_break_before=False,
                   first_line_indent_cm=None):
    """Construit un <w:p> avec style optionnel et texte simple."""
    p = wel("p")
    pPr = wel("pPr")
    if style_id:
        pPr.append(wel("pStyle", val=style_id))
    if page_break_before:
        pPr.append(wel("pageBreakBefore"))
    if align:
        pPr.append(wel("jc", val=align))
    if first_line_indent_cm:
        ind = wel("ind")
        # 567 twips = 1 cm
        ind.set(qn("w:firstLine"), str(int(first_line_indent_cm * 567)))
        pPr.append(ind)
    if len(pPr):
        p.append(pPr)
    if text:
        add_text_run(p, text)
    return p

def make_section_break_paragraph(reference_sectPr):
    """Crée un paragraphe ne contenant qu'un sectPr (saut de section).
    Utilise la copie de reference_sectPr pour préserver les en-têtes/pieds."""
    p = wel("p")
    pPr = wel("pPr")
    pPr.append(deepcopy(reference_sectPr))
    p.append(pPr)
    return p

def make_caption_paragraph(kind, text):
    """Construit une légende avec champ SEQ dynamique.
    kind : 'Figure' ou 'Tableau'
    text : suffixe après « Figure NN : » (le NN sera généré par Word)"""
    style_id = STYLES["cap_fig"] if kind == "Figure" else STYLES["cap_tab"]
    p = wel("p")
    pPr = wel("pPr")
    pPr.append(wel("pStyle", val=style_id))
    pPr.append(wel("jc", val="center"))
    p.append(pPr)
    # « Figure » + champ SEQ + « : Texte »
    add_text_run(p, f"{kind} ")
    # Champ SEQ : <w:fldSimple w:instr=" SEQ Figure \* ARABIC ">
    fld = wel("fldSimple")
    fld.set(qn("w:instr"), f" SEQ {kind} \\* ARABIC ")
    # cache visible (sera recalculé par Word à l'ouverture)
    r = wel("r")
    t = wel("t")
    t.text = "?"
    r.append(t)
    fld.append(r)
    p.append(fld)
    add_text_run(p, f" : {text}")
    return p

def make_centered_image_paragraph(rId, width_cm=15.0):
    """Construit un paragraphe centré contenant l'image référencée par rId."""
    p = wel("p")
    pPr = wel("pPr")
    pPr.append(wel("jc", val="center"))
    p.append(pPr)
    # Calcule la taille en EMU (1 cm = 360000 EMU)
    cx = int(width_cm * 360000)
    # Préserve le ratio par défaut → on ajustera height plus tard
    # Pour l'instant approximation 4/5 H = 80% W
    cy = int(cx * 0.75)
    r = wel("r")
    # On va injecter un drawing standard
    drawing_xml = f"""<w:drawing xmlns:w="{W}">
      <wp:inline distT="0" distB="0" distL="0" distR="0"
                 xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="0" name="Picture"/>
        <wp:cNvGraphicFramePr>
          <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                                noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
        <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
            <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:nvPicPr>
                <pic:cNvPr id="0" name="img"/>
                <pic:cNvPicPr/>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                         r:embed="{rId}"/>
                <a:stretch>
                  <a:fillRect/>
                </a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm>
                  <a:off x="0" y="0"/>
                  <a:ext cx="{cx}" cy="{cy}"/>
                </a:xfrm>
                <a:prstGeom prst="rect">
                  <a:avLst/>
                </a:prstGeom>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:inline>
    </w:drawing>"""
    r.append(etree.fromstring(drawing_xml))
    p.append(r)
    return p

# ──────────────────────────────────────────────────────────────────────
# Construction des tableaux (Backlog & Tests)
# ──────────────────────────────────────────────────────────────────────
def make_table(headers, rows, *, col_widths_cm=None):
    """Construit un <w:tbl> imitant le style des tableaux des chapitres
    précédents (en-tête grisé bleu, bordures noires fines, texte 10 pt)."""
    tbl = wel("tbl")
    tblPr = wel("tblPr")
    tblPr.append(wel("tblStyle", val="Grilleducoloration1Accentuation1"))
    tblPr.append(wel("tblW", w="0", type="auto"))
    tblPr.append(wel("jc", val="center"))
    # Look = en-tête
    tblLook = wel("tblLook", val="04A0", firstRow="1", lastRow="0",
                  firstColumn="0", lastColumn="0", noHBand="0", noVBand="1")
    tblPr.append(tblLook)
    tbl.append(tblPr)

    # tblGrid
    tblGrid = wel("tblGrid")
    if col_widths_cm:
        for w_cm in col_widths_cm:
            tblGrid.append(wel("gridCol", w=str(int(w_cm * 567))))
    else:
        for _ in headers:
            tblGrid.append(wel("gridCol", w="2000"))
    tbl.append(tblGrid)

    def make_cell(text, *, bold=False, shading_fill=None, font_size=10):
        tc = wel("tc")
        tcPr = wel("tcPr")
        # Borders
        tcBorders = wel("tcBorders")
        for side in ("top", "left", "bottom", "right"):
            b = wel(side, val="single", sz="6", space="0", color="000000")
            tcBorders.append(b)
        tcPr.append(tcBorders)
        if shading_fill:
            tcPr.append(wel("shd", val="clear", color="auto", fill=shading_fill))
        tcPr.append(wel("vAlign", val="center"))
        tc.append(tcPr)
        p = wel("p")
        pPr = wel("pPr")
        pPr.append(wel("jc", val="left" if not bold else "center"))
        spc = wel("spacing", before="40", after="40", line="240", lineRule="auto")
        pPr.append(spc)
        p.append(pPr)
        r = wel("r")
        rPr = wel("rPr")
        if bold:
            rPr.append(wel("b"))
        rPr.append(wel("sz", val=str(font_size * 2)))
        rPr.append(wel("szCs", val=str(font_size * 2)))
        rPr.append(wel("color", val="000000"))
        r.append(rPr)
        t = wel("t")
        t.set(qn("xml:space"), "preserve")
        t.text = text or ""
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    # Header row with shading (light blue)
    tr_h = wel("tr")
    trPr = wel("trPr")
    trPr.append(wel("tblHeader"))
    tr_h.append(trPr)
    for h in headers:
        tr_h.append(make_cell(h, bold=True, shading_fill="D9E2F3"))
    tbl.append(tr_h)

    # Body rows
    for row in rows:
        tr = wel("tr")
        for v in row:
            tr.append(make_cell(str(v)))
        tbl.append(tr)
    return tbl

# ──────────────────────────────────────────────────────────────────────
# CONTENU — Chapitre 7 et 8
# ──────────────────────────────────────────────────────────────────────
S4_BACKLOG = [
    ("US4.1",  "Consulter les bons de commande",
     "Lister les BC paginés, filtrer par statut DO_Valide, rechercher par référence ou client.",
     "Melek Lahmar", "3 j"),
    ("US4.2",  "Consulter le détail d'un bon de commande",
     "Afficher les lignes F_DOCLIGNE, le client enrichi, les montants TTC et le mode de paiement.",
     "Tawfik Siala", "2 j"),
    ("US4.3",  "Modifier le statut d'un BC",
     "PUT /api/confirmateur/commandes/{piece}/status avec contrôle du workflow.",
     "Melek Lahmar", "3 j"),
    ("US4.4",  "Transformer un BC en bon de livraison",
     "POST /transform-to-bl : copie des lignes, anti-doublon, création F_LIVRAISON, transaction SQL.",
     "Melek Lahmar", "4 j"),
    ("US4.5",  "Exporter un BL vers Sage X3",
     "Envoi non bloquant via SageX3Client ; sageSuccess et sageHttpStatus journalisés.",
     "Tawfik Siala", "3 j"),
    ("US4.6",  "Affecter automatiquement un livreur",
     "AssignmentService : filtre Haversine par zone du client puis tri par file la plus courte.",
     "Melek Lahmar", "3 j"),
    ("US4.7",  "Traiter les devis B2B côté confirmateur",
     "GET /confirmateur/devis, POST /devis/{piece}/transform-to-bc.",
     "Tawfik Siala", "3 j"),
    ("US4.8",  "Consulter les livraisons disponibles",
     "GET /api/livreur/orders/available filtré par gouvernorat et délégation du livreur.",
     "Tawfik Siala", "3 j"),
    ("US4.9",  "Prendre en charge une livraison",
     "POST /assign avec gestion du conflit 409 si déjà prise par un autre livreur.",
     "Tawfik Siala", "2 j"),
    ("US4.10", "Diffuser la position GPS en temps réel",
     "POST /start-heading + boucle /location/ping 10 s + /ping-batch idempotent (ClientActionId).",
     "Melek Lahmar", "4 j"),
    ("US4.11", "Mettre à jour le statut d'une livraison",
     "Huit codes (0 Confirme à 7 DepotPret) ; encaissement COD automatique au statut 2 (Livre).",
     "Tawfik Siala", "3 j"),
    ("US4.12", "Gérer la caisse COD du livreur",
     "Encaisse, MontantEncaisse, EncaisseAt sur F_LIVRAISON ; remise au dépôt et historique.",
     "Melek Lahmar", "3 j"),
    ("US4.13", "Optimiser la tournée du livreur",
     "Plus-proche-voisin + Haversine, ETA cumulatif (35 km/h), affichage Mapbox + OSRM.",
     "Melek Lahmar", "3 j"),
    ("US4.14", "Créer et suivre une réclamation",
     "POST /api/reclamations avec photos compressées ; TrackingStateCard 4 états côté client.",
     "Tawfik Siala", "3 j"),
    ("US4.15", "Traiter une réclamation côté confirmatrice",
     "Verrou + grâce 5 s SignalR, chat NouveauMessage, escalade automatique au seuil 3 tentatives.",
     "Melek Lahmar", "3 j"),
    ("US4.16", "Gérer les missions de transit inter-dépôts",
     "Double scan code-barres mobile_scanner, réception partielle, fenêtre d'annulation 10 min.",
     "Tawfik Siala", "4 j"),
    ("US4.17", "Superviser l'activité terrain",
     "Carte temps réel des livreurs, heatmap 90 jours, réaffectation de mission, alertes par sévérité.",
     "Melek Lahmar", "4 j"),
]

S4_TESTS = [
    ("TF4.1",  "Consultation BC confirmatrice",
     "Les BC s'affichent avec le bon statut, le client et le montant TTC.", "Validé"),
    ("TF4.2",  "Transformation BC en BL sans doublon",
     "Le BL est créé en transaction ; si déjà existant, sa référence est renvoyée.", "Validé"),
    ("TF4.3",  "Sage X3 indisponible (export non bloquant)",
     "Le BL est créé localement même si Sage répond en erreur ; sageSuccess = false.", "Validé"),
    ("TF4.4",  "Affectation automatique du livreur",
     "Le livreur retenu est de zone et possède la file la plus courte.", "Validé"),
    ("TF4.5",  "Pool filtré par zone",
     "Seules les livraisons du gouvernorat et de la délégation du livreur apparaissent.", "Validé"),
    ("TF4.6",  "Conflit de prise en charge (409)",
     "Le second livreur reçoit 409 Conflict ; aucun doublon F_LIVRAISON.", "Validé"),
    ("TF4.7",  "Diffusion GPS SignalR",
     "Position reçue côté client en moins d'une seconde, indicateur de fraîcheur correct.", "Validé"),
    ("TF4.8",  "Mode hors ligne et /ping-batch",
     "Les positions sont mises en file locale puis envoyées en lot à la reconnexion.", "Validé"),
    ("TF4.9",  "Encaissement COD automatique",
     "Encaisse = true et MontantEncaisse = DO_NetAPayer dès le statut 2 (Livre).", "Validé"),
    ("TF4.10", "TrackingStateCard adaptative",
     "Bascule AT_DEPOT → IN_DELIVERY_QUEUE → HEADING_TO_YOU → TERMINAL conforme.", "Validé"),
    ("TF4.11", "Réclamation SignalR NouveauCas",
     "La réclamation apparaît côté confirmatrice sans rechargement de page.", "Validé"),
    ("TF4.12", "Grâce 5 s à la déconnexion",
     "Reconnexion < 5 s : verrou conservé ; au-delà : libération + CommandeLiberee.", "Validé"),
    ("TF4.13", "Escalade au seuil de 3 tentatives",
     "À la 3ᵉ tentative échouée, l'événement SeuilTentativesAtteint est émis.", "Validé"),
    ("TF4.14", "Reprogrammation par créneau",
     "Les créneaux MATIN, APRES_MIDI et SOIR sont sauvegardés et respectés.", "Validé"),
    ("TF4.15", "Scan pickup transit",
     "L'article passe en EN_TRANSIT, PickedUpAt + GPS et audit log écrits.", "Validé"),
    ("TF4.16", "Réception partielle",
     "TRANSIT_PARTIELLEMENT_RECU déclenche une alerte vers le groupe superviseurs.", "Validé"),
    ("TF4.17", "Annulation pickup dans 10 minutes",
     "POST /revert-pickup réussit dans la fenêtre et crée un audit log REVERT.", "Validé"),
    ("TF4.18", "Optimisation de tournée",
     "L'ordre obtenu minimise la distance totale ; ETA cumulatifs cohérents.", "Validé"),
    ("TF4.19", "Heatmap 90 jours",
     "Les cellules de 500 m affichent l'intensité réelle des échecs.", "Validé"),
    ("TF4.20", "Réaffectation de mission de transit",
     "Le superviseur peut transférer une mission bloquée vers un autre livreur.", "Validé"),
]

S5_BACKLOG = [
    ("US5.1",  "Vue d'ensemble multi-modules",
     "GET /api/dashboard/overview : KPI consolidés, alertes, série temporelle 7 jours.",
     "Melek Lahmar", "3 j"),
    ("US5.2",  "Analyse des ventes",
     "GET /api/dashboard/sales : CA par période, top produits, répartition B2C/B2B.",
     "Tawfik Siala", "3 j"),
    ("US5.3",  "Analyse des commandes",
     "GET /api/dashboard/orders : statuts, délais, segmentation par zone et par livreur.",
     "Melek Lahmar", "3 j"),
    ("US5.4",  "Analyse des produits",
     "GET /api/dashboard/products : top ventes, ruptures, marges indicatives.",
     "Tawfik Siala", "3 j"),
    ("US5.5",  "Analyse des stocks et dépôts",
     "GET /api/dashboard/stock + /depots : niveaux par dépôt, alertes critiques.",
     "Melek Lahmar", "3 j"),
    ("US5.6",  "Analyse logistique",
     "GET /api/dashboard/logistics : taux de livraison par zone, heatmap, délais moyens.",
     "Tawfik Siala", "3 j"),
    ("US5.7",  "Analyse des livreurs",
     "GET /api/dashboard/drivers : classement, taux de réussite, caisse COD.",
     "Melek Lahmar", "3 j"),
    ("US5.8",  "Analyse des clients",
     "GET /api/dashboard/clients : fidélité, panier moyen, segmentation B2C/B2B.",
     "Tawfik Siala", "3 j"),
    ("US5.9",  "Analyse des réclamations",
     "GET /api/dashboard/reclamations : motifs, SLA, performance des confirmatrices.",
     "Melek Lahmar", "3 j"),
    ("US5.10", "Performance des confirmatrices",
     "GET /api/dashboard/confirmateur : cas actifs, taux de résolution, temps moyen.",
     "Tawfik Siala", "2 j"),
    ("US5.11", "État de la synchronisation Sage X3",
     "GET /api/dashboard/sync : dernière exécution, intégrité, articles/dépôts/stocks à jour.",
     "Melek Lahmar", "2 j"),
    ("US5.12", "Insights stratégiques automatiques",
     "GET /api/dashboard/strategic-insights : tendances, anomalies, recommandations.",
     "Tawfik Siala", "2 j"),
    ("US5.13", "Insights proactifs du chatbot",
     "ProactiveInsightsJob (Hangfire, toutes les 30 min) alimente F_CHATBOT_INSIGHTS.",
     "Melek Lahmar", "2 j"),
    ("US5.14", "Export Excel ClosedXML",
     "GET /api/admin/orders/export?format=xlsx : entêtes en gras, largeurs auto, plafond 10 000 lignes.",
     "Melek Lahmar", "2 j"),
    ("US5.15", "Export PDF QuestPDF",
     "GET /api/admin/orders/export?format=pdf : mise en page A4 avec en-tête, table et pagination.",
     "Tawfik Siala", "2 j"),
    ("US5.16", "Chatbot conversationnel /ask",
     "Orchestrator interne : Groq Router → Exécuteur → Groq Formatter (llama-3.3-70b-versatile).",
     "Tawfik Siala", "3 j"),
    ("US5.17", "Streaming SSE /ask-stream",
     "Phases routing → data → chunks → done pour un rendu mot-à-mot côté React.",
     "Melek Lahmar", "2 j"),
    ("US5.18", "Workflow n8n d'orchestration",
     "admin-chatbot-workflow-v3.json : webhook → langue → backend → insights → merge → response.",
     "Tawfik Siala", "2 j"),
    ("US5.19", "Prédictions ML.NET",
     "PredictionService : trois modèles (risque, churn, routing) consommés via /predict.",
     "Melek Lahmar", "3 j"),
    ("US5.20", "Chatbot vocal sur mobile",
     "voice_buttons.dart : reconnaissance vocale speech_to_text et lecture flutter_tts.",
     "Tawfik Siala", "2 j"),
    ("US5.21", "Conversations et historique chatbot",
     "F_CHATBOT_MESSAGE par session, filtre par date et par utilisateur.",
     "Melek Lahmar", "2 j"),
    ("US5.22", "Rafraîchissement de la base de connaissances",
     "KbGeneratorService (IHostedService) régénère le contexte exploité par Groq.",
     "Tawfik Siala", "2 j"),
]

S5_TESTS = [
    ("TF5.1",  "Cohérence des KPI de la vue d'ensemble",
     "Les KPI correspondent à un recalcul SQL direct sur les mêmes filtres.", "Validé"),
    ("TF5.2",  "Dashboard ventes — chiffre d'affaires",
     "Les montants reflètent uniquement les commandes livrées (LI_Statut = 2).", "Validé"),
    ("TF5.3",  "Dashboard logistique — taux de livraison",
     "Le taux est cohérent avec la répartition réelle des statuts en base.", "Validé"),
    ("TF5.4",  "Dashboard réclamations — délais",
     "Le délai moyen est conforme à CreatedAt → CloturedAt.", "Validé"),
    ("TF5.5",  "Dashboard sync Sage X3",
     "Les compteurs d'articles, dépôts et stocks à jour sont exacts.", "Validé"),
    ("TF5.6",  "Insights stratégiques",
     "Les insights reflètent les tendances réelles observées en base.", "Validé"),
    ("TF5.7",  "Export Excel — format et contenu",
     "Le fichier .xlsx s'ouvre correctement avec entêtes gras et colonnes ajustées.", "Validé"),
    ("TF5.8",  "Export PDF — pagination",
     "Le PDF A4 affiche l'en-tête, le tableau et la pagination.", "Validé"),
    ("TF5.9",  "Plafond MaxRows = 10 000",
     "Au-delà, le service tronque et l'utilisateur reçoit un avertissement.", "Validé"),
    ("TF5.10", "Chatbot — question simple",
     "« Quel est le CA du jour ? » renvoie un message FR et un chiffre exact.", "Validé"),
    ("TF5.11", "Chatbot — question contextuelle",
     "« Et hier ? » conserve le sujet grâce au sessionId.", "Validé"),
    ("TF5.12", "Chatbot — prédiction ML.NET",
     "L'action PREDICT renvoie un score cohérent (risque, churn ou routing).", "Validé"),
    ("TF5.13", "Streaming SSE",
     "Les quatre phases routing → data → chunks → done sont reçues dans l'ordre.", "Validé"),
    ("TF5.14", "Workflow n8n complet",
     "Les sept nœuds s'enchaînent et les insights sont fusionnés à la réponse.", "Validé"),
    ("TF5.15", "Insights proactifs Hangfire",
     "Le job de 30 min alimente correctement F_CHATBOT_INSIGHTS.", "Validé"),
    ("TF5.16", "Chatbot vocal mobile",
     "Speech-to-text capte la question et flutter_tts lit la réponse.", "Validé"),
    ("TF5.17", "Rafraîchissement de la base de connaissances",
     "POST /kb/refresh régénère la base et la rend exploitable par le router Groq.", "Validé"),
    ("TF5.18", "Panneau d'alertes par sévérité",
     "Les alertes critiques apparaissent en rouge, les avertissements en orange.", "Validé"),
]

# ──────────────────────────────────────────────────────────────────────
# CONSTRUCTION DES DEUX CHAPITRES
# ──────────────────────────────────────────────────────────────────────
def build_chapter_7(diag_rids):
    """Retourne la liste des éléments XML composant le nouveau chapitre 7."""
    els = []
    # Page break + titre chapitre
    els.append(make_paragraph(style_id=STYLES["h1"],
                              text="Chapitre 7 : Sprint 4 — Logistique, livraison et réclamations",
                              page_break_before=True))

    # 1. Introduction
    els.append(make_paragraph(STYLES["h2"], "1. Introduction"))
    els.append(make_paragraph(None,
        "Ce chapitre présente le quatrième sprint de la solution. Il met en place le cœur "
        "opérationnel de la plateforme : la confirmatrice traite les bons de commande, le "
        "livreur prend en charge ses livraisons en mode COD avec suivi GPS en temps réel, le "
        "client suit l'avancement de sa commande, le superviseur observe l'activité et le "
        "livreur de transit déplace les marchandises entre dépôts. Les réclamations sont "
        "également traitées en temps réel, avec un chat entre le client et la confirmatrice. "
        "L'objectif est de fermer la chaîne logistique de bout en bout.",
        first_line_indent_cm=0.6))

    # 2. Objectif
    els.append(make_paragraph(STYLES["h2"], "2. Objectif du sprint"))
    els.append(make_paragraph(None,
        "Le Sprint 4 s'est déroulé du 06/04/2026 au 25/04/2026. Son objectif est de relier "
        "la commande en ligne à la livraison réelle au client. Il couvre cinq ensembles de "
        "fonctionnalités : la confirmation et la transformation documentaire avec Sage X3, "
        "la livraison COD avec suivi GPS temps réel, le suivi du client à travers une carte "
        "adaptative, la gestion des réclamations avec chat en temps réel, et le transit "
        "inter-dépôts par double scan code-barres. La supervision de l'activité terrain est "
        "également mise en place pour permettre une intervention rapide en cas d'écart.",
        first_line_indent_cm=0.6))
    els.append(make_paragraph(None,
        "Le sprint mobilise les trois canaux de la solution. Les interfaces React sont "
        "utilisées par la confirmatrice et le superviseur. Les écrans Flutter sont destinés "
        "au livreur, au livreur de transit et au client. La Web API ASP.NET Core centralise "
        "les règles et expose les endpoints. SignalR assure la communication temps réel via "
        "le ReclamationHub et le SupervisorHub. La base SQL Server stocke les documents "
        "(F_DOCENTETE, F_DOCLIGNE), la livraison (F_LIVRAISON), la position du livreur "
        "(F_LIVREUR_POSITION et son historique), les réclamations (F_RECLAMATION) et les "
        "transferts (F_TRANSFERT) avec leur journal d'audit (F_TRANSFERT_AUDIT_LOG).",
        first_line_indent_cm=0.6))

    # 3. Backlog
    els.append(make_paragraph(STYLES["h2"], "3. Backlog du sprint"))
    els.append(make_paragraph(None,
        "Le tableau suivant présente les vingt histoires utilisateurs traitées pendant ce "
        "sprint. Elles couvrent l'ensemble des rôles concernés et respectent la répartition "
        "habituelle du travail entre les deux membres du binôme.",
        first_line_indent_cm=0.6))
    els.append(make_caption_paragraph("Tableau", "Backlog du Sprint 4"))
    els.append(make_table(
        ["ID", "Histoire utilisateur", "Tâches techniques", "Responsable", "Estim."],
        S4_BACKLOG,
        col_widths_cm=[1.4, 4.5, 6.0, 2.5, 1.5]))
    els.append(make_paragraph(None, ""))  # spacer

    # 4. Analyse et conception
    els.append(make_paragraph(STYLES["h2"], "4. Analyse et conception"))
    els.append(make_paragraph(None,
        "L'analyse de ce sprint précise les interactions entre les cinq acteurs du terrain "
        "et les composants de la solution. Le diagramme de cas d'utilisation présente une "
        "vue d'ensemble. Les diagrammes de séquence détaillent les quatre flux les plus "
        "structurants : la transformation documentaire avec Sage X3, la diffusion GPS et "
        "l'encaissement COD, la réclamation temps réel et le double scan du transit.",
        first_line_indent_cm=0.6))

    els.append(make_paragraph(STYLES["h3"], "4.1 Diagramme de cas d'utilisation du Sprint 4"))
    els.append(make_paragraph(None,
        "Le diagramme regroupe les fonctionnalités en cinq paquets correspondant aux grands "
        "domaines fonctionnels du sprint : confirmation, livraison COD, transit inter-dépôts, "
        "réclamations temps réel et supervision.",
        first_line_indent_cm=0.6))
    els.append(make_centered_image_paragraph(diag_rids["s4_usecase"], width_cm=14.5))
    els.append(make_caption_paragraph("Figure", "Diagramme de cas d'utilisation du Sprint 4"))

    els.append(make_paragraph(STYLES["h3"], "4.2 Description des cas d'utilisation"))
    els.append(make_paragraph(None,
        "La confirmatrice consulte les bons de commande, ajuste leur statut et déclenche la "
        "transformation BC vers BL, opération qui crée le bon de livraison en local et "
        "l'envoie à Sage X3 sans bloquer le processus. Elle traite également les devis B2B "
        "et les réclamations clients. Le livreur consulte le pool des livraisons filtré par "
        "sa zone, prend en charge une livraison disponible, diffuse sa position GPS, met à "
        "jour le statut (huit codes de 0 à 7) et gère sa caisse COD. Le livreur de transit "
        "se concentre sur les missions de transfert entre dépôts par double scan de "
        "code-barres. Le client suit sa commande grâce à une carte adaptative en quatre "
        "états et peut déposer une réclamation. Le superviseur surveille l'activité en temps "
        "réel, consulte la heatmap des échecs et réaffecte les missions de transit lorsque "
        "c'est nécessaire.",
        first_line_indent_cm=0.6))

    els.append(make_paragraph(STYLES["h3"], "4.3 Diagrammes de séquence"))
    els.append(make_paragraph(None,
        "Quatre flux temps réel ont été modélisés. Ils précisent les échanges entre "
        "l'interface, la Web API, la base SQL Server, le hub SignalR et, lorsque c'est "
        "nécessaire, les services externes Sage X3 et n8n.",
        first_line_indent_cm=0.6))

    els.append(make_paragraph(STYLES["h4"], "4.3.1 Transformation BC en bon de livraison"))
    els.append(make_paragraph(None,
        "La méthode TransformBcToBl de ConfirmateurController vérifie d'abord l'existence "
        "du BC. Si aucun BL n'a déjà été créé, elle ouvre une transaction, insère l'entête "
        "F_DOCENTETE avec DO_Type = 1 et la référence formée de « BL » suivi de la date et "
        "d'un compteur, puis recopie les lignes F_DOCLIGNE. Elle crée ensuite la livraison "
        "F_LIVRAISON avec LI_Statut = 0 (Confirme) et passe le BC en statut CONFIRME avant "
        "le COMMIT. L'AssignmentService sélectionne un livreur dont la zone couvre l'adresse "
        "du client (test Haversine) et dont la file de livraisons en cours est la plus "
        "courte. L'envoi vers Sage X3 est asynchrone : la réponse expose sageSent, "
        "sageSuccess et sageHttpStatus pour permettre la supervision.",
        first_line_indent_cm=0.6))
    els.append(make_centered_image_paragraph(diag_rids["s4_seq_bctbl"], width_cm=15.0))
    els.append(make_caption_paragraph("Figure", "Diagramme de séquence — Transformation BC en BL"))

    els.append(make_paragraph(STYLES["h4"], "4.3.2 Livraison COD et diffusion GPS via SignalR"))
    els.append(make_paragraph(None,
        "Le livreur démarre la livraison en appelant /start-heading. La Web API marque la "
        "livraison comme active (IsActiveDelivery = true) et envoie l'événement "
        "DeliveryStarted au groupe SignalR client-{userId}. Le client bascule alors sur "
        "l'état HEADING_TO_YOU. Toutes les 10 secondes, le Geolocator de Flutter envoie un "
        "ping contenant la latitude, la longitude et un identifiant ClientActionId. La Web "
        "API effectue un UPSERT dans F_LIVREUR_POSITION, ajoute une ligne dans "
        "F_LIVREUR_POSITION_HISTORY, calcule la distance Haversine vers le client et émet "
        "l'événement LivreurPositionUpdate. Si la distance descend sous 500 mètres, un "
        "événement de proximité est aussi envoyé. En l'absence de réseau, les positions "
        "sont stockées localement dans une file Hive et envoyées plus tard via "
        "/ping-batch grâce à l'idempotence garantie par ClientActionId. Lorsque le livreur "
        "valide la livraison (statut 2), F_LIVRAISON est mise à jour avec Encaisse = true "
        "et MontantEncaisse = DO_NetAPayer ; le client est notifié et la diffusion GPS "
        "s'arrête.",
        first_line_indent_cm=0.6))
    els.append(make_centered_image_paragraph(diag_rids["s4_seq_gps"], width_cm=14.5))
    els.append(make_caption_paragraph("Figure", "Diagramme de séquence — Livraison COD et diffusion GPS"))

    els.append(make_paragraph(STYLES["h4"], "4.3.3 Réclamation client avec SignalR et grâce de 5 secondes"))
    els.append(make_paragraph(None,
        "La création d'une réclamation par un client provoque l'envoi de l'événement "
        "NouveauCas au groupe confirmateurs. La confirmatrice qui prend en charge le cas "
        "ouvre une session F_CONFIRMATRICE_SESSION et verrouille le cas. La discussion se "
        "déroule via l'événement NouveauMessage diffusé au groupe client-{userId} et au "
        "groupe confirmateurs. La particularité de ce flux est la gestion de la grâce de "
        "5 secondes à la déconnexion. La méthode OnDisconnectedAsync du ReclamationHub "
        "décrémente le compteur de connexions de la confirmatrice. Si c'est la dernière "
        "connexion, un timer GracePeriod = TimeSpan.FromSeconds(5) est armé. Si la "
        "confirmatrice se reconnecte dans cet intervalle, le timer est annulé et le verrou "
        "conservé. Si la grâce expire, le cas redevient libre et l'événement "
        "CommandeLiberee est diffusé. Lorsque trois tentatives de livraison ont échoué "
        "(F_RECLAMATION_TENTATIVE), l'événement SeuilTentativesAtteint est émis vers les "
        "superviseurs.",
        first_line_indent_cm=0.6))
    els.append(make_centered_image_paragraph(diag_rids["s4_seq_reclamation"], width_cm=14.5))
    els.append(make_caption_paragraph("Figure", "Diagramme de séquence — Réclamation client avec SignalR"))

    els.append(make_paragraph(STYLES["h4"], "4.3.4 Transit inter-dépôts par double scan code-barres"))
    els.append(make_paragraph(None,
        "Le livreur de transit ouvre la mission dans transit_mission_details_screen. Le "
        "scanner transit_barcode_scanner_screen utilise le plugin mobile_scanner pour lire "
        "les codes. Un premier scan (POST /scan-pickup) fait passer la ligne F_TRANSFERT de "
        "EN_ATTENTE_TRANSIT à EN_TRANSIT, enregistre PickedUpAt et la position GPS, et "
        "écrit un journal F_TRANSFERT_AUDIT_LOG avec ActionType = PICKUP. Le second scan "
        "(POST /scan-delivery) au dépôt destination bascule la ligne en "
        "RECU_DEPOT_DESTINE. Une fenêtre d'annulation de dix minutes "
        "(RevertWindowMinutes = 10) permet de corriger une erreur, avec un nouveau journal "
        "d'audit. Si la quantité reçue est inférieure à la quantité attendue, l'application "
        "déclenche POST /scan-partial : la ligne passe en TRANSIT_PARTIELLEMENT_RECU et une "
        "alerte est envoyée au groupe superviseurs.",
        first_line_indent_cm=0.6))
    els.append(make_centered_image_paragraph(diag_rids["s4_seq_transit"], width_cm=14.5))
    els.append(make_caption_paragraph("Figure", "Diagramme de séquence — Scan codes-barres transit inter-dépôts"))

    # 5. Réalisation
    els.append(make_paragraph(STYLES["h2"], "5. Réalisation"))
    els.append(make_paragraph(None,
        "La réalisation utilise les trois canaux de la solution. Les écrans React de la "
        "confirmatrice et du superviseur consomment les endpoints via Axios. Les écrans "
        "Flutter passent par un ApiClient qui injecte automatiquement le jeton JWT stocké "
        "en sécurisé (flutter_secure_storage). Les notifications sonores des scans utilisent "
        "audioplayers et les vibrations sont gérées par le plugin du scanner. Les photos des "
        "réclamations sont compressées par flutter_image_compress avant l'envoi pour réduire "
        "le coût réseau. Les sections suivantes présentent les principaux écrans réalisés.",
        first_line_indent_cm=0.6))

    realisations_s4 = [
        ("5.1 Interface confirmatrice — liste des commandes",
         "L'écran ConfirmateurOrdersPage liste les bons de commande paginés. La "
         "confirmatrice peut filtrer par statut (EN_ATTENTE, CONFIRME, EN_LIVRAISON, "
         "LIVRE), rechercher par référence ou par client et consulter le détail. Chaque "
         "ligne affiche la référence DO_Piece, le client, la date, le montant TTC et le "
         "libellé du statut.",
         "Interface confirmatrice — liste des commandes"),
        ("5.2 Interface confirmatrice — transformation BC en BL",
         "Lorsque le bon de commande est confirmé, l'écran ConfirmateurOrderDetailsPage "
         "propose le bouton « Transformer en BL ». Le résultat de l'opération est affiché : "
         "référence du nouveau BL, indicateur alreadyExists si un BL existe déjà, sageSent "
         "et sageSuccess pour le suivi de la synchronisation Sage X3. L'identifiant du "
         "livreur automatiquement affecté est également rappelé.",
         "Interface confirmatrice — transformation BC en BL"),
        ("5.3 Interface confirmatrice — gestion des réclamations",
         "L'écran des réclamations affiche les cas par ordre de priorité et d'ancienneté. "
         "Les nouvelles réclamations apparaissent en temps réel grâce à l'événement "
         "SignalR NouveauCas, sans rechargement de page. La confirmatrice peut prendre en "
         "charge un cas, modifier son statut, envoyer une demande de correction d'adresse, "
         "créer un bon d'échange, ajouter une note interne, réaffecter le cas et chatter "
         "avec le client. L'historique des tentatives de livraison est consultable, avec "
         "mise en évidence du seuil de trois tentatives.",
         "Interface confirmatrice — gestion des réclamations"),
        ("5.4 Interface livreur — livraisons disponibles",
         "L'écran new_orders_screen présente les bons de livraison filtrés par "
         "gouvernorat et délégation du livreur. Pour chaque livraison, l'application "
         "affiche l'adresse, le montant COD attendu et les coordonnées du client. La prise "
         "en charge se fait d'un seul appui ; si un autre livreur prend la livraison en "
         "parallèle, l'API renvoie un code 409 Conflict et l'utilisateur en est informé.",
         "Interface livreur — livraisons disponibles"),
        ("5.5 Interface livreur — suivi GPS et carte de livraison",
         "La carte affiche la position du livreur (mise à jour via SignalR), les arrêts "
         "restants et l'itinéraire optimisé. L'algorithme du plus-proche-voisin avec "
         "distance Haversine ordonne les arrêts, et OSRM fournit un calcul d'ETA plus "
         "précis. La feuille LiveDeliveryMapSheet expose au client la position du livreur "
         "lorsqu'il est en route, avec un indicateur de fraîcheur du dernier ping.",
         "Interface livreur — suivi GPS et carte de livraison"),
        ("5.6 Interface livreur — caisse COD",
         "L'écran livreur_stats_screen présente la caisse du livreur : montant total "
         "encaissé, nombre de livraisons effectuées, taux de réussite et répartition par "
         "statut. L'historique des remises au dépôt est accessible. Les sous-états "
         "DepotEnCoursDePreparation (6) et DepotPret (7) sont visibles uniquement par le "
         "livreur, le client voyant un état agrégé « au dépôt ».",
         "Interface livreur — caisse COD"),
        ("5.7 Interface suivi commande client (TrackingStateCard)",
         "Le composant TrackingStateCard prend l'apparence dictée par l'endpoint "
         "/api/client/orders/{piece}/tracking-state. En AT_DEPOT, la carte est "
         "informative. En IN_DELIVERY_QUEUE, le client sait qu'un livreur va le prendre "
         "en charge. En HEADING_TO_YOU, la carte devient une vue cartographique avec la "
         "position GPS du livreur, l'ETA en minutes (calculé par Haversine à 40 km/h) et "
         "un bouton d'appel ou de SMS. En TERMINAL, la livraison est finalisée.",
         "Interface client — TrackingStateCard adaptative"),
        ("5.8 Interface de transit inter-dépôts et scan codes-barres",
         "L'écran transit_home_screen regroupe les missions en trois onglets (en attente, "
         "en cours, historique). Le détail d'une mission liste les articles attendus. Le "
         "scanner active la caméra ; chaque scan donne un retour sonore (audioplayers) et "
         "tactile. L'application gère également la réception partielle et la fenêtre "
         "d'annulation de dix minutes après un pickup.",
         "Interface de transit inter-dépôts et scan codes-barres"),
        ("5.9 Interface superviseur",
         "Le superviseur dispose d'une carte temps réel des livreurs actifs, d'un tableau "
         "de bord des alertes (retard, échec, zone non couverte, écart de caisse) et "
         "d'une heatmap construite par LivreurMapController.HeatMap. Les cellules sont "
         "d'environ 500 mètres et l'intensité représente le ratio (retours + reports) / "
         "livraisons sur les 90 derniers jours. Le superviseur peut aussi réaffecter une "
         "mission de transit grâce à l'endpoint /api/supervisor/transferts/{id}/reassign.",
         "Interface superviseur"),
    ]
    for title, body_text, fig_text in realisations_s4:
        els.append(make_paragraph(STYLES["h3"], title))
        els.append(make_paragraph(None, body_text, first_line_indent_cm=0.6))
        # placeholder paragraph (cadre vide en italique pour la capture à intégrer)
        ph = make_paragraph(None, "[ Emplacement réservé pour la capture d'écran ]",
                            align="center")
        # style italic via run prop
        for r in ph.findall(qn("w:r")):
            rPr = wel("rPr")
            rPr.append(wel("i"))
            rPr.append(wel("color", val="888888"))
            r.insert(0, rPr)
        els.append(ph)
        els.append(make_caption_paragraph("Figure", fig_text))

    # 6. Tests et validation
    els.append(make_paragraph(STYLES["h2"], "6. Tests et validation"))
    els.append(make_paragraph(None,
        "Les tests fonctionnels ont été menés tout au long du sprint sur des jeux de "
        "données réalistes (jeu de départ Sage X3 et seed Tunisie : neuf gouvernorats avec "
        "leurs délégations). Les cas limites — Sage X3 indisponible, conflit de prise en "
        "charge, perte de connexion du livreur, déconnexion fugace de la confirmatrice, "
        "scans erronés — ont été vérifiés en plus des cas nominaux.",
        first_line_indent_cm=0.6))
    els.append(make_caption_paragraph("Tableau", "Tests fonctionnels du Sprint 4"))
    els.append(make_table(
        ["ID", "Objectif", "Résultat attendu", "Statut"],
        S4_TESTS,
        col_widths_cm=[1.4, 4.5, 7.5, 1.5]))
    els.append(make_paragraph(None, ""))

    # 7. Conclusion
    els.append(make_paragraph(STYLES["h2"], "7. Conclusion"))
    els.append(make_paragraph(None,
        "Le Sprint 4 referme la chaîne logistique de la plateforme. Une commande passée "
        "en ligne peut désormais être confirmée, transformée en bon de livraison "
        "synchronisé avec Sage X3, affectée automatiquement à un livreur, suivie en temps "
        "réel par le client, livrée avec encaissement COD automatique, puis suivie d'une "
        "éventuelle réclamation traitée en conversation directe. Le transit inter-dépôts "
        "apporte la traçabilité des mouvements internes par double scan. Le superviseur "
        "dispose d'outils pour suivre l'activité et intervenir en cas d'écart. La "
        "résilience opérationnelle a été soignée : file locale hors ligne, idempotence "
        "des pings, grâce SignalR, synchronisation Sage non bloquante, audit log et jobs "
        "Hangfire d'escalade. Le sprint suivant exploite ce socle pour offrir à "
        "l'administrateur une couche complète de pilotage et d'aide à la décision.",
        first_line_indent_cm=0.6))

    return els

def build_chapter_8(diag_rids):
    els = []
    els.append(make_paragraph(STYLES["h1"],
        "Chapitre 8 : Sprint 5 — Tableaux de bord, exports et chatbot",
        page_break_before=True))

    # 1. Introduction
    els.append(make_paragraph(STYLES["h2"], "1. Introduction"))
    els.append(make_paragraph(None,
        "Ce chapitre présente le cinquième et dernier sprint de la solution. À ce stade, "
        "l'ensemble des fonctionnalités opérationnelles — catalogue, commandes, paiement, "
        "confirmation, livraison, réclamations, transit — est en place. Le Sprint 5 ajoute "
        "la couche de pilotage et d'aide à la décision destinée à l'administrateur. Trois "
        "axes sont traités en parallèle : les tableaux de bord multi-dimensionnels qui "
        "agrègent les données de tous les modules précédents, les exports Excel et PDF qui "
        "permettent d'extraire ces données vers des outils externes, et un chatbot "
        "administrateur capable d'interroger la base en langage naturel grâce à n8n, à un "
        "modèle de langage Groq et à trois modèles ML.NET.",
        first_line_indent_cm=0.6))

    els.append(make_paragraph(STYLES["h2"], "2. Objectif du sprint"))
    els.append(make_paragraph(None,
        "Le Sprint 5 s'est déroulé du 27/04/2026 au 23/05/2026. Son objectif est de "
        "transformer la solution en une plateforme pilotable. L'administrateur doit pouvoir "
        "consulter les indicateurs clés, identifier les anomalies, extraire les données et "
        "obtenir des réponses synthétiques à des questions en français. Le périmètre couvre "
        "onze tableaux de bord, deux services d'export, un chatbot conversationnel "
        "multimodal (texte sur web et vocal sur mobile) et un mécanisme d'insights "
        "proactifs qui détecte automatiquement les tendances et les écarts.",
        first_line_indent_cm=0.6))

    els.append(make_paragraph(STYLES["h2"], "3. Backlog du sprint"))
    els.append(make_paragraph(None,
        "Le tableau ci-après présente les vingt-deux histoires utilisateurs du sprint. "
        "Elles se répartissent entre les tableaux de bord, les exports, le chatbot et les "
        "services d'arrière-plan associés.",
        first_line_indent_cm=0.6))
    els.append(make_caption_paragraph("Tableau", "Backlog du Sprint 5"))
    els.append(make_table(
        ["ID", "Histoire utilisateur", "Tâches techniques", "Responsable", "Estim."],
        S5_BACKLOG,
        col_widths_cm=[1.4, 4.5, 6.0, 2.5, 1.5]))
    els.append(make_paragraph(None, ""))

    # 4. Analyse et conception
    els.append(make_paragraph(STYLES["h2"], "4. Analyse et conception"))
    els.append(make_paragraph(None,
        "L'analyse de ce sprint porte sur des fonctionnalités transverses qui réutilisent "
        "les données produites par les sprints précédents. Les tableaux de bord s'appuient "
        "sur un service d'agrégation unique. Les exports sont centralisés dans un service "
        "générique. Le chatbot est orchestré par un pipeline interne en trois étapes "
        "(Router, Exécuteur, Formatter) et complété par un workflow n8n qui injecte des "
        "insights proactifs.",
        first_line_indent_cm=0.6))

    els.append(make_paragraph(STYLES["h3"], "4.1 Diagramme de cas d'utilisation du Sprint 5"))
    els.append(make_paragraph(None,
        "Le diagramme regroupe les fonctionnalités en quatre paquets : tableaux de bord, "
        "exports, chatbot et pilotage avancé. Il fait apparaître les trois acteurs "
        "techniques qui interviennent en arrière-plan : le workflow n8n pour "
        "l'orchestration, le modèle Groq llama-3.3-70b-versatile pour le routage et la mise "
        "en forme des réponses, et ML.NET pour les prédictions. Sage X3 reste la source du "
        "tableau de bord de synchronisation.",
        first_line_indent_cm=0.6))
    els.append(make_centered_image_paragraph(diag_rids["s5_usecase"], width_cm=14.5))
    els.append(make_caption_paragraph("Figure", "Diagramme de cas d'utilisation du Sprint 5"))

    els.append(make_paragraph(STYLES["h3"], "4.2 Description des cas d'utilisation"))
    els.append(make_paragraph(None,
        "L'administrateur accède à la vue d'ensemble, à l'analyse des ventes, des "
        "commandes, des produits, des stocks et des dépôts, de la logistique, des "
        "livreurs, des clients, des réclamations, des confirmatrices et de la "
        "synchronisation Sage X3. Il consulte les insights stratégiques et les insights "
        "proactifs générés en arrière-plan, déclenche les exports Excel et PDF, dialogue "
        "avec le chatbot en mode texte ou voix, parcourt l'historique des conversations et "
        "peut rafraîchir manuellement la base de connaissances utilisée par le router "
        "Groq.",
        first_line_indent_cm=0.6))

    els.append(make_paragraph(STYLES["h3"], "4.3 Diagrammes de séquence"))
    els.append(make_paragraph(None,
        "Deux flux ont été modélisés en détail : la chaîne complète du chatbot "
        "administrateur (React → n8n → Backend → Groq → ML.NET) et la génération des "
        "exports Excel et PDF par le service ExportService.",
        first_line_indent_cm=0.6))

    els.append(make_paragraph(STYLES["h4"], "4.3.1 Diagramme de séquence — Chatbot administrateur"))
    els.append(make_paragraph(None,
        "L'administrateur saisit sa question depuis la page ChatbotSandboxPage. Le payload "
        "(question, sessionId, locale) est envoyé au webhook n8n /admin-chat-v3. Le premier "
        "nœud détecte la langue (français, arabe ou « tounsi ») puis transmet la requête au "
        "backend via POST /api/admin/chat/ask. L'AdminChatOrchestratorService applique "
        "alors le pipeline interne : un premier appel à Groq joue le rôle de routeur et "
        "choisit l'action (KB, QUERY, ANALYZE, PREDICT ou CHITCHAT). L'action est exécutée : "
        "interrogation SQL pour QUERY, modèle ML.NET pour PREDICT, lecture de la base de "
        "connaissances pour KB. Un second appel à Groq met en forme la réponse en français "
        "et propose éventuellement une structure de tableau et de graphique. Le message est "
        "journalisé dans F_CHATBOT_MESSAGE. Côté n8n, un nœud complémentaire récupère les "
        "insights proactifs en attente (F_CHATBOT_INSIGHTS) et les fusionne à la réponse "
        "avant le retour à React. Une variante de cet endpoint, /ask-stream, renvoie un "
        "flux SSE en quatre phases (routing, data, chunks, done) pour un rendu mot-à-mot.",
        first_line_indent_cm=0.6))
    els.append(make_centered_image_paragraph(diag_rids["s5_seq_chatbot"], width_cm=15.0))
    els.append(make_caption_paragraph("Figure", "Diagramme de séquence — Chatbot administrateur (n8n)"))

    els.append(make_paragraph(STYLES["h4"], "4.3.2 Diagramme de séquence — Export Excel et PDF"))
    els.append(make_paragraph(None,
        "L'administrateur déclenche un export depuis l'interface de synthèse. La route "
        "GET /api/admin/orders/export reçoit deux paramètres : format (xlsx ou pdf) et "
        "period. La méthode parse la période, projette F_DOCENTETES dans la limite de "
        "MaxRows = 10 000 lignes et délègue la sérialisation à ExportService. ClosedXML "
        "construit le classeur Excel avec entêtes en gras et colonnes ajustées "
        "automatiquement. QuestPDF construit un document A4 avec en-tête contextualisé, un "
        "tableau paginé et un pied de page. Le flux binaire est renvoyé en téléchargement "
        "direct via l'en-tête Content-Disposition. Le même service est réutilisé pour "
        "/api/admin/reclamations/export.",
        first_line_indent_cm=0.6))
    els.append(make_centered_image_paragraph(diag_rids["s5_seq_export"], width_cm=14.0))
    els.append(make_caption_paragraph("Figure", "Diagramme de séquence — Export Excel et PDF"))

    # 5. Réalisation
    els.append(make_paragraph(STYLES["h2"], "5. Réalisation"))
    els.append(make_paragraph(None,
        "Les onze tableaux de bord partagent une même infrastructure côté React : le "
        "composant DashboardAnalyticsPage est paramétré par une clé pageKey qui détermine "
        "l'endpoint à appeler, la configuration des KPI et la liste des alertes à "
        "afficher. Le rendu des graphiques utilise Recharts. Les exports sont déclenchés "
        "depuis les écrans de synthèse via un appel HTTP qui produit un téléchargement "
        "direct. Le chatbot est accessible depuis une fenêtre flottante (ChatbotFab) sur "
        "toutes les pages d'administration.",
        first_line_indent_cm=0.6))

    realisations_s5 = [
        ("5.1 Interface du tableau de bord global",
         "L'écran AdminOverviewDashboardPage agrège les indicateurs clés : chiffre "
         "d'affaires, nombre de livraisons, réclamations ouvertes, livreurs actifs et "
         "tendance hebdomadaire. Le panneau d'alertes AlertPanel met en évidence les "
         "anomalies par sévérité (avertissement, critique). Les insights stratégiques et "
         "les insights proactifs alimentés par Hangfire toutes les 30 minutes y sont "
         "également visibles.",
         "Interface du tableau de bord global"),
        ("5.2 Interface du dashboard ventes et commandes",
         "AdminSalesDashboardPage présente le chiffre d'affaires sur la période, le top "
         "produits et la répartition B2C / B2B. AdminOrdersDashboardPage détaille les "
         "statuts des commandes, les délais moyens (BC → BL → livraison) et la "
         "segmentation géographique. Les graphiques sont produits par Recharts à partir du "
         "JSON retourné par la Web API.",
         "Interface du dashboard ventes et commandes"),
        ("5.3 Interface du dashboard logistique et livreurs",
         "AdminLogisticsDashboardPage présente le taux de livraison par zone, la heatmap "
         "des échecs sur 90 jours et le délai moyen de livraison. "
         "AdminDriversDashboardPage expose le classement des livreurs (taux de réussite, "
         "volume, caisse COD) et leurs alertes individuelles. Les huit codes de statut "
         "sont restitués séparément pour permettre une lecture précise.",
         "Interface du dashboard logistique et livreurs"),
        ("5.4 Interface du dashboard réclamations et confirmatrices",
         "AdminReclamationsDashboardPage croise les réclamations par motif "
         "(COLIS_ENDOMMAGE, NON_LIVRE, MAUVAIS_ARTICLE, NUMERO_INCORRECT, "
         "ADRESSE_INCORRECTE) et par statut (ENVOYEE, EN_COURS, CLOTUREE, REFUSEE). Elle "
         "suit également les performances des confirmatrices : cas actifs, taux de "
         "résolution, temps moyen et durée des sessions F_CONFIRMATRICE_SESSION.",
         "Interface du dashboard réclamations et confirmatrices"),
        ("5.5 Interface du dashboard synchronisation Sage X3",
         "AdminSyncDashboardPage présente l'état des données synchronisées : nombre "
         "d'articles, de catalogues, de dépôts et de stocks à jour, alertes d'intégrité "
         "(clients sans adresse, articles sans famille, stocks négatifs) et horodatage de "
         "la dernière synchronisation. Elle permet à l'administrateur de vérifier en "
         "permanence la cohérence entre la plateforme et l'ERP.",
         "Interface du dashboard synchronisation Sage X3"),
        ("5.6 Interface des exports Excel et PDF",
         "L'écran de synthèse propose, à côté de chaque liste, deux boutons « Exporter en "
         "Excel » et « Exporter en PDF ». Les fichiers sont générés par ExportService "
         "(ClosedXML 0.104.1 et QuestPDF 2024.7.3) dans la limite de 10 000 lignes par "
         "export. La nomenclature des fichiers inclut la période sélectionnée pour "
         "faciliter le classement.",
         "Interface des exports Excel et PDF"),
        ("5.7 Interface du chatbot administrateur",
         "La page ChatbotSandboxPage offre une zone de saisie, l'historique du fil de "
         "conversation et le rendu structuré des réponses (message, tableau, graphique). "
         "Le bouton ChatbotFab permet d'ouvrir le chatbot depuis n'importe quelle page "
         "d'administration. Les exemples de questions traitées sont nombreux : « CA du "
         "jour », « Top 5 produits du mois », « Livreurs à plus de 10 % de retours ». Une "
         "bascule active le streaming SSE pour afficher la réponse mot à mot.",
         "Interface du chatbot administrateur (n8n)"),
        ("5.8 Interface des conversations et insights chatbot",
         "ChatbotConversationsPage parcourt l'historique des sessions, filtrable par date "
         "et par utilisateur. ChatbotInsightsPage présente les insights proactifs sous "
         "forme de cartes : chaque insight expose un titre, un constat chiffré et une "
         "action suggérée. Un bouton de feedback alimente la qualité des futures "
         "recommandations. L'endpoint /api/admin/chat/kb/refresh permet de régénérer la "
         "base de connaissances qui sert de contexte au router Groq.",
         "Interface des conversations et insights chatbot"),
    ]
    for title, body_text, fig_text in realisations_s5:
        els.append(make_paragraph(STYLES["h3"], title))
        els.append(make_paragraph(None, body_text, first_line_indent_cm=0.6))
        ph = make_paragraph(None, "[ Emplacement réservé pour la capture d'écran ]",
                            align="center")
        for r in ph.findall(qn("w:r")):
            rPr = wel("rPr")
            rPr.append(wel("i"))
            rPr.append(wel("color", val="888888"))
            r.insert(0, rPr)
        els.append(ph)
        els.append(make_caption_paragraph("Figure", fig_text))

    # 6. Tests et validation
    els.append(make_paragraph(STYLES["h2"], "6. Tests et validation"))
    els.append(make_paragraph(None,
        "Les tests vérifient la cohérence des indicateurs par rapport aux données "
        "sources, la conformité des artefacts d'export aux spécifications et la qualité "
        "des réponses du chatbot. Les tests automatisés du projet Web-Api.Tests (xUnit + "
        "EF InMemory) couvrent les modules sensibles (paiements virtuels, devis B2B, "
        "favoris, géofencing). Les tests manuels suivants complètent l'évaluation "
        "fonctionnelle du sprint.",
        first_line_indent_cm=0.6))
    els.append(make_caption_paragraph("Tableau", "Tests fonctionnels du Sprint 5"))
    els.append(make_table(
        ["ID", "Objectif", "Résultat attendu", "Statut"],
        S5_TESTS,
        col_widths_cm=[1.4, 4.5, 7.5, 1.5]))
    els.append(make_paragraph(None, ""))

    # 7. Conclusion
    els.append(make_paragraph(STYLES["h2"], "7. Conclusion"))
    els.append(make_paragraph(None,
        "Le Sprint 5 referme le projet en plaçant l'administrateur au centre d'une "
        "plateforme instrumentée. Les onze tableaux de bord offrent une vision unifiée et "
        "homogène — un contrat unique ProDashboardPageResponseDto et un composant React "
        "unique. Les exports Excel et PDF construits avec ClosedXML et QuestPDF "
        "permettent de matérialiser cette vision dans des supports partageables. Le "
        "chatbot orchestré par n8n, par le pipeline interne Router-Exécuteur-Formatter et "
        "par trois modèles ML.NET, agit comme une couche d'accès conversationnelle (texte "
        "sur le web, voix sur le mobile) qui rapproche les données métier des questions "
        "opérationnelles quotidiennes. Les insights proactifs, calculés en arrière-plan "
        "par Hangfire, anticipent les besoins de l'administrateur. Avec ce cinquième "
        "sprint, la solution atteint un niveau de maturité opérationnelle compatible avec "
        "une démonstration finale et une mise en production progressive.",
        first_line_indent_cm=0.6))

    return els

# ──────────────────────────────────────────────────────────────────────
# Pipeline principal
# ──────────────────────────────────────────────────────────────────────
def main():
    print("[1/6] Ouverture du Copie 0...")
    doc = Document(COPIE0)
    body = doc.element.body

    print("[2/6] Localisation des chapitres 7 et 8 (anciens)...")
    # Trouver les bornes par texte des Titre1 paragraphs
    def style_of(p_el):
        pPr = p_el.find(qn('w:pPr'))
        if pPr is None: return ''
        pStyle = pPr.find(qn('w:pStyle'))
        return pStyle.get(qn('w:val')) if pStyle is not None else ''
    def text_of(p_el):
        return ''.join(t.text or '' for t in p_el.findall('.//' + qn('w:t')))

    indices = {}
    for i, child in enumerate(body):
        if child.tag.endswith('}p'):
            st = style_of(child)
            txt = text_of(child).strip()
            if st == 'Titre1':
                if txt.startswith('Chapitre 7') and 'chap7_h' not in indices:
                    indices['chap7_h'] = i
                elif txt.startswith('Chapitre 8') and 'chap8_h' not in indices:
                    indices['chap8_h'] = i
                elif txt.startswith('Conclusion') and 'conclusion' not in indices:
                    indices['conclusion'] = i
            # detect sectPr (used to find end of chapter sections)
        sectPr = child.find('.//' + qn('w:sectPr')) if child.tag.endswith('}p') else None
        if sectPr is not None:
            # Note the body index for each sectPr
            indices.setdefault('sectPrs', []).append(i)

    print(f"  Chapitre 7 en-tête à l'index {indices['chap7_h']}")
    print(f"  Chapitre 8 en-tête à l'index {indices['chap8_h']}")
    print(f"  Conclusion à l'index {indices['conclusion']}")

    # Find the sectPr that ends chap 7 (= sectPr just before chap 8 header) and
    # the sectPr that ends chap 8 (= sectPr just before conclusion header)
    sect_end_chap7 = None
    sect_end_chap8 = None
    for spi in indices['sectPrs']:
        if indices['chap7_h'] < spi < indices['chap8_h']:
            sect_end_chap7 = spi
        elif indices['chap8_h'] < spi < indices['conclusion']:
            sect_end_chap8 = spi
    print(f"  sectPr end of old Ch.7 = {sect_end_chap7}")
    print(f"  sectPr end of old Ch.8 = {sect_end_chap8}")

    # Capture sectPr templates (we will reuse for new chapters)
    sectPr_chap7_template = deepcopy(body[sect_end_chap7].find('.//' + qn('w:sectPr')))
    sectPr_chap8_template = deepcopy(body[sect_end_chap8].find('.//' + qn('w:sectPr')))

    # The page-break paragraph immediately preceding chap 7 (index chap7_h - 1)
    # is an empty paragraph (its only role is to insert a break). Keep it intact;
    # we'll add our new chapter title with pageBreakBefore so the break is preserved.

    # Calculate ranges to DELETE:
    # - chap 7 body excluding the preceding empty para (chap7_h - 1) AND excluding
    #   the trailing sectPr paragraph (sect_end_chap7).
    #   → delete indices [chap7_h - 1, sect_end_chap7 - 1] inclusive
    # - chap 8 body similarly: [chap8_h - 1, sect_end_chap8 - 1] inclusive
    #
    # Then keep sect_end_chap7 (it will end new chap 7 section)
    # and sect_end_chap8 (it will end new chap 8 section).

    del_chap7_from = indices['chap7_h'] - 1  # the empty break para
    del_chap7_to   = sect_end_chap7 - 1
    del_chap8_from = indices['chap8_h'] - 1
    del_chap8_to   = sect_end_chap8 - 1

    print(f"\n[3/6] Insertion des images (création des relations)...")

    # We add pictures at end of doc using doc.add_picture, then capture rIds, then
    # delete the temp paragraphs.
    diagrams = [
        ("s4_usecase",       DIAG / "s4_usecase.png"),
        ("s4_seq_bctbl",     DIAG / "s4_seq_bctbl.png"),
        ("s4_seq_gps",       DIAG / "s4_seq_gps.png"),
        ("s4_seq_reclamation", DIAG / "s4_seq_reclamation.png"),
        ("s4_seq_transit",   DIAG / "s4_seq_transit.png"),
        ("s5_usecase",       DIAG / "s5_usecase.png"),
        ("s5_seq_chatbot",   DIAG / "s5_seq_chatbot.png"),
        ("s5_seq_export",    DIAG / "s5_seq_export.png"),
    ]

    diag_rids = {}
    temp_paragraphs = []
    for name, path in diagrams:
        p = doc.add_paragraph()
        run = p.add_run()
        pic = run.add_picture(str(path), width=Cm(14.5))
        # capture rId from the drawing
        drawing = run._element.find('.//' + qn('w:drawing'))
        blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        diag_rids[name] = rId
        temp_paragraphs.append(p._element)
        print(f"  {name} → {rId}")

    # Remove the temp picture paragraphs (we have the rIds)
    for tp in temp_paragraphs:
        tp.getparent().remove(tp)

    print(f"\n[4/6] Construction des nouveaux chapitres 7 et 8...")
    new_chap7 = build_chapter_7(diag_rids)
    new_chap8 = build_chapter_8(diag_rids)
    print(f"  Chapitre 7 : {len(new_chap7)} éléments")
    print(f"  Chapitre 8 : {len(new_chap8)} éléments")

    print(f"\n[5/6] Suppression et remplacement...")
    # Re-locate boundaries (indices may have shifted) — actually since temp paragraphs
    # were ADDED then REMOVED, indices should be unchanged. But to be safe, re-find.
    body = doc.element.body
    indices2 = {}
    for i, child in enumerate(body):
        if child.tag.endswith('}p'):
            st = style_of(child)
            txt = text_of(child).strip()
            if st == 'Titre1':
                if txt.startswith('Chapitre 7') and 'chap7_h' not in indices2:
                    indices2['chap7_h'] = i
                elif txt.startswith('Chapitre 8') and 'chap8_h' not in indices2:
                    indices2['chap8_h'] = i
                elif txt.startswith('Conclusion') and 'conclusion' not in indices2:
                    indices2['conclusion'] = i
        if child.tag.endswith('}p'):
            sectPr = child.find('.//' + qn('w:sectPr'))
            if sectPr is not None:
                indices2.setdefault('sectPrs', []).append(i)

    sect_end_chap7 = None
    sect_end_chap8 = None
    for spi in indices2['sectPrs']:
        if indices2['chap7_h'] < spi < indices2['chap8_h']:
            sect_end_chap7 = spi
        elif indices2['chap8_h'] < spi < indices2['conclusion']:
            sect_end_chap8 = spi

    # Delete old chap 8 first (so chap 7 indices remain valid)
    del_8_from = indices2['chap8_h'] - 1
    del_8_to   = sect_end_chap8 - 1
    for i in range(del_8_to, del_8_from - 1, -1):
        body.remove(body[i])

    # Now delete old chap 7 (indices unchanged since we deleted after)
    del_7_from = indices2['chap7_h'] - 1
    del_7_to   = sect_end_chap7 - 1
    for i in range(del_7_to, del_7_from - 1, -1):
        body.remove(body[i])

    # Now find anchors for insertion
    # After deletion, the sectPr that ENDED chap 7 should be at the position
    # right after where chap 7 was deleted.
    # Re-find sectPrs and chap titles
    indices3 = {'sectPrs': []}
    for i, child in enumerate(body):
        if child.tag.endswith('}p'):
            sectPr = child.find('.//' + qn('w:sectPr'))
            if sectPr is not None:
                indices3['sectPrs'].append(i)
            st = style_of(child)
            txt = text_of(child).strip()
            if st == 'Titre1' and txt.startswith('Conclusion'):
                indices3['conclusion'] = i

    # The sectPr that previously ended chap 7 should now be the first sectPr
    # whose body index is < conclusion (since chap 7 was just deleted, sectPr_end_chap7
    # is now where the deletion started)
    # Find the FIRST sectPr that appears after del_7_from - 1 and before del_8_from (now collapsed)
    # Easier: After deleting both chapters, the document structure is:
    # [...chap 6 content + sectPr ending chap 6 (originally at 1115)]
    # [sectPr_chap7_template-bearing paragraph]  ← was at sect_end_chap7
    # [sectPr_chap8_template-bearing paragraph]  ← was at sect_end_chap8
    # [empty break para before conclusion]
    # [Conclusion Titre1]
    # [...]
    # Find the two consecutive sectPr paragraphs that remained (the chap 7 and chap 8 enders)
    # We identify them: they are sectPrs whose index < conclusion.
    sectPrs_before_concl = [i for i in indices3['sectPrs'] if i < indices3['conclusion']]
    # The last 2 are the ones that ended chap 7 and chap 8 (now back-to-back)
    sect_pos_for_chap7_end = sectPrs_before_concl[-2]
    sect_pos_for_chap8_end = sectPrs_before_concl[-1]

    print(f"  Insertion chap 7 avant index {sect_pos_for_chap7_end}")
    # Insert new_chap7 elements BEFORE sect_pos_for_chap7_end
    insert_anchor = body[sect_pos_for_chap7_end]
    for el in new_chap7:
        insert_anchor.addprevious(el)

    # Re-find chap 8 anchor (indices have shifted)
    indices4 = {'sectPrs': []}
    for i, child in enumerate(body):
        if child.tag.endswith('}p'):
            sectPr = child.find('.//' + qn('w:sectPr'))
            if sectPr is not None:
                indices4['sectPrs'].append(i)
            st = style_of(child)
            txt = text_of(child).strip()
            if st == 'Titre1' and txt.startswith('Conclusion'):
                indices4['conclusion'] = i
    sectPrs_before_concl = [i for i in indices4['sectPrs'] if i < indices4['conclusion']]
    sect_pos_for_chap8_end = sectPrs_before_concl[-1]
    print(f"  Insertion chap 8 avant index {sect_pos_for_chap8_end}")
    insert_anchor = body[sect_pos_for_chap8_end]
    for el in new_chap8:
        insert_anchor.addprevious(el)

    print(f"\n[6/6] Activation de la mise à jour automatique des champs...")
    # Modify settings.xml: add w:updateFields w:val="true"
    settings_part = doc.settings.element
    # Remove existing updateFields if any
    for uf in settings_part.findall(qn('w:updateFields')):
        settings_part.remove(uf)
    uf = wel('updateFields', val='true')
    settings_part.insert(0, uf)

    # Save
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"\n✅ Rapport final généré : {OUTPUT}")
    print(f"   Taille : {OUTPUT.stat().st_size // 1024} Ko")

if __name__ == "__main__":
    main()
