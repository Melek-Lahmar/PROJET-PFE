#!/usr/bin/env python3
"""Génère les chapitres 7 et 8 du Rapport PFE : Sprints 4 et 5.
Diagrammes PlantUML rendus en PNG puis insérés dans un document Word
respectant strictement le plan fourni par l'étudiant.
"""

import os
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

DIAG_DIR = Path("/tmp/diagrams_final")
PLANTUML_JAR = "/tmp/plantuml.jar"
OUTPUT_DOCX = Path("/home/user/PROJET-PFE/Rapport/Chapitres_7_8_Corriges.docx")

# ─────────────────────────────────────────────────────────────────────
# Skin commun PlantUML — sobre, lisible pour un rapport académique
# ─────────────────────────────────────────────────────────────────────
SKIN = """
skinparam backgroundColor white
skinparam handwritten false
skinparam shadowing false
skinparam roundCorner 4
skinparam DefaultFontName Arial
skinparam DefaultFontSize 12
skinparam TitleFontStyle bold
skinparam TitleFontSize 14
skinparam ArrowColor #355070
skinparam ActorBorderColor #355070
skinparam ActorBackgroundColor #EAF1F8
skinparam UsecaseBorderColor #355070
skinparam UsecaseBackgroundColor #F4F9FD
skinparam RectangleBorderColor #355070
skinparam RectangleBackgroundColor #FBFCFE
skinparam ParticipantBorderColor #355070
skinparam ParticipantBackgroundColor #EAF1F8
skinparam SequenceLifeLineBorderColor #6C7A93
skinparam SequenceMessageAlign center
skinparam NoteBackgroundColor #FFF8DC
skinparam NoteBorderColor #B8860B
"""

PUML = {}

# ═════════════════════════════════════════════════════════════════════
#  SPRINT 4 — Diagrammes
# ═════════════════════════════════════════════════════════════════════

PUML["s4_usecase.puml"] = """@startuml
title Diagramme de cas d'utilisation — Sprint 4
__SKIN__
left to right direction
skinparam packageStyle rectangle

actor "Confirmatrice" as CF
actor "Livreur COD" as LV
actor "Livreur Transit" as LT
actor "Client" as CL
actor "Superviseur" as SV

rectangle "Sprint 4 — Logistique, livraison & réclamations" {
  package "Confirmation & documents Sage X3" {
    usecase "Consulter les bons de commande" as UC1
    usecase "Modifier le statut d'un BC" as UC2
    usecase "Transformer BC en BL\\n(+ sync Sage X3)" as UC3
    usecase "Affecter automatiquement\\nun livreur" as UC4
    usecase "Traiter les devis B2B" as UC5
  }

  package "Livraison COD & GPS temps réel" {
    usecase "Consulter le pool\\nde livraisons (zone)" as UC6
    usecase "Prendre en charge\\nune livraison" as UC7
    usecase "Démarrer la diffusion GPS\\n(start-heading)" as UC8
    usecase "Mettre à jour le statut\\n(Livré / Retour / Reporté)" as UC9
    usecase "Gérer la caisse COD\\n(encaissement auto)" as UC10
    usecase "Optimiser la tournée\\n(plus-proche-voisin)" as UC11
    usecase "Demander une correction\\nd'adresse / téléphone" as UC12
  }

  package "Transit inter-dépôts" {
    usecase "Consulter mes missions" as UC13
    usecase "Scanner pickup\\nau dépôt source" as UC14
    usecase "Scanner delivery\\nau dépôt destination" as UC15
    usecase "Réception partielle" as UC16
  }

  package "Réclamations temps réel" {
    usecase "Créer une réclamation" as UC17
    usecase "Suivre sa commande\\n(TrackingStateCard)" as UC18
    usecase "Prendre en charge un cas" as UC19
    usecase "Chatter en temps réel\\n(SignalR)" as UC20
  }

  package "Supervision" {
    usecase "Superviser les livreurs\\nactifs" as UC21
    usecase "Consulter la heatmap\\ndes échecs (90 j)" as UC22
    usecase "Réaffecter une mission\\nde transit" as UC23
  }
}

CF --> UC1
CF --> UC2
CF --> UC3
CF --> UC4
CF --> UC5
CF --> UC19
CF --> UC20

LV --> UC6
LV --> UC7
LV --> UC8
LV --> UC9
LV --> UC10
LV --> UC11
LV --> UC12

LT --> UC13
LT --> UC14
LT --> UC15
LT --> UC16

CL --> UC17
CL --> UC18
CL --> UC20

SV --> UC21
SV --> UC22
SV --> UC23

UC3 .> UC4 : <<include>>
UC8 .> UC7 : <<include>>
UC10 .> UC9 : <<include>>
UC20 .> UC17 : <<include>>
UC22 .> UC21 : <<extend>>
@enduml
"""

PUML["s4_seq_bctbl.puml"] = """@startuml
title Diagramme de séquence — Transformation BC en BL et affectation automatique
__SKIN__
participant "Confirmatrice" as CF
participant "React\\n(ConfirmateurOrders\\nDetailsPage)" as R
participant "ConfirmateurController\\n/api/confirmateur" as CC
participant "AssignmentService" as AS
participant "SageX3Client" as SAGE
database  "SQL Server\\nF_DOCENTETE\\nF_DOCLIGNE\\nF_LIVRAISON" as DB

CF -> R : Clic « Transformer en BL »
R -> CC : POST /commandes/{piece}/transform-to-bl
activate CC

CC -> DB : SELECT F_DOCENTETE\\n(WHERE DO_Piece = piece AND DO_Type = 0)
DB --> CC : Bon de commande (BC)

alt BC introuvable
  CC --> R : 404 NotFound
else BL déjà créé pour ce BC
  CC -> DB : SELECT BL existant (DO_Type = 1)
  DB --> CC : Référence BL
  CC --> R : 200 OK { blPiece, alreadyExists = true }
else Création du BL
  CC -> DB : BEGIN TRANSACTION
  CC -> DB : INSERT F_DOCENTETE\\n(DO_Type = 1, DO_Piece = "BL" + yyMMddHHmm + n)
  CC -> DB : INSERT F_DOCLIGNE (copie des lignes)
  CC -> DB : INSERT F_LIVRAISON (LI_Statut = 0 Confirmé)
  CC -> DB : UPDATE F_DOCENTETE (BC) SET DO_Valide = CONFIRME
  CC -> DB : COMMIT

  CC -> AS : ChoisirLivreur(zone, codeArticle…)
  AS -> DB : SELECT livreurs (zones couvertes,\\nstock disponible, file la plus courte)
  AS -> AS : Filtre Haversine zone\\n+ tri par charge croissante
  AS --> CC : userId du livreur ou null

  alt Livreur trouvé
    CC -> DB : UPDATE F_LIVRAISON SET LivreurUserId
    CC -> CC : SignalR « NouvelleLivraisonAffectee »\\n→ user-{livreurId}
  end

  CC -> SAGE : POST /api/sage/documents (BL)
  SAGE --> CC : { sageSuccess, httpStatus }
  CC --> R : 201 Created { blPiece, sageSent, sageSuccess }
end

deactivate CC
R --> CF : Notification + redirection vers la liste des BL

note over CC, SAGE
  L'envoi vers Sage X3 est **non bloquant** :
  le BL est créé localement même si Sage est
  indisponible. Le statut sageSuccess est
  historisé dans le tableau de bord de synchronisation.
end note
@enduml
"""

PUML["s4_seq_gps.puml"] = """@startuml
title Diagramme de séquence — Livraison active, GPS temps réel et encaissement COD
__SKIN__
actor "Livreur" as L
participant "App Flutter\\n(LivreurStatsScreen)" as F
participant "LivreurActive\\nDeliveryController" as LADC
participant "LivreurController" as LC
participant "ReclamationHub\\n(SignalR)" as SR
participant "Client" as CL
database  "SQL Server\\nF_LIVRAISON\\nF_LIVREUR_POSITION" as DB

== Démarrage de la livraison active ==
L -> F : Démarre la livraison
F -> LADC : POST /api/livreur/orders/{piece}/start-heading
LADC -> DB : UPDATE F_LIVRAISON\\nSET IsActiveDelivery = true, IsBroadcasting = true
LADC -> SR : DeliveryStarted → groupe client-{userId}
SR -> CL : État → HEADING_TO_YOU
LADC --> F : 200 OK

== Diffusion GPS (toutes les 10 s) ==
loop Geolocator Flutter
  F -> F : getCurrentPosition()
  F -> LADC : POST /api/livreur/location/ping\\n{lat, lng, accuracy, ts, clientActionId}
  LADC -> DB : UPSERT F_LIVREUR_POSITION\\n(UpdatedAt = now)
  LADC -> DB : INSERT F_LIVREUR_POSITION_HISTORY
  LADC -> LADC : Distance Haversine livreur→client
  LADC -> SR : LivreurPositionUpdate(lat,lng,etaMin,distKm)\\n→ client-{userId}
  SR -> CL : MAJ TrackingStateCard\\n(GPS + ETA + indicateur de fraîcheur)
  opt distance < 0.5 km
    LADC -> SR : ProximityAlert → client-{userId}
  end
end

note over F
  Si le réseau est coupé, les positions
  sont mises en file locale puis envoyées
  via POST /location/ping-batch avec un
  ClientActionId garantissant l'idempotence.
end note

== Encaissement COD à la livraison ==
L -> F : Marquer « Livré »
F -> LC : PUT /api/livreur/orders/{piece}/status\\n{LI_Statut = 2 (Livre)}
LC -> DB : UPDATE F_LIVRAISON\\nEncaisse = true, MontantEncaisse = DO_NetAPayer,\\nEncaisseAt = UtcNow
LC -> SR : StatutCommandeChange → client + livreur
LC -> LADC : POST /stop-heading (broadcast OFF)
LC --> F : 200 OK
F --> L : Confirmation
@enduml
"""

PUML["s4_seq_reclamation.puml"] = """@startuml
title Diagramme de séquence — Réclamation client avec SignalR (grâce 5 s)
__SKIN__
actor "Client" as C
participant "App Flutter" as F
participant "Reclamations\\nController" as RC
participant "ReclamationHub\\n(SignalR)" as SH
actor "Confirmatrice" as CF
database  "SQL Server\\nF_RECLAMATION\\nF_CONFIRMATRICE_SESSION" as DB

C -> F : Crée une réclamation\\n(motif + photos)
F -> RC : POST /api/reclamations\\n{motif: COLIS_ENDOMMAGE, doPiece, photos}
RC -> DB : INSERT F_RECLAMATION\\n(Statut = ENVOYEE, TypeCas = RECLAMATION, VisibleClient = true)
RC -> SH : Diffuser NouveauCas\\n→ groupe "confirmateurs"
SH -> CF : Réclamation visible (sans rafraîchissement)

CF -> RC : POST /reclamations/{id}/prendre-en-charge
RC -> DB : UPDATE confirmateurId, Statut = EN_COURS
RC -> DB : INSERT F_CONFIRMATRICE_SESSION (ouverture)
RC -> SH : CommandePriseEnCharge\\n→ "confirmateurs"
RC -> SH : CasAssigne → client-{userId}
SH -> C : Chat ouvert + identité confirmatrice

== Chat temps réel ==
loop Echange Client ↔ Confirmatrice
  C -> RC : POST /reclamations/{id}/messages
  RC -> DB : INSERT F_RECLAMATION_MESSAGE
  RC -> SH : NouveauMessage → "confirmateurs"
  CF -> RC : POST /reclamations/{id}/messages
  RC -> DB : INSERT F_RECLAMATION_MESSAGE
  RC -> SH : NouveauMessage → client-{userId}
end

== Déconnexion : période de grâce de 5 secondes ==
note over SH
  OnDisconnectedAsync décrémente le compteur
  de connexions de la confirmatrice. Si c'est
  la dernière connexion, le cas n'est pas
  libéré immédiatement : un timer de
  **GracePeriod = TimeSpan.FromSeconds(5)**
  est armé. Si la confirmatrice se reconnecte
  dans cet intervalle (changement Wi-Fi / 4G),
  le timer est annulé, le cas reste verrouillé.
end note

alt Reconnexion ≤ 5 s
  CF -> SH : Reconnect (compteur++)
  SH -> SH : Annule le timer de libération
else Pas de reconnexion
  SH -> DB : UPDATE Statut = ENVOYEE\\n(confirmateurId = null)
  SH -> SH : Ferme F_CONFIRMATRICE_SESSION
  SH -> SH : Diffuse CommandeLiberee\\n→ "confirmateurs"
end

== Clôture ==
CF -> RC : PUT /reclamations/{id}/status (CLOTUREE)
RC -> DB : UPDATE Statut = CLOTUREE
RC -> SH : StatutCasChange → client-{userId}
SH -> C : Notification finale
@enduml
"""

PUML["s4_seq_transit.puml"] = """@startuml
title Diagramme de séquence — Transit inter-dépôts par double scan
__SKIN__
actor "Livreur Transit" as LT
participant "App Flutter\\n(transit_mission_\\ndetails_screen)" as F
participant "mobile_scanner\\n(transit_barcode_\\nscanner_screen)" as SC
participant "TransitController\\n/api/transit" as TC
participant "Supervisor\\nHub" as SH
database  "SQL Server\\nF_TRANSFERT\\nF_TRANSFERT_AUDIT_LOG" as DB

LT -> F : Ouvre la mission
F -> TC : GET /my-missions/{id}
TC -> DB : SELECT F_TRANSFERT (lignes attendues)
DB --> TC : Articles, quantités, dépôts
TC --> F : MissionDto (statut EN_ATTENTE_TRANSIT)

== Scan PICKUP (dépôt source) ==
LT -> SC : Active le scanner caméra
SC -> SC : Decode barcode (EAN13)
SC -> TC : POST /scan-pickup\\n{missionId, barcode, gps}
TC -> DB : SELECT ligne (ArRef ↔ barcode,\\nSourceDepotNo correspondant)
alt Article appartient à la mission
  TC -> DB : UPDATE statut = EN_TRANSIT,\\nPickedUpAt, PickupGpsLat/Lng
  TC -> DB : INSERT F_TRANSFERT_AUDIT_LOG (PICKUP)
  TC --> SC : 200 { ok = true, son ✓ }
  SC --> LT : Vibration + son confirmation
else Article inconnu / mauvais dépôt
  TC --> SC : 400 { ok = false, raison }
  SC --> LT : Son d'erreur + message
end

note over TC
  Fenêtre d'annulation : POST /revert-pickup
  autorisé pendant **RevertWindowMinutes = 10**
  après PickedUpAt, avec écriture d'audit.
end note

== Scan DELIVERY (dépôt destination) ==
LT -> SC : Re-scanne à l'arrivée
SC -> TC : POST /scan-delivery\\n{missionId, barcode, gps}
TC -> DB : UPDATE statut = RECU_DEPOT_DESTINE,\\nDeliveredAt, DeliveryGpsLat/Lng
TC -> DB : INSERT F_TRANSFERT_AUDIT_LOG (DELIVERY)
TC --> F : 200 OK

alt Réception partielle (quantité < attendue)
  LT -> F : Saisir reçus = N (< quantite)
  F -> TC : POST /missions/{id}/scan-partial\\n{receivedQty}
  TC -> DB : UPDATE statut = TRANSIT_PARTIELLEMENT_RECU
  TC -> SH : NouvelleAlerte → "superviseurs"
end

opt Toutes les lignes RECU_DEPOT_DESTINE
  TC -> DB : UPDATE F_TRANSFERT_ENTETE statut = TRANSIT_TERMINE
  TC --> F : 200 OK
  F --> LT : Mission clôturée
end
@enduml
"""

# ═════════════════════════════════════════════════════════════════════
#  SPRINT 5 — Diagrammes
# ═════════════════════════════════════════════════════════════════════

PUML["s5_usecase.puml"] = """@startuml
title Diagramme de cas d'utilisation — Sprint 5
__SKIN__
left to right direction
skinparam packageStyle rectangle

actor "Administrateur" as A
actor "Groq LLM\\n(llama-3.3-70b)" as LLM
actor "n8n Workflow" as N8N
actor "Sage X3" as SX

rectangle "Sprint 5 — Pilotage, exports & chatbot" {
  package "Tableaux de bord multi-dimensionnels" {
    usecase "Vue d'ensemble (KPI, alertes)" as UC1
    usecase "Ventes (CA, top produits)" as UC2
    usecase "Commandes (statuts, délais)" as UC3
    usecase "Produits & stocks" as UC4
    usecase "Logistique (zones, heatmap)" as UC5
    usecase "Livreurs (performance, COD)" as UC6
    usecase "Clients (B2B/B2C, fidélité)" as UC7
    usecase "Réclamations (motifs, SLA)" as UC8
    usecase "Confirmatrices (cas, temps)" as UC9
    usecase "Synchronisation Sage X3" as UC10
    usecase "Insights stratégiques" as UC11
  }

  package "Exports & rapports" {
    usecase "Exporter commandes\\n(PDF / Excel)" as UC12
    usecase "Exporter réclamations\\n(PDF / Excel)" as UC13
  }

  package "Chatbot administrateur" {
    usecase "Poser une question NL\\n(/ask)" as UC14
    usecase "Streaming SSE\\n(/ask-stream)" as UC15
    usecase "Historique des conversations" as UC16
    usecase "Insights proactifs" as UC17
    usecase "Rafraîchir la base de connaissances" as UC18
  }
}

A --> UC1
A --> UC2
A --> UC3
A --> UC4
A --> UC5
A --> UC6
A --> UC7
A --> UC8
A --> UC9
A --> UC10
A --> UC11
A --> UC12
A --> UC13
A --> UC14
A --> UC15
A --> UC16
A --> UC17
A --> UC18

N8N --> UC14 : <<orchestre>>
LLM --> UC14 : <<formate>>
LLM --> UC17 : <<génère>>
SX --> UC10 : <<source>>

UC15 .> UC14 : <<extend>>
UC11 .> UC1  : <<include>>
UC17 .> UC14 : <<include>>
@enduml
"""

PUML["s5_seq_chatbot.puml"] = """@startuml
title Diagramme de séquence — Chatbot administrateur (React → n8n → Groq)
__SKIN__
actor "Administrateur" as A
participant "ChatbotSandboxPage\\n(React)" as R
participant "Webhook n8n\\n/admin-chat-v3" as W
participant "Node Détecter langue\\n(FR / AR / Tounsi)" as LANG
participant "AdminChatController\\n/api/admin/chat/ask" as CC
participant "ChatOrchestrator\\n(Router → Exécuteur → Formatter)" as ORCH
participant "GroqClient\\nllama-3.3-70b" as LLM
database  "SQL Server\\nF_CHATBOT_*,\\nF_DOCENTETE…" as DB

A -> R : « Quels livreurs ont le plus de retours ? »
R -> W : POST /admin-chat-v3\\n{question, sessionId, locale}
W -> LANG : Détection langue
LANG --> W : language = "fr"

W -> CC : POST /api/admin/chat/ask\\n(Bearer JWT + X-Chat-Api-Key)
CC -> ORCH : AskAsync(question, sessionId, userId)

ORCH -> LLM : Router (Groq) :\\n→ action ∈ {KB, QUERY, ANALYZE, PREDICT}
LLM --> ORCH : action = "QUERY"\\n(metric = retours, subject = livreurs)

ORCH -> DB : SELECT livreurs + GROUP BY\\nstats retour 30 j
DB --> ORCH : Rows + agrégats

ORCH -> LLM : Formatter (Groq) :\\ngénère message FR + chart spec
LLM --> ORCH : { message, rows, chart, suggestions }

ORCH --> CC : ChatAskResponseDto
CC -> DB : INSERT F_CHATBOT_MESSAGE\\n(question, response, sessionId)
CC --> W : 200 OK + payload

W -> W : Node « Insights proactifs »\\nGET /api/admin/chat/insights/pending
W -> W : Node « Merge réponse + insights »
W --> R : 200 OK { message, rows, chart, insights[] }
R --> A : Affichage texte + tableau + graphe Recharts

note over R, ORCH
  Variante streaming : POST /ask-stream renvoie
  un flux SSE en quatre phases — **routing → data
  → chunks → done** — pour un rendu mot-à-mot
  côté React (EventSource).
end note
@enduml
"""

PUML["s5_seq_export.puml"] = """@startuml
title Diagramme de séquence — Export Excel / PDF (QuestPDF + ClosedXML)
__SKIN__
actor "Administrateur" as A
participant "AdminSummaryPage\\n(React)" as R
participant "AdminSummary\\nController" as ASC
participant "ExportService" as ES
database  "SQL Server" as DB

A -> R : Clic « Exporter »\\n(type, période, format)
R -> ASC : GET /api/admin/orders/export\\n?format=xlsx&period=30d
ASC -> ASC : ParsePeriod(period) → (from, to)
ASC -> DB : SELECT F_DOCENTETES\\nWHERE cbCreation ∈ [from, to[\\n.Take(MaxRows = 10 000)
DB --> ASC : Lignes (DO_Piece, DO_Tiers,\\nDO_Valide, DO_Date, ville, DO_NetAPayer)

alt format == "xlsx"
  ASC -> ES : ExportToExcel(headers, rows)
  ES -> ES : XLWorkbook (ClosedXML) :\\nentêtes gras, AdjustToContents
  ES --> ASC : byte[] (.xlsx)
  ASC --> R : 200 + Content-Disposition\\napplication/vnd…spreadsheetml
else format == "pdf"
  ASC -> ES : ExportToPdf(title, period, headers, rows)
  ES -> ES : Document.Create (QuestPDF) :\\nA4, en-tête, table, pagination
  ES --> ASC : byte[] (.pdf)
  ASC --> R : 200 + application/pdf
end
R --> A : Téléchargement direct du fichier

note over ES
  Plafond **MaxRows = 10 000** lignes par export
  (avertissement utilisateur au-delà). Service
  générique réutilisé par /reclamations/export
  et toute future ressource via (headers, rows).
end note
@enduml
"""

# ─────────────────────────────────────────────────────────────────────
# 1) Écrire les .puml
# ─────────────────────────────────────────────────────────────────────
def write_puml():
    DIAG_DIR.mkdir(parents=True, exist_ok=True)
    for name, content in PUML.items():
        (DIAG_DIR / name).write_text(content.replace("__SKIN__", SKIN), encoding="utf-8")
        print(f"  wrote {name}")

# ─────────────────────────────────────────────────────────────────────
# 2) Générer les PNG via PlantUML
# ─────────────────────────────────────────────────────────────────────
def render_pngs():
    print("  Rendu PlantUML…")
    res = subprocess.run(
        ["java", "-jar", PLANTUML_JAR, "-tpng", "-charset", "UTF-8", str(DIAG_DIR / "*.puml")],
        capture_output=True, text=True,
    )
    if res.stdout: print(res.stdout)
    if res.stderr: print(res.stderr)
    pngs = sorted(DIAG_DIR.glob("*.png"))
    print(f"  {len(pngs)} PNGs : {[p.name for p in pngs]}")
    return pngs

# ─────────────────────────────────────────────────────────────────────
# 3) Construire le docx
# ─────────────────────────────────────────────────────────────────────
NAVY = RGBColor(0x35, 0x50, 0x70)

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(text)
    run.font.color.rgb = NAVY
    run.font.size = Pt(18)
    run.bold = True

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(text)
    run.font.color.rgb = NAVY
    run.font.size = Pt(14)
    run.bold = True

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    run = p.add_run(text)
    run.font.color.rgb = NAVY
    run.font.size = Pt(12)
    run.bold = True

def add_h4(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 4"]
    run = p.add_run(text)
    run.font.color.rgb = NAVY
    run.bold = True

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs: run.font.size = Pt(11)
    return p

def add_figure(doc, png_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(png_path), width=Cm(15.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY

def add_placeholder(doc, label):
    """Cadre vide pour future capture d'écran."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(15)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(f"\n\n[ Emplacement réservé — {label} ]\n\n")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    # bordures grises pointillées
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'dashed')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:color'), 'AAAAAA')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY

def style_table(table):
    table.style = "Light Grid Accent 1"
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

# ─────────────────────────────────────────────────────────────────────
#  CONTENU — Chapitre 7 : Sprint 4
# ─────────────────────────────────────────────────────────────────────
S4_BACKLOG = [
    ("US4.1",  "Consulter les BC en attente de confirmation",
     "Filtrage par statut (EN_ATTENTE / CONFIRME / EN_LIVRAISON / LIVRE), pagination, recherche par référence ou client.", "Melek Lahmar", "3 j"),
    ("US4.2",  "Consulter le détail d'une commande",
     "Affichage des lignes (F_DOCLIGNE), des informations client enrichies, des montants TTC et du mode de paiement.", "Tawfik Siala", "2 j"),
    ("US4.3",  "Modifier le statut d'un BC",
     "PUT /api/confirmateur/commandes/{piece}/status avec workflow contrôlé (DO_Valide 0 → 3).", "Melek Lahmar", "3 j"),
    ("US4.4",  "Transformer un BC en BL",
     "POST /transform-to-bl : création F_DOCENTETE DO_Type=1, copie des lignes, contrôle anti-doublon, transaction SQL.", "Melek Lahmar", "4 j"),
    ("US4.5",  "Exporter un BL vers Sage X3",
     "Sérialisation SageDocEntetePayload, POST asynchrone non bloquant, journalisation du sageSuccess.", "Tawfik Siala", "3 j"),
    ("US4.6",  "Traiter les devis B2B côté confirmateur",
     "Endpoints /confirmateur/devis : consultation, transformation devis → BC, suivi des révisions.", "Melek Lahmar", "3 j"),
    ("US4.7",  "Consulter les livraisons disponibles",
     "GET /api/livreur/orders/available filtré par gouvernorat / délégation du livreur connecté.", "Tawfik Siala", "3 j"),
    ("US4.8",  "Prendre en charge une livraison",
     "POST /assign : création F_LIVRAISON avec gestion du conflit 409 (déjà prise par un autre livreur).", "Tawfik Siala", "2 j"),
    ("US4.9",  "Démarrer et arrêter la diffusion GPS",
     "/start-heading + /location/ping (10 s) + /ping-batch idempotent, hub SignalR ReclamationHub.", "Melek Lahmar", "4 j"),
    ("US4.10", "Mettre à jour le statut d'une livraison",
     "Codes LI_Statut : 1 EnLivraison, 2 Livre (auto-COD), 3 Retour, 4 Depot, 5 Reporte ; validation des motifs.", "Tawfik Siala", "3 j"),
    ("US4.11", "Gérer la caisse COD",
     "Encaissement automatique au statut 2 (Encaisse, MontantEncaisse, EncaisseAt), remise au dépôt et historique.", "Melek Lahmar", "3 j"),
    ("US4.12", "Créer et suivre une réclamation",
     "POST /api/reclamations (motif, photos), SignalR NouveauCas vers les confirmatrices, suivi via TrackingStateCard.", "Tawfik Siala", "3 j"),
    ("US4.13", "Prendre en charge et traiter une réclamation",
     "Verrou pessimiste avec période de grâce de 5 s à la déconnexion, chat temps réel NouveauMessage.", "Melek Lahmar", "3 j"),
    ("US4.14", "Demander une correction d'adresse / téléphone",
     "Type DEMANDE_LIVREUR escaladé vers la confirmatrice ; appliquer la correction sur F_DOCENTETE.", "Tawfik Siala", "2 j"),
    ("US4.15", "Superviser les livreurs actifs",
     "Carte temps réel, heatmap 90 j (LivreurMapController), alertes (retard, échec, zone non couverte).", "Melek Lahmar", "4 j"),
    ("US4.16", "Gérer les missions de transit inter-dépôts",
     "Double scan pickup/delivery (mobile_scanner), réception partielle, audit log et annulation 10 min.", "Tawfik Siala", "4 j"),
    ("US4.17", "Optimiser la tournée",
     "Algorithme du plus-proche-voisin + Haversine, ETA cumulatif à 35 km/h (LivreurMapController.OptimizeTournee).", "Melek Lahmar", "3 j"),
]

S4_TESTS = [
    ("TF4.1", "Consultation BC confirmatrice", "Les BC s'affichent avec leur statut, leur client et leur montant.", "Validé"),
    ("TF4.2", "Transformation BC → BL sans doublon", "Création du BL et anti-doublon vérifié ; le BL existant est renvoyé.", "Validé"),
    ("TF4.3", "Sage X3 indisponible (non bloquant)", "Le BL est créé localement ; sageSuccess = false dans la réponse et le dashboard.", "Validé"),
    ("TF4.4", "Affectation automatique du livreur", "Le livreur retenu est celui de zone à charge minimale.", "Validé"),
    ("TF4.5", "Pool filtré par zone livreur", "Seules les livraisons de son gouvernorat / délégation apparaissent.", "Validé"),
    ("TF4.6", "Prise en charge concurrentielle", "Le deuxième livreur reçoit 409 Conflict ; aucun doublon F_LIVRAISON.", "Validé"),
    ("TF4.7", "Diffusion GPS SignalR", "Position reçue côté client en < 1 s, fraîcheur < 30 s.", "Validé"),
    ("TF4.8", "Mode hors-ligne /ping-batch", "Positions envoyées en lot à la reconnexion, idempotence par ClientActionId.", "Validé"),
    ("TF4.9", "Encaissement COD automatique", "Encaisse=true et MontantEncaisse=DO_NetAPayer dès LI_Statut=2.", "Validé"),
    ("TF4.10", "TrackingStateCard adaptative", "Bascule AT_DEPOT → IN_DELIVERY_QUEUE → HEADING_TO_YOU → TERMINAL conforme.", "Validé"),
    ("TF4.11", "Réclamation SignalR NouveauCas", "Réclamation visible chez la confirmatrice sans rechargement.", "Validé"),
    ("TF4.12", "Grâce de 5 s à la déconnexion", "Reconnexion ≤ 5 s : le cas reste verrouillé ; > 5 s : libération + CommandeLiberee.", "Validé"),
    ("TF4.13", "Scan pickup transit", "L'article passe à EN_TRANSIT, PickedUpAt et GPS enregistrés.", "Validé"),
    ("TF4.14", "Scan delivery + réception partielle", "TRANSIT_PARTIELLEMENT_RECU avec alerte au superviseur.", "Validé"),
    ("TF4.15", "Optimisation tournée plus-proche-voisin", "Ordre minimisant la distance totale, ETA cumulatifs cohérents.", "Validé"),
    ("TF4.16", "Heatmap 90 jours", "Cellules ≈ 500 m, intensité = ratio (retour+report)/livraisons.", "Validé"),
]

S5_BACKLOG = [
    ("US5.1",  "Tableau de bord d'ensemble", "GET /api/dashboard/overview : KPI agrégés (CA, livraisons, réclamations, alertes).", "Melek Lahmar", "3 j"),
    ("US5.2",  "Dashboard ventes", "GET /api/dashboard/sales : CA, top produits, séries temporelles.", "Tawfik Siala", "3 j"),
    ("US5.3",  "Dashboard commandes", "GET /api/dashboard/orders : répartition par statut, délais moyens.", "Melek Lahmar / Tawfik Siala", "3 j"),
    ("US5.4",  "Dashboard produits & stocks", "GET /api/dashboard/products, /stock, /depots : ruptures, top vendus.", "Melek Lahmar", "3 j"),
    ("US5.5",  "Dashboard logistique", "GET /api/dashboard/logistics : taux livraison, zones, heatmap.", "Tawfik Siala", "3 j"),
    ("US5.6",  "Dashboard livreurs", "GET /api/dashboard/drivers : performances, classement, caisse COD.", "Melek Lahmar / Tawfik Siala", "3 j"),
    ("US5.7",  "Dashboard clients", "GET /api/dashboard/clients : B2B/B2C, fidélité, panier moyen.", "Melek Lahmar", "3 j"),
    ("US5.8",  "Dashboard réclamations", "GET /api/dashboard/reclamations : motifs, SLA, taux de résolution.", "Tawfik Siala", "3 j"),
    ("US5.9",  "Dashboard confirmatrices", "GET /api/dashboard/confirmateur : cas actifs, temps moyen.", "Melek Lahmar / Tawfik Siala", "3 j"),
    ("US5.10", "Dashboard sync Sage X3", "GET /api/dashboard/admin-sync : santé des données, dernières synchros.", "Melek Lahmar", "3 j"),
    ("US5.11", "Insights stratégiques", "GET /api/dashboard/strategic-insights : tendances et recommandations.", "Tawfik Siala", "2 j"),
    ("US5.12", "Export Excel (ClosedXML)", "GET /api/admin/orders/export?format=xlsx : .xlsx jusqu'à 10 000 lignes.", "Melek Lahmar", "2 j"),
    ("US5.13", "Export PDF (QuestPDF)", "GET /api/admin/orders/export?format=pdf : .pdf paginé A4.", "Tawfik Siala", "2 j"),
    ("US5.14", "Chatbot administrateur /ask", "POST /api/admin/chat/ask : pipeline Router → Exécuteur → Formatter (Groq).", "Tawfik Siala", "3 j"),
    ("US5.15", "Streaming SSE /ask-stream", "Phases routing → data → chunks → done pour rendu mot-à-mot.", "Melek Lahmar", "2 j"),
    ("US5.16", "Workflow n8n orchestrateur", "admin-chatbot-workflow-v3.json : 5 nœuds (webhook → langue → backend → insights → merge → respond).", "Tawfik Siala", "2 j"),
    ("US5.17", "Conversations & insights chatbot", "Historique F_CHATBOT_MESSAGE, insights F_CHATBOT_INSIGHTS.", "Melek Lahmar / Tawfik Siala", "2 j"),
    ("US5.18", "Livrables & démonstration", "README, Swagger, vidéo de démo et rapport final.", "Melek Lahmar", "2 j"),
]

S5_TESTS = [
    ("TF5.1",  "Cohérence KPI vue d'ensemble", "Les KPI correspondent à un calcul SQL direct sur les mêmes filtres.", "Validé"),
    ("TF5.2",  "Dashboard ventes — CA", "Les montants reflètent uniquement les commandes livrées (LI_Statut=2).", "Validé"),
    ("TF5.3",  "Dashboard logistique — taux", "Le taux de livraison = livrées / (livrées + retours + reportées).", "Validé"),
    ("TF5.4",  "Dashboard réclamations — SLA", "Le délai de clôture moyen est cohérent avec F_RECLAMATION.CreatedAt → CloturedAt.", "Validé"),
    ("TF5.5",  "Dashboard sync Sage X3", "L'état d'intégrité (clients sans adresse, articles sans famille) est exact.", "Validé"),
    ("TF5.6",  "Export Excel /orders/export", "Fichier .xlsx ouvrable, entêtes en gras, ≤ 10 000 lignes.", "Validé"),
    ("TF5.7",  "Export PDF /orders/export", "PDF A4 paginé, en-tête + période + total cohérent.", "Validé"),
    ("TF5.8",  "Plafond MaxRows", "Au-delà de 10 000 lignes, le service tronque et signale l'avertissement.", "Validé"),
    ("TF5.9",  "Chatbot /ask question simple", "« CA du jour ? » → réponse FR + chiffre correct.", "Validé"),
    ("TF5.10", "Chatbot /ask question contextuelle", "« Et hier ? » conserve le sujet via sessionId.", "Validé"),
    ("TF5.11", "Streaming SSE /ask-stream", "Phases reçues dans l'ordre routing → data → chunks → done.", "Validé"),
    ("TF5.12", "Workflow n8n", "Les 5 nœuds s'enchaînent correctement, insights fusionnés.", "Validé"),
    ("TF5.13", "Insights F_CHATBOT_INSIGHTS", "Le bandeau insights affiche les lignes DismissedAt = null.", "Validé"),
    ("TF5.14", "AlertPanel sévérité", "Les alertes critiques sont mises en évidence (rouge).", "Validé"),
]

def build_chap7(doc):
    # === Page break + titre chapitre
    doc.add_page_break()
    add_h1(doc, "Chapitre 7 : Sprint 4 — Logistique, livraison et réclamations")

    # 1. Introduction
    add_h2(doc, "1. Introduction")
    add_para(doc,
        "Ce septième chapitre est consacré au quatrième sprint, qui matérialise le cœur opérationnel "
        "de la plateforme : la circulation des bons de commande depuis la confirmatrice jusqu'à la "
        "livraison effective au client, en passant par la gestion en temps réel des incidents. Après les "
        "sprints précédents qui ont mis en place la base produit, l'authentification, le catalogue et "
        "les commandes en ligne, ce sprint introduit cinq nouveaux acteurs sur le terrain — la "
        "confirmatrice, le livreur COD, le livreur de transit, le client (en suivi) et le superviseur — "
        "et trois flux structurants : la transformation documentaire BC → BL synchronisée avec Sage X3, "
        "le suivi GPS du livreur diffusé en SignalR vers le client, et la conversation temps réel autour "
        "des réclamations. L'architecture cliente est multi-canal : interfaces React pour la confirmatrice "
        "et le superviseur, application Flutter pour le livreur et le client, le tout convergeant vers "
        "l'API ASP.NET Core et la base SQL Server, conformément au pattern adopté dans les chapitres "
        "précédents."
    )

    # 2. Objectif du sprint
    add_h2(doc, "2. Objectif du sprint")
    add_para(doc,
        "Le Sprint 4 s'est déroulé du 06/04/2026 au 25/04/2026 sur trois itérations hebdomadaires. "
        "Son objectif fonctionnel est de fermer la chaîne logistique de bout en bout, en garantissant "
        "trois propriétés non-fonctionnelles : la cohérence documentaire avec Sage X3 (transformation "
        "BC → BL idempotente, anti-doublon, synchronisation non bloquante), la temporalité du suivi "
        "client (latence GPS sous la seconde via SignalR, indicateur de fraîcheur), et la résilience "
        "des interactions temps réel (file locale offline côté Flutter, période de grâce de 5 secondes "
        "à la déconnexion pour ne pas libérer prématurément un cas pris en charge)."
    )
    add_para(doc, "Le périmètre couvre, en synthèse, les capacités suivantes :")
    add_bullet(doc, "Confirmation : consultation, modification de statut, transformation BC → BL, affectation automatique du livreur, export Sage X3.")
    add_bullet(doc, "Livraison COD : pool filtré par zone, prise en charge, diffusion GPS, optimisation de tournée, encaissement automatique.")
    add_bullet(doc, "Suivi client : carte adaptative TrackingStateCard exposant quatre états successifs.")
    add_bullet(doc, "Réclamations : création, prise en charge avec verrou, chat SignalR, escalade.")
    add_bullet(doc, "Transit inter-dépôts : missions, double scan code-barres, réception partielle, audit.")
    add_bullet(doc, "Supervision : carte temps réel, heatmap 90 jours, alertes, réaffectation de mission de transit.")

    # 3. Backlog du sprint
    add_h2(doc, "3. Backlog du sprint")
    add_para(doc,
        "Le tableau ci-après présente les dix-sept histoires utilisateurs retenues, leur découpe en "
        "tâches techniques, leur responsable et leur estimation en jours-homme."
    )
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    for i, h in enumerate(["ID", "Histoire utilisateur", "Tâches techniques", "Responsable", "Estim."]):
        hdr[i].text = h
    for row in S4_BACKLOG:
        cells = table.add_row().cells
        for i, v in enumerate(row): cells[i].text = v
    style_table(table)
    add_caption(doc, "Tableau 7.1 — Backlog du Sprint 4")

    # 4. Analyse et conception
    add_h2(doc, "4. Analyse et conception")
    add_para(doc,
        "L'analyse de ce sprint articule les cinq acteurs autour de quatre flux temps réel. "
        "Les diagrammes ci-dessous ont été réalisés en PlantUML à partir du code effectivement "
        "présent dans la solution (contrôleurs ASP.NET Core, hub SignalR et écrans Flutter) afin "
        "de garantir la cohérence entre la documentation et l'implémentation."
    )

    add_h3(doc, "4.1 Diagramme de cas d'utilisation du Sprint 4")
    add_para(doc,
        "Le diagramme de cas d'utilisation ci-dessous regroupe les fonctionnalités par paquet "
        "métier (Confirmation, Livraison COD, Transit, Réclamations, Supervision) et fait apparaître "
        "les relations <<include>> et <<extend>> entre cas (par exemple, « Démarrer la diffusion GPS » "
        "présuppose « Prendre en charge la livraison »)."
    )
    add_figure(doc, DIAG_DIR / "s4_usecase.png",
               "Figure 7.1 — Diagramme de cas d'utilisation du Sprint 4")

    add_h3(doc, "4.2 Diagrammes de séquence du Sprint 4")
    add_para(doc,
        "Quatre flux ont été modélisés : la transformation BC → BL avec affectation automatique et "
        "synchronisation Sage X3, la diffusion GPS et l'encaissement COD, la conversation de "
        "réclamation avec sa période de grâce de 5 secondes, et enfin le double scan du transit "
        "inter-dépôts."
    )

    add_h4(doc, "4.2.1 Transformation BC → BL et synchronisation Sage X3")
    add_para(doc,
        "La méthode TransformBcToBl du ConfirmateurController vérifie en transaction l'existence du "
        "BC (DO_Type = 0), s'assure qu'aucun BL n'a déjà été produit pour cette pièce, puis insère "
        "F_DOCENTETE (DO_Type = 1) et la copie des lignes F_DOCLIGNE. La référence du BL suit le "
        "format « BL » + yyMMddHHmm + n. Un livreur est sélectionné par l'AssignmentService selon "
        "deux critères pondérés : couverture de la zone du client (test Haversine) et longueur de "
        "file la plus courte. L'envoi vers Sage X3 est non bloquant : la création locale aboutit "
        "même si Sage est indisponible, le statut sageSuccess étant remonté dans le tableau de bord "
        "de synchronisation."
    )
    add_figure(doc, DIAG_DIR / "s4_seq_bctbl.png",
               "Figure 7.2 — Séquence : transformation BC → BL et synchronisation Sage X3")

    add_h4(doc, "4.2.2 Livraison COD, diffusion GPS et encaissement")
    add_para(doc,
        "L'appel POST /api/livreur/orders/{piece}/start-heading bascule la livraison en mode actif "
        "(IsActiveDelivery = true) et notifie le client via SignalR. Le Geolocator Flutter capture "
        "ensuite la position toutes les dix secondes ; chaque ping effectue un UPSERT dans "
        "F_LIVREUR_POSITION, un INSERT dans F_LIVREUR_POSITION_HISTORY et déclenche un événement "
        "LivreurPositionUpdate sur le groupe client-{userId}. Une distance Haversine inférieure à "
        "500 m déclenche une alerte de proximité. En cas de coupure réseau, l'application met les "
        "positions en file locale et les expédie par lots via /ping-batch en utilisant un "
        "ClientActionId pour garantir l'idempotence. Lorsque le livreur marque la livraison « Livré » "
        "(LI_Statut = 2), l'API applique automatiquement Encaisse = true, MontantEncaisse = "
        "DO_NetAPayer et EncaisseAt = UtcNow, alimentant la caisse COD du livreur."
    )
    add_figure(doc, DIAG_DIR / "s4_seq_gps.png",
               "Figure 7.3 — Séquence : livraison active, GPS temps réel et encaissement COD")

    add_h4(doc, "4.2.3 Réclamation client avec période de grâce SignalR")
    add_para(doc,
        "Le ReclamationHub maintient un compteur de connexions par confirmatrice. Lorsqu'une "
        "réclamation est créée, l'événement NouveauCas est diffusé au groupe « confirmateurs ». La "
        "prise en charge insère une session F_CONFIRMATRICE_SESSION et verrouille le cas. La "
        "subtilité du sprint réside dans la méthode OnDisconnectedAsync : si c'est la dernière "
        "connexion de la confirmatrice, le cas n'est pas libéré immédiatement ; un timer "
        "GracePeriod = TimeSpan.FromSeconds(5) est armé. Si une reconnexion survient pendant cet "
        "intervalle — typiquement un basculement Wi-Fi → 4G — le timer est annulé et le verrou "
        "préservé. Au-delà, le cas redevient ENVOYEE et un événement CommandeLiberee est diffusé."
    )
    add_figure(doc, DIAG_DIR / "s4_seq_reclamation.png",
               "Figure 7.4 — Séquence : réclamation client, chat SignalR et grâce 5 s")

    add_h4(doc, "4.2.4 Transit inter-dépôts par double scan")
    add_para(doc,
        "Le module transit s'appuie sur le plugin Flutter mobile_scanner. Le premier scan "
        "(POST /scan-pickup) fait passer la ligne F_TRANSFERT de EN_ATTENTE_TRANSIT à EN_TRANSIT, "
        "enregistre PickedUpAt et les coordonnées GPS de récupération. Le second scan "
        "(POST /scan-delivery) bascule la ligne en RECU_DEPOT_DESTINE avec DeliveredAt. Une "
        "fenêtre de RevertWindowMinutes = 10 minutes autorise l'annulation d'un pickup erroné, "
        "tracée dans F_TRANSFERT_AUDIT_LOG. Lorsqu'une partie seulement des articles est reçue, "
        "le statut TRANSIT_PARTIELLEMENT_RECU déclenche une alerte vers le groupe « superviseurs »."
    )
    add_figure(doc, DIAG_DIR / "s4_seq_transit.png",
               "Figure 7.5 — Séquence : transit inter-dépôts par double scan code-barres")

    # 5. Réalisation
    add_h2(doc, "5. Réalisation")
    add_para(doc,
        "Les interfaces utilisateurs ci-dessous concrétisent les cas d'utilisation. Les écrans "
        "React (confirmatrice et superviseur) consomment les endpoints via Axios encapsulé dans des "
        "hooks React Query, tandis que les écrans Flutter (livreur et client) passent par ApiClient "
        "avec injection du JWT. Les emplacements réservés correspondent aux captures à intégrer "
        "lors de la version finale."
    )

    sections_real = [
        ("5.1 Interface confirmatrice — liste des bons de commande",
         "Construite sur ConfirmateurOrdersPage.tsx, cette page liste les BC paginés avec filtres "
         "(statut, période, montant), recherche par référence ou par client, et accès direct au "
         "détail. Les colonnes affichent DO_Piece, le client résolu, DO_Date, DO_NetAPayer et le "
         "libellé de DO_Valide. Les actions disponibles sont contextuelles : modification du statut "
         "depuis le menu inline, ou ouverture du détail pour la transformation en BL.",
         "capture liste des bons de commande confirmatrice",
         "Figure 7.6 — Interface confirmatrice : liste des bons de commande"),
        ("5.2 Interface confirmatrice — transformation BC en BL",
         "L'écran ConfirmateurOrderDetailsPage.tsx propose, lorsque le BC est confirmé, le bouton "
         "« Transformer en BL ». La réponse de l'API expose blPiece, alreadyExists, sageSent et "
         "sageSuccess : un message contextuel s'affiche en cas de doublon, d'échec partiel ou de "
         "succès complet. L'identifiant du livreur affecté automatiquement est rappelé.",
         "capture de la transformation BC → BL",
         "Figure 7.7 — Interface confirmatrice : transformation BC → BL"),
        ("5.3 Interface confirmatrice — gestion des réclamations",
         "L'interface des réclamations affiche les cas en cours, classés par priorité et par "
         "ancienneté. L'événement SignalR NouveauCas fait apparaître les nouveaux cas sans "
         "rafraîchissement. La confirmatrice peut prendre en charge un cas (verrou avec grâce 5 s), "
         "modifier son statut, créer un bon d'échange, demander une correction d'adresse ou "
         "réassigner. Le panneau de chat reçoit chaque NouveauMessage en temps réel.",
         "capture de la gestion des réclamations + chat SignalR",
         "Figure 7.8 — Interface confirmatrice : gestion des réclamations et chat SignalR"),
        ("5.4 Interface livreur — livraisons disponibles",
         "L'écran new_orders_screen.dart affiche les BL du pool filtrés par gouvernorat et "
         "délégation. Chaque carte expose l'adresse, le montant COD attendu et les coordonnées du "
         "client. Le bouton « Prendre en charge » déclenche POST /assign et bascule la livraison "
         "vers l'onglet « Mes livraisons » ; un retour 409 affiche un toast en cas de conflit avec "
         "un autre livreur.",
         "capture du pool livraisons côté livreur",
         "Figure 7.9 — Interface livreur : livraisons disponibles"),
        ("5.5 Interface livreur — suivi GPS et carte de tournée",
         "La carte intégrée superpose la position du livreur (mise à jour SignalR), les "
         "destinations restantes et l'itinéraire ordonné par l'algorithme du plus-proche-voisin. "
         "Pour chaque arrêt, le panneau affiche DistanceFromPreviousKm, CumulativeDistanceKm et "
         "CumulativeEtaMinutes (vitesse de référence 35 km/h). Côté client, le composant "
         "LiveDeliveryMapSheet expose la position GPS du livreur en HEADING_TO_YOU.",
         "capture du suivi GPS et carte de tournée",
         "Figure 7.10 — Interface livreur : suivi GPS et optimisation de tournée"),
        ("5.6 Interface livreur — caisse COD",
         "L'écran livreur_stats_screen.dart agrège le montant encaissé du jour, les livraisons "
         "réalisées, le taux de réussite et la répartition par statut. L'historique des remises au "
         "dépôt est consultable, ainsi que les écarts éventuels signalés par le superviseur lors du "
         "rapprochement de caisse.",
         "capture de la caisse COD du livreur",
         "Figure 7.11 — Interface livreur : caisse COD et statistiques"),
        ("5.7 Interface client — TrackingStateCard adaptative",
         "Le composant TrackingStateCard prend la forme déterminée par l'endpoint "
         "/api/client/orders/{piece}/tracking-state : AT_DEPOT et IN_DELIVERY_QUEUE en mode "
         "informatif, HEADING_TO_YOU en mode carte (GPS du livreur, ETA Haversine à 40 km/h, "
         "indicateur de fraîcheur vert/orange/rouge selon l'âge du dernier ping), et TERMINAL "
         "(livré / retour / refusé). Des raccourcis d'appel et de SMS au livreur sont disponibles "
         "en HEADING_TO_YOU.",
         "capture de la TrackingStateCard sur les quatre états",
         "Figure 7.12 — Interface client : TrackingStateCard adaptative"),
        ("5.8 Interface de transit inter-dépôts et scan code-barres",
         "transit_home_screen.dart segmente les missions en trois onglets : en attente, en cours "
         "et historique. transit_mission_details_screen.dart liste les articles attendus avec leur "
         "quantité. transit_barcode_scanner_screen.dart active la caméra via mobile_scanner ; "
         "chaque scan reçoit un retour sonore et tactile, et la fenêtre d'annulation de dix minutes "
         "est exposée sur la fiche de la ligne.",
         "capture du scanner de transit inter-dépôts",
         "Figure 7.13 — Interface transit : missions et scanner code-barres"),
        ("5.9 Interface superviseur — carte, alertes et heatmap",
         "Le superviseur dispose d'une carte temps réel des livreurs actifs (avec dernier ping), "
         "d'un tableau de bord des alertes (retard, échec, zone non couverte, écart de caisse) et "
         "d'une heatmap construite par LivreurMapController.HeatMap : agrégation par cellule "
         "d'environ 500 m, intensité égale au ratio (retours + reports) / livraisons sur une "
         "fenêtre glissante de 90 jours.",
         "capture de la console superviseur",
         "Figure 7.14 — Interface superviseur : livreurs actifs et heatmap 90 j"),
    ]
    for title, body, ph_label, caption in sections_real:
        add_h3(doc, title)
        add_para(doc, body)
        add_placeholder(doc, ph_label)
        add_caption(doc, caption)

    # 6. Tests et validation
    add_h2(doc, "6. Tests et validation")
    add_para(doc,
        "Les tests fonctionnels du sprint ont été exécutés à mesure que les fonctionnalités "
        "étaient développées. Chaque cas a été vérifié à la fois sur des données réalistes (jeu de "
        "départ Sage X3) et sur des cas limites identifiés pendant les itérations : indisponibilité "
        "de Sage, conflit de prise en charge, perte de connexion du livreur, déconnexion fugace de "
        "la confirmatrice, scans erronés."
    )
    t = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["ID", "Objectif", "Résultat attendu", "Statut"]):
        t.rows[0].cells[i].text = h
    for row in S4_TESTS:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = v
    style_table(t)
    add_caption(doc, "Tableau 7.2 — Tests fonctionnels du Sprint 4")

    # 7. Conclusion
    add_h2(doc, "7. Conclusion")
    add_para(doc,
        "Le Sprint 4 referme la chaîne logistique : un bon de commande créé en ligne peut désormais "
        "être confirmé, transformé en bon de livraison synchronisé avec Sage X3, affecté "
        "automatiquement à un livreur de zone, suivi en temps réel par le client via SignalR, livré "
        "avec encaissement COD automatique, et traité a posteriori par une réclamation conversée et "
        "verrouillée. Le module de transit inter-dépôts apporte la traçabilité fine des mouvements "
        "de stock entre points de stockage, tandis que la console superviseur synthétise l'activité "
        "en cours. La résilience opérationnelle a été particulièrement soignée — file locale "
        "Flutter avec idempotence, période de grâce de 5 secondes côté SignalR, synchronisation "
        "Sage X3 non bloquante — afin d'absorber les aléas du terrain. Le sprint suivant capitalise "
        "sur ce socle pour exposer une couche de pilotage transverse à l'administrateur."
    )

def build_chap8(doc):
    doc.add_page_break()
    add_h1(doc, "Chapitre 8 : Sprint 5 — Tableaux de bord, exports et chatbot")

    # 1. Introduction
    add_h2(doc, "1. Introduction")
    add_para(doc,
        "Ce huitième et dernier chapitre fonctionnel est consacré au cinquième sprint, dédié à la "
        "couche de pilotage et d'aide à la décision. À ce stade du projet, l'ensemble des "
        "fonctionnalités métier — catalogue, commandes, paiement, confirmation, livraison, "
        "réclamations, transit — est opérationnel ; le défi du Sprint 5 consiste à transformer cette "
        "masse de données opérationnelle en une plateforme pilotable, exploitable et interrogeable "
        "par l'administrateur. Trois axes structurent la livraison : onze tableaux de bord "
        "multi-dimensionnels servis par un service d'agrégation unique, des exports Excel et PDF "
        "génériques s'appuyant sur ClosedXML et QuestPDF, et un chatbot conversationnel piloté par "
        "un workflow n8n et un LLM Groq llama-3.3-70b chargé du routage et de la mise en forme des "
        "réponses."
    )

    # 2. Objectif du sprint
    add_h2(doc, "2. Objectif du sprint")
    add_para(doc,
        "Le Sprint 5 s'est déroulé du 27/04/2026 au 23/05/2026. Son objectif est triple : doter "
        "l'administrateur d'une lecture transversale et cohérente de la performance "
        "opérationnelle, lui permettre d'exporter de manière fiable les données brutes à des fins "
        "d'audit ou de partage hors plateforme, et lui ouvrir un canal conversationnel "
        "lui évitant la navigation entre dashboards lorsqu'une question ponctuelle se pose. Sur le "
        "plan non fonctionnel, l'enjeu est l'homogénéité (toutes les pages exposent le même contrat "
        "ProDashboardPageResponseDto : ExecutiveSummary, KPI, distributions, alertes, insights), la "
        "robustesse des exports (plafond MaxRows = 10 000 lignes, en-tête / pagination en PDF) et "
        "la qualité des réponses du chatbot (routage déterministe via Groq, formatage strict)."
    )

    # 3. Backlog du sprint
    add_h2(doc, "3. Backlog du sprint")
    add_para(doc,
        "Le tableau ci-après présente les dix-huit histoires utilisateurs traitées au cours du "
        "sprint."
    )
    t = doc.add_table(rows=1, cols=5)
    for i, h in enumerate(["ID", "Histoire utilisateur", "Tâches techniques", "Responsable", "Estim."]):
        t.rows[0].cells[i].text = h
    for row in S5_BACKLOG:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = v
    style_table(t)
    add_caption(doc, "Tableau 8.1 — Backlog du Sprint 5")

    # 4. Analyse et conception
    add_h2(doc, "4. Analyse et conception")
    add_para(doc,
        "L'analyse de ce sprint porte sur des fonctionnalités transverses : les dashboards "
        "réutilisent des projections sur F_DOCENTETE, F_LIVRAISON, F_RECLAMATION, F_TRANSFERT et "
        "les vues issues de Sage X3 ; les exports partagent un service unique (ExportService) ; le "
        "chatbot orchestre quatre actions (KB / QUERY / ANALYZE / PREDICT) via un même contrat de "
        "réponse."
    )

    add_h3(doc, "4.1 Diagramme de cas d'utilisation du Sprint 5")
    add_para(doc,
        "Le diagramme de cas d'utilisation ci-dessous regroupe les fonctionnalités par paquet "
        "(tableaux de bord, exports, chatbot) et illustre les trois acteurs techniques qui "
        "interviennent en arrière-plan du chatbot : le workflow n8n (orchestration), le LLM Groq "
        "(routage et formatage), et le système Sage X3 (source pour le dashboard de "
        "synchronisation)."
    )
    add_figure(doc, DIAG_DIR / "s5_usecase.png",
               "Figure 8.1 — Diagramme de cas d'utilisation du Sprint 5")

    add_h3(doc, "4.2 Diagramme de séquence — Chatbot administrateur")
    add_para(doc,
        "L'administrateur saisit sa question dans ChatbotSandboxPage ; le payload est envoyé au "
        "webhook n8n /admin-chat-v3. Le premier nœud n8n détecte la langue (français, arabe ou "
        "« tounsi »), puis transmet la requête à AdminChatController.AskAsync. "
        "L'orchestrateur appelle Groq llama-3.3-70b une première fois pour le routage (choix de "
        "l'action : KB documentaire, QUERY métriques, ANALYZE statistiques ou PREDICT ML.NET), "
        "exécute l'action correspondante en interrogeant la base, puis appelle Groq une seconde "
        "fois pour formater la réponse (texte + structure de tableau et de graphique). Le message "
        "est journalisé dans F_CHATBOT_MESSAGE. Côté n8n, un nœud complémentaire récupère les "
        "insights proactifs en attente (F_CHATBOT_INSIGHTS) et les fusionne à la réponse avant "
        "le retour vers l'interface React. Une variante /ask-stream renvoie un flux SSE en quatre "
        "phases (routing → data → chunks → done) pour un rendu mot-à-mot."
    )
    add_figure(doc, DIAG_DIR / "s5_seq_chatbot.png",
               "Figure 8.2 — Séquence : chatbot administrateur (React → n8n → Groq)")

    add_h3(doc, "4.3 Diagramme de séquence — Export Excel / PDF")
    add_para(doc,
        "L'administrateur déclenche l'export depuis l'écran de synthèse : la route "
        "GET /api/admin/orders/export?format=xlsx|pdf parse la période, projette F_DOCENTETES dans "
        "la limite de MaxRows lignes, et délègue la sérialisation à ExportService. ClosedXML "
        "produit un classeur Excel avec en-têtes en gras et largeurs automatiques ; QuestPDF "
        "produit un document A4 paginé avec un en-tête de période et un tableau ajusté. Le flux "
        "binaire est renvoyé en téléchargement direct via Content-Disposition. Le même service "
        "est réutilisé pour /api/admin/reclamations/export."
    )
    add_figure(doc, DIAG_DIR / "s5_seq_export.png",
               "Figure 8.3 — Séquence : génération d'un export Excel ou PDF")

    # 5. Réalisation
    add_h2(doc, "5. Réalisation")
    add_para(doc,
        "Les pages React s'appuient toutes sur le composant générique DashboardAnalyticsPage, "
        "paramétré par une clé (pageKey) qui pointe vers l'endpoint d'agrégation et la "
        "configuration d'affichage. Cette factorisation garantit une expérience cohérente sur "
        "l'ensemble des dashboards : structure de page, panneau d'alertes, mise en page des KPI "
        "et bouton d'actualisation manuelle."
    )

    sections_real = [
        ("5.1 Tableau de bord d'ensemble",
         "Construit sur AdminOverviewDashboardPage.tsx, ce tableau de bord agrège les indicateurs "
         "clés : chiffre d'affaires de la période, nombre de livraisons réalisées, réclamations "
         "ouvertes, livreurs actifs et tendance hebdomadaire. Le panneau d'alertes "
         "(AlertPanel.tsx) met en évidence les anomalies par sévérité (warning, critical) — par "
         "exemple un nombre élevé de BC en attente, des stocks critiques ou des cas réclamations "
         "non assignés.",
         "capture du tableau de bord d'ensemble",
         "Figure 8.4 — Interface : tableau de bord d'ensemble"),
        ("5.2 Tableau de bord ventes et commandes",
         "AdminSalesDashboardPage.tsx et AdminOrdersDashboardPage.tsx exposent respectivement la "
         "performance commerciale (CA par période, top produits, répartition B2C / B2B) et la "
         "santé du flux commandes (répartition par statut DO_Valide, délai moyen BC → BL → "
         "livraison, segmentation par zone). Les graphiques sont implémentés avec Recharts.",
         "capture des dashboards ventes et commandes",
         "Figure 8.5 — Interface : dashboards ventes et commandes"),
        ("5.3 Tableau de bord logistique et livreurs",
         "AdminLogisticsDashboardPage et AdminDriversDashboardPage présentent le taux de "
         "livraison par zone, le délai moyen, la heatmap des échecs 90 jours, le classement des "
         "livreurs et leurs caisses COD respectives. La répartition fine des statuts "
         "(LI_Statut 2 livré, 3 retour, 4 dépôt, 5 reporté) permet d'isoler immédiatement les "
         "zones ou livreurs en difficulté.",
         "capture des dashboards logistique et livreurs",
         "Figure 8.6 — Interface : dashboards logistique et livreurs"),
        ("5.4 Tableau de bord réclamations et confirmatrices",
         "AdminReclamationsDashboardPage croise les réclamations par motif "
         "(COLIS_ENDOMMAGE, NON_LIVRE, MAUVAIS_ARTICLE, NUMERO_INCORRECT, ADRESSE_INCORRECTE) et "
         "par statut (ENVOYEE, EN_COURS, CLOTUREE, REFUSEE). AdminConfirmateurDashboardPage suit "
         "la charge de chaque confirmatrice : cas actifs, taux de résolution et temps moyen de "
         "traitement.",
         "capture des dashboards réclamations et confirmatrices",
         "Figure 8.7 — Interface : dashboards réclamations et confirmatrices"),
        ("5.5 Tableau de bord synchronisation Sage X3",
         "AdminSyncDashboardPage (alias /api/dashboard/admin-sync) présente l'état des données "
         "synchronisées avec Sage X3 : nombre d'articles, dépôts, stocks et clients, ainsi que "
         "des alertes d'intégrité telles que les clients dépourvus d'adresse ou les articles "
         "sans famille. Le panneau « Last sync status » affiche l'horodatage de la dernière "
         "exécution.",
         "capture du dashboard synchronisation Sage X3",
         "Figure 8.8 — Interface : dashboard synchronisation Sage X3"),
        ("5.6 Exports Excel et PDF",
         "Les endpoints /api/admin/orders/export et /api/admin/reclamations/export acceptent un "
         "paramètre format (xlsx ou pdf) et une période. Le service ExportService.cs centralise la "
         "production : ClosedXML pour Excel (en-têtes en gras, largeurs automatiques) et QuestPDF "
         "pour PDF (mise en page A4 avec en-tête + pagination). Le plafond MaxRows = 10 000 lignes "
         "protège la mémoire serveur et prévient l'utilisateur dans l'interface.",
         "capture de la page Exports",
         "Figure 8.9 — Interface : exports Excel et PDF"),
        ("5.7 Chatbot administrateur — Sandbox",
         "ChatbotSandboxPage.tsx accueille la conversation interactive : zone de saisie, "
         "historique du fil et rendu structuré des réponses (message, tableau, graphique). Les "
         "exemples typiques traités par le pipeline Router → Exécuteur → Formatter incluent "
         "« CA du jour », « Top 5 produits du mois », « Livreurs à plus de 10 % de retours », "
         "« Réclamations ouvertes depuis plus de 48 h ». Une bascule SSE active le rendu "
         "mot-à-mot via /ask-stream.",
         "capture du chatbot administrateur (sandbox)",
         "Figure 8.10 — Interface : chatbot administrateur (Sandbox)"),
        ("5.8 Chatbot — Conversations et insights",
         "ChatbotConversationsPage parcourt les sessions historiques (F_CHATBOT_MESSAGE) avec "
         "filtres par date et par utilisateur. ChatbotInsightsPage agrège les insights proactifs "
         "(F_CHATBOT_INSIGHTS) : chaque carte expose un titre, un constat chiffré et une action "
         "suggérée. Le bouton « feedback » alimente l'évaluation qualitative des insights, et "
         "/kb/refresh régénère la base de connaissances exploitée par le routeur Groq.",
         "capture des conversations et insights chatbot",
         "Figure 8.11 — Interface : conversations et insights chatbot"),
    ]
    for title, body, ph_label, caption in sections_real:
        add_h3(doc, title)
        add_para(doc, body)
        add_placeholder(doc, ph_label)
        add_caption(doc, caption)

    # 6. Tests et validation
    add_h2(doc, "6. Tests et validation")
    add_para(doc,
        "Les tests du sprint vérifient la cohérence des données présentées par rapport aux "
        "données sources (recalcul SQL direct), la conformité des artefacts d'export aux "
        "spécifications (format, plafond MaxRows, pagination PDF) et la qualité des réponses du "
        "chatbot (langue, fidélité, persistance de la session)."
    )
    t = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["ID", "Objectif", "Résultat attendu", "Statut"]):
        t.rows[0].cells[i].text = h
    for row in S5_TESTS:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = v
    style_table(t)
    add_caption(doc, "Tableau 8.2 — Tests fonctionnels du Sprint 5")

    # 7. Conclusion
    add_h2(doc, "7. Conclusion")
    add_para(doc,
        "Le Sprint 5 referme le projet en plaçant l'administrateur au centre d'une plateforme "
        "complètement instrumentée. Les onze tableaux de bord offrent une vision unifiée et "
        "homogène — un contrat unique ProDashboardPageResponseDto, un composant React unique — "
        "tandis que les exports Excel et PDF, construits sur ClosedXML et QuestPDF, permettent de "
        "matérialiser cette vision dans des supports partageables ou auditables. Le chatbot, "
        "orchestré par n8n et adossé à Groq llama-3.3-70b, agit comme une couche d'accès "
        "conversationnelle qui rapproche les données métier des questions opérationnelles "
        "quotidiennes, complétée par des insights proactifs et un mode streaming SSE. Avec ce "
        "cinquième sprint, la solution atteint un niveau de maturité opérationnelle compatible "
        "avec une démonstration et une mise en production : le chapitre suivant pourra dresser le "
        "bilan global, les enseignements et les perspectives du projet."
    )

# ─────────────────────────────────────────────────────────────────────
# main
# ─────────────────────────────────────────────────────────────────────
def main():
    print("1/3  Écriture des fichiers PUML…")
    write_puml()
    print("2/3  Rendu PNG…")
    render_pngs()
    print("3/3  Construction du docx…")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    build_chap7(doc)
    build_chap8(doc)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"\n✅ docx généré : {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
