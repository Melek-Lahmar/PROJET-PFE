#!/usr/bin/env python3
"""
build_final_v2.py  —  Reconstruction complète du rapport PFE
  1. Génère 18 diagrammes PlantUML (use case + séquence) style Sprint 1
  2. Part du document propre (sans diagrammes)
  3. Insère les PNG aux emplacements [Insérer ici...]
  4. Transforme les tableaux backlog (5 colonnes, colonne Modifications)
  5. Corrige les noms Melek / Tawfik
  6. Sauvegarde Rapport_PFE_Final_Modifie.docx
"""

import os, subprocess
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

# ─── Chemins ───────────────────────────────────────────────────────────────────
DIAG_DIR   = Path("/tmp/diagrams")
PLANTUML   = "/tmp/plantuml.jar"
CLEAN_DOC  = "/tmp/Rapport_PFE_clean.docx"
OUT_DOC    = "/home/user/PROJET-PFE/Rapport/Rapport_PFE_Final_Modifie.docx"
DIAG_DIR.mkdir(exist_ok=True)

# ─── Skinparams communs (identiques Sprint 1) ──────────────────────────────────
SK_SEQ = """\
skinparam backgroundColor white
skinparam handwritten false
skinparam shadowing false
skinparam roundCorner 0
skinparam DefaultFontName Arial
skinparam DefaultFontSize 11
skinparam TitleFontStyle bold
skinparam TitleFontSize 13
skinparam SequenceMessageAlign center
skinparam responseMessageBelowArrow true
skinparam ParticipantPadding 20
skinparam BoxPadding 10
skinparam SequenceLifeLineBorderColor #555555
skinparam SequenceArrowColor #000000
skinparam ParticipantBackgroundColor #FFFFFF
skinparam ParticipantBorderColor #000000
skinparam ActorBackgroundColor #FFFFFF
skinparam ActorBorderColor #000000
skinparam NoteBackgroundColor #FFFFFF
skinparam NoteBorderColor #000000
"""

SK_UC = """\
skinparam backgroundColor white
skinparam handwritten false
skinparam shadowing false
skinparam roundCorner 0
skinparam DefaultFontName Arial
skinparam DefaultFontSize 11
skinparam TitleFontStyle bold
skinparam TitleFontSize 13
skinparam ActorBackgroundColor #FFFFFF
skinparam ActorBorderColor #000000
skinparam UsecaseBackgroundColor #FFFFFF
skinparam UsecaseBorderColor #000000
skinparam UsecaseFontSize 11
skinparam ArrowColor #000000
skinparam rectangleBorderColor #000000
skinparam rectangleBackgroundColor #FFFFFF
"""

# ─── Définitions PlantUML ──────────────────────────────────────────────────────
DIAGRAMS = {}

# ══════════════════════════════════════════════════════════════
# SPRINT 2
# ══════════════════════════════════════════════════════════════
DIAGRAMS["s2_usecase"] = f"""@startuml
title Diagramme de cas d'utilisation — Sprint 2
{SK_UC}
left to right direction

actor "Visiteur" as V
actor "Client B2C" as CB2C
actor "Client B2B" as CB2B
actor "Système Konnect" as KON

rectangle "Système e-commerce" {{
  usecase "Consulter le catalogue" as UC1
  usecase "Rechercher un produit" as UC2
  usecase "Filtrer les produits" as UC3
  usecase "Consulter le détail produit" as UC4
  usecase "Gérer ses favoris" as UC5
  usecase "Gérer le panier" as UC6
  usecase "Passer une commande" as UC7
  usecase "Payer en ligne (Konnect)" as UC8
  usecase "Suivre sa commande" as UC9
  usecase "Gérer ses adresses" as UC10
  usecase "Demander un devis B2B" as UC11
}}

V      --> UC1
V      --> UC2
CB2C   --> UC1
CB2C   --> UC2
CB2C   --> UC3
CB2C   --> UC4
CB2C   --> UC5
CB2C   --> UC6
CB2C   --> UC7
CB2C   --> UC9
CB2C   --> UC10
CB2B   --> UC11
KON    --> UC8

UC2  ..> UC1 : <<include>>
UC3  ..> UC1 : <<include>>
UC7  ..> UC6 : <<include>>
UC8  ..> UC7 : <<include>>
@enduml"""

DIAGRAMS["s2_seq_catalogue"] = f"""@startuml
title Diagramme de séquence — Consultation du catalogue
{SK_SEQ}
actor "Utilisateur" as U
participant "Interface React" as R
participant "ArticleController" as AC
database "SQL Server" as DB

U  -> R  : Ouvre la page catalogue
R  -> AC : GET /api/articles?page=1
AC -> DB : SELECT articles + stock
DB --> AC : Liste articles
AC --> R  : JSON {{items, total, pages}}
R  --> U  : Affiche la grille produits

U  -> R  : Applique un filtre
R  -> AC : GET /api/articles?categorie=X&minPrix=Y
AC -> DB : SELECT avec filtres
DB --> AC : Articles filtrés
AC --> R  : JSON résultats
R  --> U  : Met à jour l'affichage

U  -> R  : Clique sur un produit
R  -> AC : GET /api/articles/{{id}}
AC -> DB : SELECT détail + variantes + images
DB --> AC : Détail produit
AC --> R  : JSON détail
R  --> U  : Affiche la fiche produit
@enduml"""

DIAGRAMS["s2_seq_panier"] = f"""@startuml
title Diagramme de séquence — Gestion du panier
{SK_SEQ}
actor "Client" as C
participant "Interface React" as R
participant "PanierController" as PC
database "SQL Server" as DB

C  -> R  : Clique "Ajouter au panier"
R  -> PC : POST /api/panier/ajouter
PC -> DB : Vérifier stock disponible
DB --> PC : Stock OK
PC -> DB : INSERT/UPDATE PanierItem
DB --> PC : OK
PC --> R  : JSON {{panier, total, nbItems}}
R  --> C  : Met à jour l'icône panier

C  -> R  : Ouvre le panier
R  -> PC : GET /api/panier
PC -> DB : SELECT panier + articles
DB --> PC : Contenu du panier
PC --> R  : JSON panier complet
R  --> C  : Affiche le récapitulatif

C  -> R  : Modifie une quantité
R  -> PC : PUT /api/panier/{{itemId}}
PC -> DB : UPDATE PanierItem
DB --> PC : OK
PC --> R  : JSON panier mis à jour
R  --> C  : Actualise le total
@enduml"""

DIAGRAMS["s2_seq_commande"] = f"""@startuml
title Diagramme de séquence — Création d'une commande client
{SK_SEQ}
actor "Client" as C
participant "Interface React" as R
participant "CommandeController" as CC
database "SQL Server" as DB

C  -> R  : Valide le panier
R  -> CC : GET /api/panier
CC -> DB : Charger panier + articles
DB --> CC : Données panier
CC --> R  : Panier + total
R  --> C  : Affiche le récapitulatif

C  -> R  : Choisit l'adresse de livraison
C  -> R  : Confirme la commande
R  -> CC : POST /api/commandes
CC -> DB : Vérifier stock final
DB --> CC : Stock confirmé
CC -> DB : INSERT F_DOCENTETE (DO_Type=2)
CC -> DB : INSERT F_DOCLIGNE par article
CC -> DB : UPDATE stock (- quantité)
CC -> DB : EMPTY panier
DB --> CC : Commande créée
CC --> R  : {{commandeId, reference, statut: EN_ATTENTE}}
R  --> C  : Page de confirmation
@enduml"""

DIAGRAMS["s2_seq_konnect"] = f"""@startuml
title Diagramme de séquence — Paiement en ligne via Konnect
{SK_SEQ}
actor "Client" as C
participant "Interface React" as R
participant "PaiementController" as PC
participant "Konnect API" as KON
database "SQL Server" as DB

C   -> R   : Clique "Payer en ligne"
R   -> PC  : POST /api/paiement/initier
PC  -> DB  : Vérifier commande en attente
DB  --> PC : Commande valide
PC  -> KON : POST /payments/init
KON --> PC : {{payUrl, paymentRef}}
PC  -> DB  : UPDATE commande {{paymentRef}}
PC  --> R  : {{payUrl}}
R   --> C  : Redirect vers Konnect

C   -> KON : Saisit infos carte
KON --> C  : Confirmation paiement
C   -> KON : Confirme
KON --> C  : Redirect success URL

KON -> PC  : POST /api/paiement/webhook
PC  -> DB  : UPDATE commande {{statut: PAYEE}}
DB  --> PC : OK
PC  --> KON : 200 OK

C   -> R   : Page succès
R   -> PC  : GET /api/commandes/{{id}}
PC  -> DB  : SELECT commande
DB  --> PC : Commande payée
PC  --> R  : Détails commande
R   --> C  : Confirmation avec numéro
@enduml"""

DIAGRAMS["s2_seq_devis"] = f"""@startuml
title Diagramme de séquence — Demande de devis B2B
{SK_SEQ}
actor "Client B2B" as CB
participant "Interface React" as R
participant "DevisController" as DC
database "SQL Server" as DB
actor "Administrateur" as ADM

CB  -> R   : Accède à l'espace B2B
R   -> DC  : GET /api/b2b/produits
DC  -> DB  : SELECT articles B2B + tarifs
DB  --> DC : Catalogue B2B
DC  --> R  : JSON produits
R   --> CB : Affiche le catalogue

CB  -> R   : Soumet une demande de devis
R   -> DC  : POST /api/devis
DC  -> DB  : INSERT F_DEVIS_ENTETE {{statut: EN_ATTENTE}}
DC  -> DB  : INSERT F_DEVIS_LIGNE
DB  --> DC : Devis créé
DC  --> R  : {{devisId, reference}}
R   --> CB : Confirmation soumission

ADM -> R   : Consulte et modifie les prix
R   -> DC  : PUT /api/admin/devis/{{id}}
DC  -> DB  : UPDATE F_DEVIS_ENTETE {{statut: EN_NEGOCIATION}}
DB  --> DC : OK
DC  --> R  : Devis mis à jour

CB  -> R   : Accepte le devis
R   -> DC  : POST /api/devis/{{id}}/accepter
DC  -> DB  : UPDATE statut = ACCEPTE
DB  --> DC : OK
DC  --> R  : Devis accepté
R   --> CB : Confirmation
@enduml"""

# ══════════════════════════════════════════════════════════════
# SPRINT 3
# ══════════════════════════════════════════════════════════════
DIAGRAMS["s3_usecase"] = f"""@startuml
title Diagramme de cas d'utilisation — Sprint 3
{SK_UC}
left to right direction

actor "Administrateur" as ADM
actor "Vendeur" as VEN
actor "Client B2B" as CB2B
actor "Sage X3" as SX3

rectangle "Gestion commerciale et espace vendeur" {{
  usecase "Gérer les produits (CRUD)" as UC1
  usecase "Gérer les images Cloudinary" as UC2
  usecase "Synchroniser avec Sage X3" as UC3
  usecase "Gérer les stocks et dépôts" as UC4
  usecase "Gérer les clients B2B" as UC5
  usecase "Configurer les paramètres" as UC6
  usecase "Consulter le catalogue vendeur" as UC7
  usecase "Créer une commande vendeur" as UC8
  usecase "Gérer les devis B2B" as UC9
  usecase "Convertir devis en commande" as UC10
  usecase "Consulter les indicateurs" as UC11
}}

ADM  --> UC1
ADM  --> UC2
ADM  --> UC3
ADM  --> UC4
ADM  --> UC5
ADM  --> UC6
ADM  --> UC11
VEN  --> UC7
VEN  --> UC8
VEN  --> UC9
VEN  --> UC10
CB2B --> UC9
SX3  --> UC3

UC2  ..> UC1 : <<include>>
UC3  ..> UC4 : <<extend>>
UC10 ..> UC9 : <<include>>
@enduml"""

DIAGRAMS["s3_seq_sage"] = f"""@startuml
title Diagramme de séquence — Synchronisation Sage X3
{SK_SEQ}
actor "Administrateur" as ADM
participant "Interface React" as R
participant "SageX3Controller" as SC
participant "SageX3Service" as SS
database "SQL Server" as DB
database "Sage X3" as SX3

ADM -> R   : Clique "Synchroniser Sage X3"
R   -> SC  : POST /api/admin/sage/sync
SC  -> SS  : Lancer synchronisation()

SS  -> SX3 : SELECT F_ARTICLE actifs
SX3 --> SS : Liste articles

loop Pour chaque article
  SS -> DB : UPSERT F_ARTICLE
  DB --> SS : OK
end

SS  -> SX3 : SELECT F_TARIF
SX3 --> SS : Grilles de tarifs

loop Pour chaque tarif
  SS -> DB : UPSERT F_TARIF
  DB --> SS : OK
end

SS  -> SX3 : SELECT F_STOCK par dépôt
SX3 --> SS : Niveaux de stock

loop Pour chaque stock
  SS -> DB : UPSERT F_STOCK
  DB --> SS : OK
end

SS  --> SC : {{articles: N, tarifs: M, stocks: P}}
SC  -> DB  : INSERT SyncLog
SC  --> R  : {{success, stats}}
R   --> ADM : Rapport de synchronisation
@enduml"""

DIAGRAMS["s3_seq_devis"] = f"""@startuml
title Diagramme de séquence — Workflow devis B2B
{SK_SEQ}
actor "Client B2B" as CB
participant "Interface React" as R
participant "DevisController" as DC
database "SQL Server" as DB
actor "Vendeur" as VEN

CB  -> R   : Soumet demande de devis
R   -> DC  : POST /api/devis
DC  -> DB  : INSERT F_DEVIS_ENTETE {{EN_ATTENTE}}
DC  -> DB  : INSERT F_DEVIS_LIGNE
DC  --> R  : Devis créé

VEN -> R   : Modifie les prix
R   -> DC  : PUT /api/vendeur/devis/{{id}}
DC  -> DB  : UPDATE F_DEVIS_LIGNE (prix négocié)
DC  -> DB  : UPDATE statut = EN_NEGOCIATION
DB  --> DC : OK
DC  --> R  : Devis en négociation

CB  -> R   : Accepte le devis
R   -> DC  : POST /api/devis/{{id}}/accepter
DC  -> DB  : UPDATE statut = ACCEPTE
DB  --> DC : OK

VEN -> R   : Convertit en commande
R   -> DC  : POST /api/vendeur/devis/{{id}}/convertir
DC  -> DB  : INSERT F_DOCENTETE (DO_Type=2)
DC  -> DB  : INSERT F_DOCLIGNE (prix du devis)
DC  -> DB  : UPDATE statut = CONVERTI
DB  --> DC : Commande créée
DC  --> R  : {{commandeId, reference}}
R   --> VEN : Devis converti
@enduml"""

DIAGRAMS["s3_seq_vendeur"] = f"""@startuml
title Diagramme de séquence — Gestion commandes espace vendeur
{SK_SEQ}
actor "Vendeur" as VEN
participant "Interface React" as R
participant "VendeurController" as VC
database "SQL Server" as DB

VEN -> R   : Accède à l'espace vendeur
R   -> VC  : GET /api/vendeur/commandes?statut=EN_ATTENTE
VC  -> DB  : SELECT commandes + client + articles
DB  --> VC : Liste commandes
VC  --> R  : JSON commandes
R   --> VEN : Tableau de bord vendeur

VEN -> R   : Consulte le détail
R   -> VC  : GET /api/vendeur/commandes/{{id}}
VC  -> DB  : SELECT commande + lignes + adresse
DB  --> VC : Détail commande
VC  --> R  : JSON détail
R   --> VEN : Affiche détail

VEN -> R   : Confirme la commande
R   -> VC  : PUT /api/vendeur/commandes/{{id}}/confirmer
VC  -> DB  : UPDATE statut = CONFIRMEE
VC  -> DB  : INSERT HistoriqueCommande
DB  --> VC : OK
VC  --> R  : Commande confirmée
R   --> VEN : Statut mis à jour

VEN -> R   : Consulte les indicateurs
R   -> VC  : GET /api/vendeur/stats
VC  -> DB  : SELECT CA, nb commandes, top produits
DB  --> VC : Statistiques
VC  --> R  : JSON stats
R   --> VEN : Indicateurs commerciaux
@enduml"""

# ══════════════════════════════════════════════════════════════
# SPRINT 4
# ══════════════════════════════════════════════════════════════
DIAGRAMS["s4_usecase"] = f"""@startuml
title Diagramme de cas d'utilisation — Sprint 4
{SK_UC}
left to right direction

actor "Confirmatrice" as CONF
actor "Livreur" as LIV
actor "Client" as CLI
actor "Superviseur" as SUP

rectangle "Logistique COD et transit inter-dépôts" {{
  usecase "Consulter les commandes (BC)" as UC1
  usecase "Transformer BC en BL" as UC2
  usecase "Assigner un livreur" as UC3
  usecase "Gérer les réclamations" as UC4
  usecase "Consulter ses livraisons" as UC5
  usecase "Collecter paiement COD" as UC6
  usecase "Mettre à jour position GPS" as UC7
  usecase "Scanner colis (transit)" as UC8
  usecase "Suivre sa commande (GPS)" as UC9
  usecase "Déposer une réclamation" as UC10
  usecase "Voir dashboard logistique" as UC11
}}

CONF --> UC1
CONF --> UC2
CONF --> UC3
CONF --> UC4
LIV  --> UC5
LIV  --> UC6
LIV  --> UC7
LIV  --> UC8
CLI  --> UC9
CLI  --> UC10
SUP  --> UC11

UC2 ..> UC1 : <<include>>
UC6 ..> UC5 : <<include>>
UC7 ..> UC5 : <<include>>
@enduml"""

DIAGRAMS["s4_seq_bctbl"] = f"""@startuml
title Diagramme de séquence — Transformation BC en BL
{SK_SEQ}
actor "Confirmatrice" as CONF
participant "Interface React" as R
participant "ConfirmateurController" as CC
database "SQL Server" as DB

CONF -> R   : Consulte les BC confirmés
R   -> CC  : GET /api/confirmateur/commandes?statut=CONFIRMEE
CC  -> DB  : SELECT F_DOCENTETE (DO_Type=2)
DB  --> CC : Liste BC
CC  --> R  : JSON commandes
R   --> CONF : Liste des BC

CONF -> R   : Clique "Transformer en BL"
R   -> CC  : POST /api/confirmateur/commandes/{{id}}/transformer-bl
CC  -> DB  : Vérifier doublon BL

alt Doublon détecté
  DB  --> CC : BL déjà existant
  CC  --> R  : Erreur: doublon
  R   --> CONF : Message d'erreur
else Pas de doublon
  DB  --> CC : OK
  CC  -> DB  : INSERT F_DOCENTETE (DO_Type=1)
  CC  -> DB  : INSERT F_DOCLIGNE (copie BC)
  CC  -> DB  : UPDATE BC {{statut: TRANSFORME}}
  DB  --> CC : BL créé
  CC  --> R  : {{blId, reference, statut: EN_LIVRAISON}}
  R   --> CONF : BL créé — assignation livreur
end
@enduml"""

DIAGRAMS["s4_seq_gps"] = f"""@startuml
title Diagramme de séquence — Livraison COD et tracking GPS (SignalR)
{SK_SEQ}
actor "Livreur" as LIV
participant "Application Flutter" as FL
participant "SignalR Hub" as SH
participant "LivreurController" as LC
database "SQL Server" as DB
actor "Client" as CLI
participant "Interface React" as RC

LIV -> FL  : Démarre la livraison
FL  -> LC  : POST /api/livreur/livraisons/{{id}}/demarrer
LC  -> DB  : UPDATE BL {{statut: EN_COURS}}
LC  -> SH  : Notifier "client-{{id}}": StatutChange
SH  --> RC : StatutChange — EN_ROUTE
RC  --> CLI : "Livreur en route"

loop Tracking GPS (toutes les 10s)
  FL  -> FL  : GPS.getCurrentPosition()
  FL  -> LC  : POST /api/livreur/ping-batch
  LC  -> DB  : UPSERT F_LIVREUR_POSITION
  LC  -> SH  : Broadcast LocationUpdate {{lat, lng}}
  SH  --> RC : Position mise à jour
  RC  --> CLI : Carte Mapbox actualisée
end

LIV -> FL  : Collecte paiement COD
FL  -> LC  : POST /api/livreur/livraisons/{{id}}/livrer
LC  -> DB  : UPDATE BL {{statut: LIVREE}}
LC  -> DB  : INSERT CaisseCOD {{montant}}
LC  -> SH  : Notifier "client-{{id}}": LIVREE
SH  --> RC : Livraison terminée
RC  --> CLI : Confirmation livraison
@enduml"""

DIAGRAMS["s4_seq_reclamation"] = f"""@startuml
title Diagramme de séquence — Réclamation et gestion SignalR
{SK_SEQ}
actor "Client" as CLI
participant "Interface React" as RC
participant "SignalR Hub" as SH
participant "ReclamationController" as REC
database "SQL Server" as DB
actor "Confirmatrice" as CONF
participant "Interface Confirmatrice" as RCONF

CLI  -> RC   : Soumet une réclamation
RC   -> REC  : POST /api/reclamations
REC  -> DB   : INSERT Reclamation {{OUVERTE}}
REC  -> SH   : Notifier "confirmateurs": NouveauCas
SH   --> RCONF : NouveauCas — alerte
RCONF --> CONF : Nouvelle réclamation reçue
REC  --> RC  : Réclamation créée
RC   --> CLI : Confirmation soumission

CONF -> RCONF : Prend en charge
RCONF -> REC  : PUT /api/confirmateur/reclamations/{{id}}/assigner
REC  -> DB   : UPDATE {{EN_TRAITEMENT, assigneeId}}
REC  -> SH   : Notifier "client-{{id}}": CasAssigne
SH   --> RC  : CasAssigne
RC   --> CLI : "Réclamation prise en charge"

CONF -> RCONF : Résout la réclamation
RCONF -> REC  : PUT /api/confirmateur/reclamations/{{id}}/resoudre
REC  -> DB   : UPDATE {{statut: RESOLUE}}
REC  -> SH   : Notifier "client-{{id}}": StatutChange RESOLUE
SH   --> RC  : Réclamation résolue
RC   --> CLI : Confirmation résolution
@enduml"""

DIAGRAMS["s4_seq_transit"] = f"""@startuml
title Diagramme de séquence — Scan transit inter-dépôts
{SK_SEQ}
actor "Opérateur" as OP
participant "Application Flutter" as FL
participant "TransitController" as TC
database "SQL Server" as DB
actor "Superviseur" as SUP
participant "Interface Admin" as RA

OP  -> FL  : Accède au module transit
FL  -> TC  : GET /api/transit/transferts?statut=EN_ATTENTE
TC  -> DB  : SELECT F_TRANSFERTS
DB  --> TC : Liste transferts
TC  --> FL : JSON transferts
FL  --> OP : Affiche transferts à scanner

OP  -> FL  : Scanne le code-barres
FL  -> FL  : mobile_scanner.scan()
FL  -> TC  : POST /api/transit/scanner
TC  -> DB  : SELECT colis par barcode
DB  --> TC : Colis trouvé
TC  -> DB  : UPDATE F_TRANSFERTS {{RECU_AU_DEPOT}}
TC  -> DB  : INSERT ScanLog
DB  --> TC : OK
TC  --> FL : {{success, colisInfo}}
FL  --> OP : Scan validé

OP  -> FL  : Valide le lot
FL  -> TC  : POST /api/transit/valider-lot
TC  -> DB  : UPDATE F_TRANSFERTS (statut: VALIDE)
DB  --> TC : Lot validé
TC  --> FL : Confirmation
FL  --> OP : Tous les colis validés

SUP -> RA  : Consulte le suivi transit
RA  -> TC  : GET /api/admin/transit/stats
TC  -> DB  : SELECT stats par dépôt
DB  --> TC : Statistiques
TC  --> RA : JSON stats
RA  --> SUP : Tableau de bord transit
@enduml"""

# ══════════════════════════════════════════════════════════════
# SPRINT 5
# ══════════════════════════════════════════════════════════════
DIAGRAMS["s5_usecase"] = f"""@startuml
title Diagramme de cas d'utilisation — Sprint 5
{SK_UC}
left to right direction

actor "Administrateur" as ADM
actor "Superviseur" as SUP
actor "n8n / Groq LLM" as LLM

rectangle "Pilotage, tableaux de bord et chatbot" {{
  usecase "Consulter le dashboard global" as UC1
  usecase "Analyser les ventes et commandes" as UC2
  usecase "Suivre les indicateurs logistiques" as UC3
  usecase "Analyser les réclamations" as UC4
  usecase "Dashboard synchronisation Sage X3" as UC5
  usecase "Exporter en Excel" as UC6
  usecase "Exporter en PDF" as UC7
  usecase "Utiliser le chatbot" as UC8
  usecase "Consulter l'historique chatbot" as UC9
  usecase "Configurer les alertes" as UC10
}}

ADM --> UC1
ADM --> UC2
ADM --> UC3
ADM --> UC4
ADM --> UC5
ADM --> UC6
ADM --> UC7
ADM --> UC8
ADM --> UC9
ADM --> UC10
SUP --> UC1
SUP --> UC3
SUP --> UC4
LLM --> UC8

UC2 ..> UC1 : <<include>>
UC3 ..> UC1 : <<include>>
UC7 ..> UC6 : <<extend>>
@enduml"""

DIAGRAMS["s5_seq_chatbot"] = f"""@startuml
title Diagramme de séquence — Chatbot administrateur (n8n + Groq)
{SK_SEQ}
actor "Administrateur" as ADM
participant "Interface React" as R
participant "ChatbotController" as CC
participant "n8n Workflow" as N8N
participant "Groq LLM" as LLM
database "SQL Server" as DB

ADM -> R   : Saisit une question métier
R   -> CC  : POST /api/chatbot/ask
CC  -> DB  : Récupérer données contextuelles
DB  --> CC : CA, stocks, réclamations
CC  -> N8N : POST webhook {{question, data_context}}
N8N -> N8N : Construit le prompt enrichi
N8N -> LLM : POST /chat/completions (llama3)
LLM --> N8N : Réponse générée
N8N --> CC  : {{answer, sources}}
CC  -> DB   : INSERT ConversationLog
CC  --> R   : {{answer, sources}}
R   --> ADM : Affiche la réponse

ADM -> R   : Consulte l'historique
R   -> CC  : GET /api/chatbot/historique
CC  -> DB  : SELECT ConversationLog
DB  --> CC : Historique
CC  --> R  : JSON historique
R   --> ADM : Liste des conversations
@enduml"""

DIAGRAMS["s5_seq_export"] = f"""@startuml
title Diagramme de séquence — Export des données (Excel / PDF)
{SK_SEQ}
actor "Administrateur" as ADM
participant "Interface React" as R
participant "ExportController" as EC
database "SQL Server" as DB

ADM -> R   : Sélectionne le type d'export
ADM -> R   : Configure les filtres (période, type)
R   -> EC  : POST /api/exports/generer
EC  -> DB  : SELECT données selon filtres
DB  --> EC : Dataset

alt Export Excel
  EC  -> EC  : Construire fichier XLSX (EPPlus)
  EC  --> R  : Stream fichier Excel
  R   --> ADM : Téléchargement .xlsx
else Export PDF
  EC  -> EC  : Générer PDF (iTextSharp)
  EC  --> R  : Stream fichier PDF
  R   --> ADM : Téléchargement .pdf
end

ADM -> R   : Consulte les exports précédents
R   -> EC  : GET /api/exports/historique
EC  -> DB  : SELECT ExportLog
DB  --> EC : Historique exports
EC  --> R  : JSON exports
R   --> ADM : Liste des exports
@enduml"""

# ─── Correspondance placeholder → image ──────────────────────────────────────
# Mapping par mots-clés robustes (ASCII uniquement, insensible à la casse)
PLACEHOLDER_MAP = [
    # (fragment ASCII à rechercher dans le texte normalisé,  nom_image)
    ("utilisation du Sprint 2",        "s2_usecase"),
    ("Consultation du catalogue",      "s2_seq_catalogue"),
    ("Gestion du panier",              "s2_seq_panier"),
    ("ommande client",                 "s2_seq_commande"),
    ("Paiement en ligne via Konnect",  "s2_seq_konnect"),
    ("evis B2B",                       "s2_seq_devis"),      # Sprint 2
    ("utilisation du Sprint 3",        "s3_usecase"),
    ("Synchronisation Sage",           "s3_seq_sage"),
    ("Devis B2B",                      "s3_seq_devis"),      # Sprint 3
    ("Commande vendeur",               "s3_seq_vendeur"),
    ("utilisation du Sprint 4",        "s4_usecase"),
    ("Transformation BC",              "s4_seq_bctbl"),
    ("GPS SignalR",                    "s4_seq_gps"),
    ("clamation et SignalR",           "s4_seq_reclamation"),
    ("transit",                        "s4_seq_transit"),
    ("utilisation du Sprint 5",        "s5_usecase"),
    ("Chatbot",                        "s5_seq_chatbot"),
    ("Export Excel",                   "s5_seq_export"),
]

# ─── Modifications par US ────────────────────────────────────────────────────
MODIFS = {
    # Sprint 2
    "US2.1":  "API: GET /api/articles\nReact: CataloguePage, ProductGrid\nFlutter: —",
    "US2.2":  "API: GET /api/articles?search=\nReact: SearchBar\nFlutter: SearchWidget",
    "US2.3":  "API: GET /api/articles/{id}\nReact: ProductDetailPage\nFlutter: ProductDetailScreen",
    "US2.4":  "API: POST /api/articles/comparer\nReact: CompareModal\nFlutter: —",
    "US2.5":  "API: GET/POST/DELETE /api/favoris\nReact: FavoritesPage\nFlutter: FavoritesScreen",
    "US2.6":  "API: CRUD /api/panier\nReact: CartPage, CartIcon\nFlutter: CartScreen",
    "US2.7":  "API: POST /api/commandes\nReact: CheckoutPage\nFlutter: —",
    "US2.8":  "API: POST /api/paiement/initier, webhook\nReact: PaymentPage\nFlutter: —",
    "US2.9":  "API: GET /api/commandes/{id}\nReact: OrderTrackingPage\nFlutter: OrderTrackingScreen",
    "US2.10": "API: CRUD /api/adresses\nReact: AddressManagement\nFlutter: —",
    "US2.11": "API: POST /api/devis\nReact: DevisForm\nFlutter: —",
    "US2.12": "API: POST /api/auth/login\nReact: LoginPage\nFlutter: LoginScreen",
    # Sprint 3
    "US3.1":  "API: CRUD /api/admin/articles\nReact: AdminProductsPage\nFlutter: —",
    "US3.2":  "API: POST /api/admin/articles/{id}/images\nReact: ImageUploader (Cloudinary)\nFlutter: —",
    "US3.3":  "API: POST /api/articles/sync-sage\nReact: SyncSagePage\nFlutter: —",
    "US3.4":  "API: POST /api/articles/sync-catalogues\nReact: SyncCataloguePage\nFlutter: —",
    "US3.5":  "API: POST /api/admin/depots/sync-sage\nReact: DepotsPage\nFlutter: —",
    "US3.6":  "API: POST /api/admin/stocks/sync-sage\nReact: StockPage\nFlutter: —",
    "US3.7":  "API: CRUD /api/admin/clients\nReact: ClientsB2BPage\nFlutter: —",
    "US3.8":  "API: CRUD /api/admin/b2b/remises\nReact: RemisesB2BPage\nFlutter: —",
    "US3.9":  "API: GET /api/vendeur/catalogue\nReact: VendeurDashboard\nFlutter: —",
    "US3.10": "API: POST /api/vendeur/commandes\nReact: CreateCommandePage\nFlutter: —",
    "US3.11": "API: POST /api/devis\nReact: DevisForm B2B\nFlutter: —",
    "US3.12": "API: CRUD /api/admin/devis\nReact: DevisManagementPage\nFlutter: —",
    "US3.13": "API: CRUD /api/admin/depots\nReact: DepotsMapPage (Mapbox)\nFlutter: —",
    "US3.14": "API: PUT /api/admin/homepage\nReact: HomepageConfigPage\nFlutter: —",
    "US3.15": "API: CRUD /api/admin/settings\nReact: SettingsPage\nFlutter: —",
    "US3.16": "API: GET /api/admin/dashboard\nReact: InitialDashboard\nFlutter: —",
    # Sprint 4
    "US4.1":  "API: GET /api/confirmateur/commandes\nReact: CommandesTable\nFlutter: —",
    "US4.2":  "API: GET /api/confirmateur/commandes/{id}\nReact: CommandeDetailModal\nFlutter: —",
    "US4.3":  "API: POST /api/confirmateur/.../transformer-bl\nReact: TransformerBLButton\nFlutter: —",
    "US4.4":  "API: PUT /api/confirmateur/.../assigner\nReact: LivreurSelector\nFlutter: —",
    "US4.5":  "API: GET /api/livreur/livraisons\nReact: —\nFlutter: LivraisonsListScreen",
    "US4.6":  "API: POST /api/livreur/livraisons/{id}/livrer\nReact: —\nFlutter: LivraisonDetailScreen",
    "US4.7":  "API: POST /api/livreur/ping-batch (SignalR)\nReact: TrackingMap (Mapbox)\nFlutter: GPSTrackerService",
    "US4.8":  "API: POST /api/reclamations\nReact: ReclamationForm\nFlutter: ReclamationScreen",
    "US4.9":  "API: CRUD /api/confirmateur/reclamations (SignalR)\nReact: ReclamationDashboard\nFlutter: —",
    "US4.10": "API: GET /api/livreur/caisse-cod\nReact: —\nFlutter: CaisseScreen",
    "US4.11": "API: POST /api/transit/scanner, /valider-lot\nReact: —\nFlutter: ScannerScreen (mobile_scanner)",
    "US4.12": "API: GET /api/admin/transit/stats\nReact: TransitDashboard\nFlutter: —",
    # Sprint 5
    "US5.1":  "API: GET /api/admin/dashboard\nReact: GlobalDashboard\nFlutter: —",
    "US5.2":  "API: GET /api/dashboard/ventes\nReact: VentesChart, CATimeline\nFlutter: —",
    "US5.3":  "API: GET /api/dashboard/logistique\nReact: LogistiqueDashboard\nFlutter: —",
    "US5.4":  "API: GET /api/dashboard/reclamations\nReact: ReclamationsDashboard\nFlutter: —",
    "US5.5":  "API: GET /api/dashboard/sage-x3\nReact: SageDashboard\nFlutter: —",
    "US5.6":  "API: POST /api/exports/generer (XLSX)\nReact: ExportButton\nFlutter: —",
    "US5.7":  "API: POST /api/exports/generer (PDF)\nReact: PDFExportButton\nFlutter: —",
    "US5.8":  "API: POST /api/chatbot/ask\nReact: ChatbotInterface\nFlutter: —",
    "US5.9":  "API: GET /api/chatbot/historique\nReact: ChatHistory\nFlutter: —",
    "US5.10": "API: CRUD /api/admin/alertes\nReact: AlertConfig\nFlutter: —",
}

# ─── Étape 1 : Générer les PNG ────────────────────────────────────────────────
def step1_generate():
    print("═" * 60)
    print("ÉTAPE 1 — Génération des diagrammes PlantUML")
    print("═" * 60)
    ok = 0
    for name, content in DIAGRAMS.items():
        puml = DIAG_DIR / f"{name}.puml"
        png  = DIAG_DIR / f"{name}.png"
        puml.write_text(content, encoding="utf-8")
        res = subprocess.run(
            ["java", "-jar", PLANTUML, "-tpng", str(puml), "-o", str(DIAG_DIR)],
            capture_output=True, text=True
        )
        if res.returncode == 0 and png.exists() and png.stat().st_size > 5000:
            print(f"  ✓ {name}.png  ({png.stat().st_size:,} bytes)")
            ok += 1
        else:
            err = (res.stderr or res.stdout or "")[:120]
            print(f"  ✗ {name}  →  {err}")
    print(f"\n  Total: {ok}/{len(DIAGRAMS)} diagrammes générés")
    return ok

# ─── Étape 2 : Mise à jour du document ───────────────────────────────────────

def _clear_para(para):
    """Vide un paragraphe de tout son contenu XML."""
    for child in list(para._p):
        para._p.remove(child)

def _center_para(para):
    """Ajoute un alignement centré."""
    pPr = OxmlElement("w:pPr")
    jc  = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    para._p.insert(0, pPr)

def insert_image(para, img_path, width_in=5.8):
    """Insère une image dans un paragraphe vidé et centré."""
    _clear_para(para)
    _center_para(para)
    run = para.add_run()
    run.add_picture(str(img_path), width=Inches(width_in))


# ── Remplacement colonne dans un tableau ─────────────────────────────────────
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _make_tc(text, width_dxa, bold=False, font_size=20, wrap=True):
    """Crée un élément <w:tc> complet."""
    tc  = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)
    # Marge interne
    tcMar = OxmlElement("w:tcMar")
    for side in ("top","left","bottom","right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "80" if side in ("top","bottom") else "120")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)
    tc.append(tcPr)

    p  = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    if not wrap:
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "Normal")
        pPr.append(pStyle)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    pPr.append(jc)
    p.append(pPr)

    lines = text.split("\n")
    for li, line in enumerate(lines):
        if li == 0:
            r   = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            if bold:
                rPr.append(OxmlElement("w:b"))
                rPr.append(OxmlElement("w:bCs"))
            for sz_tag in ("w:sz", "w:szCs"):
                sz = OxmlElement(sz_tag)
                sz.set(qn("w:val"), str(font_size))
                rPr.append(sz)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = line
            r.append(t)
            p.append(r)
        else:
            # Saut de ligne dans la même cellule
            br = OxmlElement("w:br")
            r2 = OxmlElement("w:r")
            rPr2 = OxmlElement("w:rPr")
            for sz_tag in ("w:sz","w:szCs"):
                sz = OxmlElement(sz_tag)
                sz.set(qn("w:val"), str(font_size))
                rPr2.append(sz)
            r2.append(rPr2)
            r_br = OxmlElement("w:r")
            r_br.append(br)
            p.append(r_br)
            r3 = OxmlElement("w:r")
            rPr3 = OxmlElement("w:rPr")
            for sz_tag in ("w:sz","w:szCs"):
                sz = OxmlElement(sz_tag)
                sz.set(qn("w:val"), str(font_size))
                rPr3.append(sz)
            r3.append(rPr3)
            t3 = OxmlElement("w:t")
            t3.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t3.text = line
            r3.append(t3)
            p.append(r3)
    tc.append(p)
    return tc


def rebuild_table_sprint2(table):
    """
    Sprint 2 — 5 colonnes : ID | User Story | Modifications | Responsable | Estimation
    (remplace la colonne 'Tâches principales' par 'Modifications')
    Largeurs (dxa) : 600 | 3000 | 3200 | 1560 | 1000  = 9360
    """
    WIDTHS = [600, 3000, 3200, 1560, 1000]
    HEADERS = ["ID", "User Story", "Modifications\n(Web API / React / Flutter)", "Responsable", "Estimation"]
    tbl = table._tbl
    rows = tbl.findall(f"{{{W}}}tr")

    for ri, row in enumerate(rows):
        cells = row.findall(f"{{{W}}}tc")
        # col 0 = ID, col 1 = US, col 2 = Tâches→Modif, col 3 = Resp, col 4 = Estim
        if ri == 0:
            texts = [HEADERS[0], HEADERS[1], HEADERS[2], HEADERS[3], HEADERS[4]]
            bold  = [True]*5
        else:
            us_id = cells[0].findtext(f".//{{{W}}}t", default="").strip()
            texts = [
                cells[0].findtext(f".//{{{W}}}t", default="").strip(),
                " ".join(t.text or "" for t in cells[1].findall(f".//{{{W}}}t")).strip(),
                MODIFS.get(us_id, "API: —\nReact: —\nFlutter: —"),
                cells[3].findtext(f".//{{{W}}}t", default="").strip(),
                cells[4].findtext(f".//{{{W}}}t", default="").strip(),
            ]
            bold = [False]*5

        # Reconstruire les cellules
        for ci, (old_tc, new_text, new_bold, width) in enumerate(zip(cells, texts, bold, WIDTHS)):
            new_tc = _make_tc(new_text, width, bold=new_bold)
            row.replace(old_tc, new_tc)


def rebuild_table_sprint345(table, sprint_prefix):
    """
    Sprint 3/4/5 — 5 colonnes : ID | User Story | Rôle | Endpoints / Modules | Modifications
    (supprime Resp. et Estim., ajoute Modifications)
    Largeurs (dxa) : 600 | 2600 | 1200 | 2500 | 2460 = 9360
    """
    WIDTHS  = [600, 2600, 1200, 2500, 2460]
    HEADERS = ["ID", "User Story", "Rôle", "Endpoints / Modules",
               "Modifications\n(Web API / React / Flutter)"]
    tbl  = table._tbl
    rows = tbl.findall(f"{{{W}}}tr")

    for ri, row in enumerate(rows):
        cells = row.findall(f"{{{W}}}tc")
        if ri == 0:
            texts = HEADERS
            bold  = [True]*5
        else:
            us_id = cells[0].findtext(f".//{{{W}}}t", default="").strip()
            texts = [
                us_id,
                " ".join(t.text or "" for t in cells[1].findall(f".//{{{W}}}t")).strip(),
                " ".join(t.text or "" for t in cells[2].findall(f".//{{{W}}}t")).strip(),
                " ".join(t.text or "" for t in cells[3].findall(f".//{{{W}}}t")).strip(),
                MODIFS.get(us_id, "API: —\nReact: —\nFlutter: —"),
            ]
            bold = [False]*5

        # Supprimer toutes les cellules existantes et reconstruire en 5
        for old_tc in list(cells):
            row.remove(old_tc)
        for text, bold_flag, width in zip(texts, bold, WIDTHS):
            row.append(_make_tc(text, width, bold=bold_flag))

    # Mettre à jour tblGrid
    tblGrid = tbl.find(f"{{{W}}}tblGrid")
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in WIDTHS:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    # Insérer après tblPr ou en premier
    tblPr = tbl.find(f"{{{W}}}tblPr")
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def fix_names(doc):
    """Corrige Melek → Melek Lahmar et Tawfik → Tawfik Siala partout."""
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            orig = run.text
            if "Melek" in run.text and "Lahmar" not in run.text:
                run.text = run.text.replace("Melek", "Melek Lahmar"); count += 1
            if "Tawfik" in run.text and "Siala" not in run.text:
                run.text = run.text.replace("Tawfik", "Tawfik Siala"); count += 1
    return count


def step2_update_doc(n_diag):
    print("\n" + "═" * 60)
    print("ÉTAPE 2 — Mise à jour du document Word")
    print("═" * 60)

    doc = Document(CLEAN_DOC)
    paras = doc.paragraphs

    # 2a. Insertion des diagrammes
    inserted = 0
    used_names = set()
    for i, para in enumerate(paras):
        text = para.text.strip()
        if "Insérer ici" not in text or "diagramme" not in text.lower():
            continue
        img_name = None
        for keyword, name in PLACEHOLDER_MAP:
            if name not in used_names and keyword.lower() in text.lower():
                img_name = name
                used_names.add(name)
                break
        if not img_name:
            print(f"  ⚠ Pas de mapping pour: {text[:70]}")
            continue
        img_path = DIAG_DIR / f"{img_name}.png"
        if not img_path.exists():
            print(f"  ✗ Image manquante: {img_name}.png")
            continue
        insert_image(para, img_path, width_in=5.8)
        print(f"  ✓ Para {i}: {img_name}.png")
        inserted += 1

    print(f"\n  {inserted}/{len(PLACEHOLDER_MAP)} diagrammes insérés")

    # 2b. Reconstruction des tableaux backlog
    print("\n  Reconstruction des tableaux backlog...")
    rebuild_table_sprint2(doc.tables[4])
    print("  ✓ Table 4  (Sprint 2 — 5 cols: ID|US|Modifications|Resp|Estim)")
    rebuild_table_sprint345(doc.tables[6],  "US3")
    print("  ✓ Table 6  (Sprint 3 — 5 cols: ID|US|Rôle|Endpoints|Modifications)")
    rebuild_table_sprint345(doc.tables[8],  "US4")
    print("  ✓ Table 8  (Sprint 4 — 5 cols: ID|US|Rôle|Endpoints|Modifications)")
    rebuild_table_sprint345(doc.tables[10], "US5")
    print("  ✓ Table 10 (Sprint 5 — 5 cols: ID|US|Rôle|Endpoints|Modifications)")

    # 2c. Correction des noms
    n_names = fix_names(doc)
    print(f"\n  ✓ Noms corrigés: {n_names} occurrences")

    # 2d. Sauvegarde
    doc.save(OUT_DOC)
    import os
    size = os.path.getsize(OUT_DOC)
    print(f"\n  ✓ Document sauvegardé: {OUT_DOC}")
    print(f"    Taille: {size/1024/1024:.1f} Mo")
    return inserted


# ─── Main ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    n1 = step1_generate()
    n2 = step2_update_doc(n1)
    print("\n" + "═" * 60)
    print(f"TERMINÉ — {n1} diagrammes générés, {n2} insérés dans le document.")
    print("═" * 60)
