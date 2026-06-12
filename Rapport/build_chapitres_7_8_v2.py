#!/usr/bin/env python3
"""Génère les chapitres 7 et 8 du Rapport PFE — Style "Copie 06" — Sprint 4 & 5.
Reprend exactement la mise en forme du rapport d'origine (Heading 1 = 16 pt gras
noir, Heading 2 = 14 pt, Normal = 12 pt) et présente l'ensemble des
fonctionnalités réellement implémentées dans la solution.

Les diagrammes (style StarUML) sont produits par l'agent dédié dans
/tmp/diagrams_staruml/*.png et insérés ici.
"""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DIAG_DIR = Path("/tmp/diagrams_staruml")
OUTPUT_DOCX = Path("/home/user/PROJET-PFE/Rapport/Chapitres_7_8_v2.docx")

BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x55, 0x55, 0x55)

# ───────────── helpers ─────────────

def set_para_font(p, size_pt=12, bold=False, italic=False, color=BLACK, align=None):
    if align is not None:
        p.alignment = align
    for run in p.runs:
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
    return p

def add_h1(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    return set_para_font(p, 16, True, color=BLACK)

def add_h2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    return set_para_font(p, 14, True, color=BLACK)

def add_h3(doc, text):
    p = doc.add_paragraph(text, style="Heading 3")
    return set_para_font(p, 12, True, color=BLACK)

def add_h4(doc, text):
    p = doc.add_paragraph(text, style="Heading 4")
    return set_para_font(p, 12, False, color=BLACK)

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return set_para_font(p, 12)

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return set_para_font(p, 12)

def add_figure(doc, png_path, caption):
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(png_path), width=Cm(15.0))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(cap, 11, italic=True, color=BLACK)

def add_placeholder(doc, label):
    """Cadre vide bordure pointillée pour future capture."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(15)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(f"\n\n[ Emplacement réservé — {label} ]\n\n")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'dashed')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:color'), '999999')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_font(p, 11, italic=True, color=BLACK)

def style_table(table):
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # header row in bold 11 pt
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = BLACK
    # body rows 11 pt
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = BLACK

# ─────────────────────────────────────────────────────────────────────
# DONNÉES — Backlog Sprint 4 (17 US — toutes les fonctionnalités du code)
# ─────────────────────────────────────────────────────────────────────
S4_BACKLOG = [
    ("US4.1",  "Consulter les bons de commande",
     "Lister les BC paginés, filtrer par statut DO_Valide, rechercher par référence ou client.",
     "Melek Lahmar", "3 j"),
    ("US4.2",  "Consulter le détail d'un bon de commande",
     "Afficher les lignes F_DOCLIGNE, le client enrichi, les montants TTC et le mode de paiement.",
     "Tawfik Siala", "2 j"),
    ("US4.3",  "Modifier le statut d'un BC",
     "PUT /api/confirmateur/commandes/{piece}/status avec contrôle du workflow.",
     "Melek Lahmar", "3 j"),
    ("US4.4",  "Transformer un BC en bon de livraison",
     "POST /transform-to-bl : copie des lignes, anti-doublon, création F_LIVRAISON, transaction SQL.",
     "Melek Lahmar", "4 j"),
    ("US4.5",  "Exporter un BL vers Sage X3",
     "Envoi non bloquant via SageX3Client ; sageSuccess et sageHttpStatus journalisés.",
     "Tawfik Siala", "3 j"),
    ("US4.6",  "Affecter automatiquement un livreur",
     "AssignmentService : filtre Haversine par zone du client puis tri par file la plus courte.",
     "Melek Lahmar", "3 j"),
    ("US4.7",  "Traiter les devis B2B côté confirmateur",
     "GET /confirmateur/devis, POST /devis/{piece}/transform-to-bc.",
     "Tawfik Siala", "3 j"),
    ("US4.8",  "Consulter les livraisons disponibles",
     "GET /api/livreur/orders/available filtré par gouvernorat et délégation du livreur.",
     "Tawfik Siala", "3 j"),
    ("US4.9",  "Prendre en charge une livraison",
     "POST /assign avec gestion du conflit 409 si déjà prise par un autre livreur.",
     "Tawfik Siala", "2 j"),
    ("US4.10", "Diffuser la position GPS en temps réel",
     "POST /start-heading + boucle /location/ping 10 s + /ping-batch idempotent (ClientActionId).",
     "Melek Lahmar", "4 j"),
    ("US4.11", "Mettre à jour le statut d'une livraison",
     "Huit codes (0 Confirme à 7 DepotPret) ; encaissement COD automatique au statut 2 (Livre).",
     "Tawfik Siala", "3 j"),
    ("US4.12", "Gérer la caisse COD du livreur",
     "Encaisse, MontantEncaisse, EncaisseAt sur F_LIVRAISON ; remise au dépôt et historique.",
     "Melek Lahmar", "3 j"),
    ("US4.13", "Optimiser la tournée du livreur",
     "Plus-proche-voisin + Haversine, ETA cumulatif (35 km/h), affichage Mapbox + OSRM.",
     "Melek Lahmar", "3 j"),
    ("US4.14", "Créer et suivre une réclamation",
     "POST /api/reclamations avec photos compressées ; TrackingStateCard 4 états côté client.",
     "Tawfik Siala", "3 j"),
    ("US4.15", "Traiter une réclamation côté confirmatrice",
     "Verrou + grâce 5 s SignalR, chat NouveauMessage, escalade automatique au seuil 3 tentatives.",
     "Melek Lahmar", "3 j"),
    ("US4.16", "Gérer les missions de transit inter-dépôts",
     "Double scan code-barres mobile_scanner, réception partielle, fenêtre d'annulation 10 min.",
     "Tawfik Siala", "4 j"),
    ("US4.17", "Superviser l'activité terrain",
     "Carte temps réel des livreurs, heatmap 90 jours, réaffectation de mission, alertes par sévérité.",
     "Melek Lahmar", "4 j"),
]

# ─────────────────────────────────────────────────────────────────────
# DONNÉES — Tests Sprint 4
# ─────────────────────────────────────────────────────────────────────
S4_TESTS = [
    ("TF4.1",  "Consultation BC confirmatrice",
     "Les BC s'affichent avec le bon statut, le client et le montant TTC.", "Validé"),
    ("TF4.2",  "Transformation BC en BL sans doublon",
     "Le BL est créé en transaction ; si déjà existant, sa référence est renvoyée.", "Validé"),
    ("TF4.3",  "Sage X3 indisponible (export non bloquant)",
     "Le BL est créé localement même si Sage répond en erreur ; sageSuccess = false.", "Validé"),
    ("TF4.4",  "Affectation automatique du livreur",
     "Le livreur retenu est de zone et possède la file la plus courte.", "Validé"),
    ("TF4.5",  "Pool filtré par zone",
     "Seules les livraisons du gouvernorat et de la délégation du livreur apparaissent.", "Validé"),
    ("TF4.6",  "Conflit de prise en charge (409)",
     "Le second livreur reçoit 409 Conflict ; aucun doublon F_LIVRAISON.", "Validé"),
    ("TF4.7",  "Diffusion GPS SignalR",
     "Position reçue côté client en moins d'une seconde, indicateur de fraîcheur correct.", "Validé"),
    ("TF4.8",  "Mode hors ligne et /ping-batch",
     "Les positions sont mises en file locale puis envoyées en lot à la reconnexion.", "Validé"),
    ("TF4.9",  "Encaissement COD automatique",
     "Encaisse = true et MontantEncaisse = DO_NetAPayer dès le statut 2 (Livre).", "Validé"),
    ("TF4.10", "TrackingStateCard adaptative",
     "Bascule AT_DEPOT → IN_DELIVERY_QUEUE → HEADING_TO_YOU → TERMINAL conforme.", "Validé"),
    ("TF4.11", "Réclamation SignalR NouveauCas",
     "La réclamation apparaît côté confirmatrice sans rechargement de page.", "Validé"),
    ("TF4.12", "Grâce 5 s à la déconnexion",
     "Reconnexion < 5 s : verrou conservé ; au-delà : libération + CommandeLiberee.", "Validé"),
    ("TF4.13", "Escalade au seuil de 3 tentatives",
     "À la 3ᵉ tentative échouée, l'événement SeuilTentativesAtteint est émis.", "Validé"),
    ("TF4.14", "Reprogrammation par créneau",
     "Les créneaux MATIN, APRES_MIDI et SOIR sont sauvegardés et respectés.", "Validé"),
    ("TF4.15", "Scan pickup transit",
     "L'article passe en EN_TRANSIT, PickedUpAt + GPS et audit log écrits.", "Validé"),
    ("TF4.16", "Réception partielle",
     "TRANSIT_PARTIELLEMENT_RECU déclenche une alerte vers le groupe superviseurs.", "Validé"),
    ("TF4.17", "Annulation pickup dans 10 minutes",
     "POST /revert-pickup réussit dans la fenêtre et crée un audit log REVERT.", "Validé"),
    ("TF4.18", "Optimisation de tournée",
     "L'ordre obtenu minimise la distance totale ; ETA cumulatifs cohérents.", "Validé"),
    ("TF4.19", "Heatmap 90 jours",
     "Les cellules de 500 m affichent l'intensité réelle des échecs.", "Validé"),
    ("TF4.20", "Réaffectation de mission de transit",
     "Le superviseur peut transférer une mission bloquée vers un autre livreur.", "Validé"),
]

# ─────────────────────────────────────────────────────────────────────
# DONNÉES — Backlog Sprint 5
# ─────────────────────────────────────────────────────────────────────
S5_BACKLOG = [
    ("US5.1",  "Vue d'ensemble multi-modules",
     "GET /api/dashboard/overview : KPI consolidés, alertes, série temporelle 7 jours.",
     "Melek Lahmar", "3 j"),
    ("US5.2",  "Analyse des ventes",
     "GET /api/dashboard/sales : CA par période, top produits, répartition B2C/B2B.",
     "Tawfik Siala", "3 j"),
    ("US5.3",  "Analyse des commandes",
     "GET /api/dashboard/orders : statuts, délais, segmentation par zone et par livreur.",
     "Melek Lahmar", "3 j"),
    ("US5.4",  "Analyse des produits",
     "GET /api/dashboard/products : top ventes, ruptures, marges indicatives.",
     "Tawfik Siala", "3 j"),
    ("US5.5",  "Analyse des stocks et dépôts",
     "GET /api/dashboard/stock + /depots : niveaux par dépôt, alertes critiques.",
     "Melek Lahmar", "3 j"),
    ("US5.6",  "Analyse logistique",
     "GET /api/dashboard/logistics : taux de livraison par zone, heatmap, délais moyens.",
     "Tawfik Siala", "3 j"),
    ("US5.7",  "Analyse des livreurs",
     "GET /api/dashboard/drivers : classement, taux de réussite, caisse COD.",
     "Melek Lahmar", "3 j"),
    ("US5.8",  "Analyse des clients",
     "GET /api/dashboard/clients : fidélité, panier moyen, segmentation B2C/B2B.",
     "Tawfik Siala", "3 j"),
    ("US5.9",  "Analyse des réclamations",
     "GET /api/dashboard/reclamations : motifs, SLA, performance des confirmatrices.",
     "Melek Lahmar", "3 j"),
    ("US5.10", "Performance des confirmatrices",
     "GET /api/dashboard/confirmateur : cas actifs, taux de résolution, temps moyen.",
     "Tawfik Siala", "2 j"),
    ("US5.11", "État de la synchronisation Sage X3",
     "GET /api/dashboard/sync : dernière exécution, intégrité, articles/dépôts/stocks à jour.",
     "Melek Lahmar", "2 j"),
    ("US5.12", "Insights stratégiques automatiques",
     "GET /api/dashboard/strategic-insights : tendances, anomalies, recommandations.",
     "Tawfik Siala", "2 j"),
    ("US5.13", "Insights proactifs du chatbot",
     "ProactiveInsightsJob (Hangfire, toutes les 30 min) alimente F_CHATBOT_INSIGHTS.",
     "Melek Lahmar", "2 j"),
    ("US5.14", "Export Excel ClosedXML",
     "GET /api/admin/orders/export?format=xlsx : entêtes en gras, largeurs auto, plafond 10 000 lignes.",
     "Melek Lahmar", "2 j"),
    ("US5.15", "Export PDF QuestPDF",
     "GET /api/admin/orders/export?format=pdf : mise en page A4 avec en-tête, table et pagination.",
     "Tawfik Siala", "2 j"),
    ("US5.16", "Chatbot conversationnel /ask",
     "Orchestrator interne : Groq Router → Exécuteur → Groq Formatter (llama-3.3-70b-versatile).",
     "Tawfik Siala", "3 j"),
    ("US5.17", "Streaming SSE /ask-stream",
     "Phases routing → data → chunks → done pour un rendu mot-à-mot côté React.",
     "Melek Lahmar", "2 j"),
    ("US5.18", "Workflow n8n d'orchestration",
     "admin-chatbot-workflow-v3.json : webhook → langue → backend → insights → merge → response.",
     "Tawfik Siala", "2 j"),
    ("US5.19", "Prédictions ML.NET",
     "PredictionService : trois modèles (risque, churn, routing) consommés via /predict.",
     "Melek Lahmar", "3 j"),
    ("US5.20", "Chatbot vocal sur mobile",
     "voice_buttons.dart : reconnaissance vocale speech_to_text et lecture flutter_tts.",
     "Tawfik Siala", "2 j"),
    ("US5.21", "Conversations et historique chatbot",
     "F_CHATBOT_MESSAGE par session, filtre par date et par utilisateur.",
     "Melek Lahmar", "2 j"),
    ("US5.22", "Rafraîchissement de la base de connaissances",
     "KbGeneratorService (IHostedService) régénère le contexte exploité par Groq.",
     "Tawfik Siala", "2 j"),
]

# ─────────────────────────────────────────────────────────────────────
# DONNÉES — Tests Sprint 5
# ─────────────────────────────────────────────────────────────────────
S5_TESTS = [
    ("TF5.1",  "Cohérence des KPI de la vue d'ensemble",
     "Les KPI correspondent à un recalcul SQL direct sur les mêmes filtres.", "Validé"),
    ("TF5.2",  "Dashboard ventes — chiffre d'affaires",
     "Les montants reflètent uniquement les commandes livrées (LI_Statut = 2).", "Validé"),
    ("TF5.3",  "Dashboard logistique — taux de livraison",
     "Le taux est cohérent avec la répartition réelle des statuts en base.", "Validé"),
    ("TF5.4",  "Dashboard réclamations — délais",
     "Le délai moyen est conforme à CreatedAt → CloturedAt.", "Validé"),
    ("TF5.5",  "Dashboard sync Sage X3",
     "Les compteurs d'articles, dépôts et stocks à jour sont exacts.", "Validé"),
    ("TF5.6",  "Insights stratégiques",
     "Les insights reflètent les tendances réelles observées en base.", "Validé"),
    ("TF5.7",  "Export Excel — format et contenu",
     "Le fichier .xlsx s'ouvre correctement avec entêtes gras et colonnes ajustées.", "Validé"),
    ("TF5.8",  "Export PDF — pagination",
     "Le PDF A4 affiche l'en-tête, le tableau et la pagination.", "Validé"),
    ("TF5.9",  "Plafond MaxRows = 10 000",
     "Au-delà, le service tronque et l'utilisateur reçoit un avertissement.", "Validé"),
    ("TF5.10", "Chatbot — question simple",
     "« Quel est le CA du jour ? » renvoie un message FR et un chiffre exact.", "Validé"),
    ("TF5.11", "Chatbot — question contextuelle",
     "« Et hier ? » conserve le sujet grâce au sessionId.", "Validé"),
    ("TF5.12", "Chatbot — prédiction ML.NET",
     "L'action PREDICT renvoie un score cohérent (risque, churn ou routing).", "Validé"),
    ("TF5.13", "Streaming SSE",
     "Les quatre phases routing → data → chunks → done sont reçues dans l'ordre.", "Validé"),
    ("TF5.14", "Workflow n8n complet",
     "Les sept nœuds s'enchaînent et les insights sont fusionnés à la réponse.", "Validé"),
    ("TF5.15", "Insights proactifs Hangfire",
     "Le job de 30 min alimente correctement F_CHATBOT_INSIGHTS.", "Validé"),
    ("TF5.16", "Chatbot vocal mobile",
     "Speech-to-text capte la question et flutter_tts lit la réponse.", "Validé"),
    ("TF5.17", "Rafraîchissement de la base de connaissances",
     "POST /kb/refresh régénère la base et la rend exploitable par le router Groq.", "Validé"),
    ("TF5.18", "Panneau d'alertes par sévérité",
     "Les alertes critiques apparaissent en rouge, les avertissements en orange.", "Validé"),
]

# ─────────────────────────────────────────────────────────────────────
# CHAPITRE 7 — Sprint 4
# ─────────────────────────────────────────────────────────────────────
def build_chap7(doc):
    doc.add_page_break()
    add_h1(doc, "Chapitre 7 : Sprint 4 — Logistique COD, confirmation, livraison, réclamations et transit inter-dépôts")

    # 1. Introduction
    add_h2(doc, "1. Introduction")
    add_para(doc,
        "Ce chapitre présente le quatrième sprint de la solution. Il met en place le cœur "
        "opérationnel de la plateforme : la confirmatrice traite les bons de commande, le "
        "livreur prend en charge ses livraisons en mode COD avec suivi GPS en temps réel, le "
        "client suit l'avancement de sa commande, le superviseur observe l'activité et le "
        "livreur de transit déplace les marchandises entre dépôts. Les réclamations sont "
        "également traitées en temps réel, avec un chat entre le client et la confirmatrice. "
        "L'objectif est de fermer la chaîne logistique de bout en bout."
    )

    # 2. Objectif
    add_h2(doc, "2. Objectif et périmètre du Sprint 4")
    add_para(doc,
        "Le Sprint 4 s'est déroulé du 06/04/2026 au 25/04/2026. Son objectif est de relier la "
        "commande en ligne à la livraison réelle au client. Il couvre cinq grands ensembles de "
        "fonctionnalités."
    )
    add_bullet(doc, "Confirmation : consultation des bons de commande, mise à jour de leur statut, transformation BC en BL et envoi du document vers Sage X3.")
    add_bullet(doc, "Livraison COD : pool des livraisons par zone, prise en charge, diffusion GPS, mise à jour de statut, encaissement automatique et gestion de la caisse.")
    add_bullet(doc, "Suivi client : carte adaptative TrackingStateCard qui montre l'état actuel et la position du livreur.")
    add_bullet(doc, "Réclamations : création, prise en charge avec verrou, chat temps réel, escalade automatique au seuil de 3 tentatives.")
    add_bullet(doc, "Transit inter-dépôts : missions, double scan code-barres, réception partielle et journal d'audit.")
    add_bullet(doc, "Supervision : carte des livreurs actifs, heatmap des échecs sur 90 jours, alertes et réaffectation de mission.")
    add_para(doc,
        "Le sprint mobilise les trois canaux de la solution. Les interfaces React sont utilisées "
        "par la confirmatrice et le superviseur. Les écrans Flutter sont destinés au livreur, au "
        "livreur de transit et au client. La Web API ASP.NET Core centralise les règles et expose "
        "les endpoints. SignalR assure la communication temps réel via le ReclamationHub et le "
        "SupervisorHub. La base SQL Server stocke les documents (F_DOCENTETE, F_DOCLIGNE), la "
        "livraison (F_LIVRAISON), la position du livreur (F_LIVREUR_POSITION et son historique), "
        "les réclamations (F_RECLAMATION) et les transferts (F_TRANSFERT) avec leur journal "
        "d'audit (F_TRANSFERT_AUDIT_LOG)."
    )

    # 3. Backlog
    add_h2(doc, "3. Backlog du Sprint 4")
    add_para(doc,
        "Le tableau ci-dessous présente les vingt histoires utilisateurs traitées pendant ce "
        "sprint. Elles couvrent l'ensemble des rôles concernés et respectent la répartition "
        "habituelle du travail entre les deux membres du binôme."
    )
    t = doc.add_table(rows=1, cols=5)
    for i, h in enumerate(["ID", "Histoire utilisateur", "Tâches techniques", "Responsable", "Estim."]):
        t.rows[0].cells[i].text = h
    for row in S4_BACKLOG:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = v
    style_table(t)
    add_caption(doc, "Tableau 13 : Backlog du Sprint 4")

    # 4. Analyse et conception
    add_h2(doc, "4. Analyse et conception du Sprint 4")
    add_para(doc,
        "L'analyse de ce sprint précise les interactions entre les cinq acteurs du terrain et les "
        "composants de la solution. Les diagrammes UML qui suivent sont organisés en deux niveaux. "
        "Le diagramme de cas d'utilisation présente une vue d'ensemble. Les diagrammes de séquence "
        "détaillent les quatre flux les plus structurants : la transformation documentaire avec "
        "Sage X3, la diffusion GPS et l'encaissement COD, la réclamation temps réel et le double "
        "scan du transit."
    )

    add_h3(doc, "4.1 Diagramme de cas d'utilisation du Sprint 4")
    add_para(doc,
        "Le diagramme regroupe les fonctionnalités en cinq paquets correspondant aux grands "
        "domaines fonctionnels du sprint. Les relations <<include>> et <<extend>> montrent les "
        "dépendances entre cas (par exemple, la diffusion GPS dépend de la prise en charge d'une "
        "livraison)."
    )
    add_figure(doc, DIAG_DIR / "s4_usecase.png",
               "Figure 53 : Diagramme de cas d'utilisation du Sprint 4")

    add_h3(doc, "4.2 Description des cas d'utilisation")
    add_para(doc,
        "La confirmatrice consulte les bons de commande, ajuste leur statut et déclenche la "
        "transformation BC vers BL, opération qui crée le bon de livraison en local et l'envoie "
        "à Sage X3 sans bloquer le processus. Elle traite également les devis B2B et les "
        "réclamations clients. Le livreur consulte le pool des livraisons filtré par sa zone, "
        "prend en charge une livraison disponible, diffuse sa position GPS, met à jour le statut "
        "(huit codes de 0 à 7) et gère sa caisse COD. Le livreur de transit se concentre sur "
        "les missions de transfert entre dépôts par double scan de code-barres. Le client suit "
        "sa commande grâce à une carte adaptative en quatre états et peut déposer une "
        "réclamation. Le superviseur surveille l'activité en temps réel, consulte la heatmap des "
        "échecs et réaffecte les missions de transit lorsque c'est nécessaire."
    )

    add_h3(doc, "4.3 Diagrammes de séquence du Sprint 4")
    add_para(doc,
        "Quatre flux temps réel ont été modélisés. Ils précisent les échanges entre l'interface, "
        "la Web API, la base SQL Server, le hub SignalR et, lorsque c'est nécessaire, les "
        "services externes Sage X3 et n8n."
    )

    add_h4(doc, "4.3.1 Transformation BC en bon de livraison")
    add_para(doc,
        "La méthode TransformBcToBl de ConfirmateurController vérifie d'abord l'existence du BC. "
        "Si aucun BL n'a déjà été créé, elle ouvre une transaction, insère l'entête F_DOCENTETE "
        "avec DO_Type = 1 et la référence formée de « BL » suivi de la date et d'un compteur, "
        "puis recopie les lignes F_DOCLIGNE. Elle crée ensuite la livraison F_LIVRAISON avec "
        "LI_Statut = 0 (Confirme) et passe le BC en statut CONFIRME avant le COMMIT. "
        "L'AssignmentService sélectionne un livreur dont la zone couvre l'adresse du client "
        "(test Haversine) et dont la file de livraisons en cours est la plus courte. L'envoi vers "
        "Sage X3 est asynchrone : la réponse expose sageSent, sageSuccess et sageHttpStatus pour "
        "permettre la supervision."
    )
    add_figure(doc, DIAG_DIR / "s4_seq_bctbl.png",
               "Figure 54 : Diagramme de séquence — Transformation BC en BL")

    add_h4(doc, "4.3.2 Livraison COD et diffusion GPS via SignalR")
    add_para(doc,
        "Le livreur démarre la livraison en appelant /start-heading. La Web API marque la "
        "livraison comme active (IsActiveDelivery = true) et envoie l'événement DeliveryStarted "
        "au groupe SignalR client-{userId}. Le client bascule alors sur l'état HEADING_TO_YOU. "
        "Toutes les 10 secondes, le Geolocator de Flutter envoie un ping contenant la latitude, "
        "la longitude et un identifiant ClientActionId. La Web API effectue un UPSERT dans "
        "F_LIVREUR_POSITION, ajoute une ligne dans F_LIVREUR_POSITION_HISTORY, calcule la "
        "distance Haversine vers le client et émet l'événement LivreurPositionUpdate. "
        "Si la distance descend sous 500 mètres, un événement de proximité est aussi envoyé. "
        "En l'absence de réseau, les positions sont stockées localement dans une file Hive et "
        "envoyées plus tard via /ping-batch grâce à l'idempotence garantie par ClientActionId. "
        "Lorsque le livreur valide la livraison (statut 2), F_LIVRAISON est mise à jour avec "
        "Encaisse = true et MontantEncaisse = DO_NetAPayer ; le client est notifié et la "
        "diffusion GPS s'arrête."
    )
    add_figure(doc, DIAG_DIR / "s4_seq_gps.png",
               "Figure 55 : Diagramme de séquence — Livraison COD et diffusion GPS")

    add_h4(doc, "4.3.3 Réclamation client avec SignalR et grâce de 5 secondes")
    add_para(doc,
        "La création d'une réclamation par un client provoque l'envoi de l'événement NouveauCas "
        "au groupe confirmateurs. La confirmatrice qui prend en charge le cas ouvre une session "
        "F_CONFIRMATRICE_SESSION et verrouille le cas. La discussion se déroule via l'événement "
        "NouveauMessage diffusé au groupe client-{userId} et au groupe confirmateurs. La "
        "particularité de ce flux est la gestion de la grâce de 5 secondes à la déconnexion. "
        "La méthode OnDisconnectedAsync du ReclamationHub décrémente le compteur de connexions de "
        "la confirmatrice. Si c'est la dernière connexion, un timer GracePeriod = "
        "TimeSpan.FromSeconds(5) est armé. Si la confirmatrice se reconnecte dans cet intervalle "
        "(typiquement un basculement Wi-Fi vers 4G), le timer est annulé et le verrou conservé. "
        "Si la grâce expire, le cas redevient libre et l'événement CommandeLiberee est diffusé. "
        "Lorsque trois tentatives de livraison ont échoué (F_RECLAMATION_TENTATIVE), l'événement "
        "SeuilTentativesAtteint est émis vers les superviseurs."
    )
    add_figure(doc, DIAG_DIR / "s4_seq_reclamation.png",
               "Figure 56 : Diagramme de séquence — Réclamation avec grâce 5 s SignalR")

    add_h4(doc, "4.3.4 Transit inter-dépôts par double scan code-barres")
    add_para(doc,
        "Le livreur de transit ouvre la mission dans transit_mission_details_screen. Le scanner "
        "transit_barcode_scanner_screen utilise le plugin mobile_scanner pour lire les codes. "
        "Un premier scan (POST /scan-pickup) fait passer la ligne F_TRANSFERT de "
        "EN_ATTENTE_TRANSIT à EN_TRANSIT, enregistre PickedUpAt et la position GPS, et écrit un "
        "journal F_TRANSFERT_AUDIT_LOG avec ActionType = PICKUP. Le second scan "
        "(POST /scan-delivery) au dépôt destination bascule la ligne en RECU_DEPOT_DESTINE. Une "
        "fenêtre d'annulation de dix minutes (RevertWindowMinutes = 10) permet de corriger une "
        "erreur, avec un nouveau journal d'audit. Si la quantité reçue est inférieure à la "
        "quantité attendue, l'application déclenche POST /scan-partial : la ligne passe en "
        "TRANSIT_PARTIELLEMENT_RECU et une alerte est envoyée au groupe superviseurs."
    )
    add_figure(doc, DIAG_DIR / "s4_seq_transit.png",
               "Figure 57 : Diagramme de séquence — Transit inter-dépôts (double scan)")

    # 5. Réalisation
    add_h2(doc, "5. Réalisation du Sprint 4")
    add_para(doc,
        "La réalisation utilise les trois canaux de la solution. Les écrans React de la "
        "confirmatrice et du superviseur consomment les endpoints via Axios. Les écrans Flutter "
        "passent par un ApiClient qui injecte automatiquement le jeton JWT stocké en sécurisé "
        "(flutter_secure_storage). Les notifications sonores des scans utilisent audioplayers et "
        "les vibrations sont gérées par le plugin du scanner. Les photos des réclamations sont "
        "compressées par flutter_image_compress avant l'envoi pour réduire le coût réseau. Les "
        "captures suivantes présentent les principaux écrans réalisés ; les emplacements vides "
        "correspondent aux images à insérer au moment de la mise au point finale."
    )

    real = [
        ("5.1 Interface confirmatrice — liste des bons de commande",
         "L'écran ConfirmateurOrdersPage liste les BC paginés. La confirmatrice peut filtrer par "
         "statut (EN_ATTENTE, CONFIRME, EN_LIVRAISON, LIVRE), rechercher par référence ou par "
         "client et consulter le détail. Chaque ligne affiche la référence DO_Piece, le client, "
         "la date, le montant TTC et le libellé du statut.",
         "liste des bons de commande confirmatrice",
         "Figure 58 : Interface confirmatrice — liste des bons de commande"),
        ("5.2 Interface confirmatrice — transformation BC en BL",
         "Lorsque le BC est confirmé, l'écran ConfirmateurOrderDetailsPage propose le bouton "
         "« Transformer en BL ». Le résultat de l'opération est affiché : référence du nouveau BL, "
         "indicateur alreadyExists si un BL existe déjà, sageSent et sageSuccess pour le suivi de "
         "la synchronisation Sage X3. L'identifiant du livreur automatiquement affecté est "
         "également rappelé.",
         "transformation BC en BL",
         "Figure 59 : Interface confirmatrice — transformation BC en BL"),
        ("5.3 Interface confirmatrice — gestion des réclamations",
         "L'écran des réclamations affiche les cas par ordre de priorité et d'ancienneté. "
         "Les nouvelles réclamations apparaissent en temps réel grâce à l'événement SignalR "
         "NouveauCas, sans rechargement de page. La confirmatrice peut prendre en charge un "
         "cas, modifier son statut, envoyer une demande de correction d'adresse, créer un bon "
         "d'échange, ajouter une note interne, réaffecter le cas et chatter avec le client. "
         "L'historique des tentatives de livraison est consultable, avec mise en évidence du "
         "seuil de trois tentatives.",
         "gestion des réclamations + chat SignalR",
         "Figure 60 : Interface confirmatrice — gestion des réclamations"),
        ("5.4 Interface livreur — livraisons disponibles",
         "L'écran new_orders_screen présente les bons de livraison filtrés par gouvernorat et "
         "délégation du livreur. Pour chaque livraison, l'application affiche l'adresse, le "
         "montant COD attendu et les coordonnées du client. La prise en charge se fait d'un seul "
         "appui ; si un autre livreur prend la livraison en parallèle, l'API renvoie un code 409 "
         "Conflict et l'utilisateur en est informé.",
         "pool livraisons côté livreur",
         "Figure 61 : Interface livreur — livraisons disponibles"),
        ("5.5 Interface livreur — suivi GPS et carte de tournée",
         "La carte affiche la position du livreur (mise à jour via SignalR), les arrêts restants "
         "et l'itinéraire optimisé. L'algorithme du plus-proche-voisin avec distance Haversine "
         "ordonne les arrêts, et OSRM fournit un calcul d'ETA plus précis. La feuille "
         "LiveDeliveryMapSheet expose au client la position du livreur lorsqu'il est en route, "
         "avec un indicateur de fraîcheur du dernier ping (vert, orange ou rouge).",
         "suivi GPS et tournée",
         "Figure 62 : Interface livreur — suivi GPS et tournée"),
        ("5.6 Interface livreur — caisse COD",
         "L'écran livreur_stats_screen présente la caisse du livreur : montant total encaissé, "
         "nombre de livraisons effectuées, taux de réussite et répartition par statut. "
         "L'historique des remises au dépôt est accessible. Les sous-états DepotEnCoursDePreparation "
         "(6) et DepotPret (7) sont visibles uniquement par le livreur, le client voyant un état "
         "agrégé « au dépôt ».",
         "caisse COD livreur",
         "Figure 63 : Interface livreur — caisse COD et statistiques"),
        ("5.7 Interface client — TrackingStateCard adaptative",
         "Le composant TrackingStateCard prend l'apparence dictée par l'endpoint "
         "/api/client/orders/{piece}/tracking-state. En AT_DEPOT, la carte est informative. En "
         "IN_DELIVERY_QUEUE, le client sait qu'un livreur va le prendre en charge. En "
         "HEADING_TO_YOU, la carte devient une vue cartographique avec la position GPS du "
         "livreur, l'ETA en minutes (calculé par Haversine à 40 km/h) et un bouton d'appel ou de "
         "SMS. En TERMINAL, la livraison est finalisée (livrée, retour ou refusée).",
         "TrackingStateCard sur les 4 états",
         "Figure 64 : Interface client — TrackingStateCard adaptative"),
        ("5.8 Interface de transit inter-dépôts et scan code-barres",
         "L'écran transit_home_screen regroupe les missions en trois onglets (en attente, en "
         "cours, historique). Le détail d'une mission liste les articles attendus. Le scanner "
         "active la caméra ; chaque scan donne un retour sonore (audioplayers) et tactile. "
         "L'application gère également la réception partielle et la fenêtre d'annulation de "
         "dix minutes après un pickup.",
         "transit et scanner",
         "Figure 65 : Interface transit — missions et scanner code-barres"),
        ("5.9 Interface superviseur",
         "Le superviseur dispose d'une carte temps réel des livreurs actifs, d'un tableau de "
         "bord des alertes (retard, échec, zone non couverte, écart de caisse) et d'une heatmap "
         "construite par LivreurMapController.HeatMap. Les cellules sont d'environ 500 mètres et "
         "l'intensité représente le ratio (retours + reports) / livraisons sur les 90 derniers "
         "jours. Le superviseur peut aussi réaffecter une mission de transit grâce à l'endpoint "
         "/api/supervisor/transferts/{id}/reassign.",
         "console superviseur",
         "Figure 66 : Interface superviseur — livreurs actifs et heatmap"),
        ("5.10 Services d'arrière-plan et qualité de service",
         "Plusieurs services tournent en arrière-plan pour garantir la qualité du flux logistique. "
         "ReclamationRedistributionHostedService redistribue les cas non assignés. "
         "StaleLockCleanupHostedService libère les verrous bloqués. Le job Hangfire "
         "TransitEscalation24hJob alerte le superviseur lorsque des missions de transit restent "
         "non prises pendant plus de 24 heures. L'ensemble de ces mécanismes assure la robustesse "
         "du sprint en situation réelle.",
         "panneau services et alertes",
         "Figure 67 : Services d'arrière-plan et tableau d'alertes"),
    ]
    for title, body, ph, cap in real:
        add_h3(doc, title)
        add_para(doc, body)
        add_placeholder(doc, ph)
        add_caption(doc, cap)

    # 6. Tests
    add_h2(doc, "6. Tests et validation du Sprint 4")
    add_para(doc,
        "Les tests fonctionnels ont été menés tout au long du sprint sur des jeux de données "
        "réalistes (jeu de départ Sage X3 et seed Tunisie : neuf gouvernorats avec leurs "
        "délégations). Les cas limites — Sage X3 indisponible, conflit de prise en charge, perte "
        "de connexion du livreur, déconnexion fugace de la confirmatrice, scans erronés — ont "
        "été vérifiés en plus des cas nominaux."
    )
    t = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["ID", "Objectif", "Résultat attendu", "Statut"]):
        t.rows[0].cells[i].text = h
    for row in S4_TESTS:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = v
    style_table(t)
    add_caption(doc, "Tableau 14 : Tests fonctionnels du Sprint 4")

    # 7. Conclusion
    add_h2(doc, "7. Conclusion")
    add_para(doc,
        "Le Sprint 4 referme la chaîne logistique de la plateforme. Une commande passée en "
        "ligne peut désormais être confirmée, transformée en bon de livraison synchronisé avec "
        "Sage X3, affectée automatiquement à un livreur, suivie en temps réel par le client, "
        "livrée avec encaissement COD automatique, puis suivie d'une éventuelle réclamation "
        "traitée en conversation directe. Le transit inter-dépôts apporte la traçabilité des "
        "mouvements internes par double scan. Le superviseur dispose d'outils pour suivre "
        "l'activité et intervenir en cas d'écart. La résilience opérationnelle a été soignée : "
        "file locale hors ligne, idempotence des pings, grâce SignalR, synchronisation Sage "
        "non bloquante, audit log et jobs Hangfire d'escalade. Le sprint suivant exploite ce "
        "socle pour offrir à l'administrateur une couche complète de pilotage et d'aide à la "
        "décision."
    )

# ─────────────────────────────────────────────────────────────────────
# CHAPITRE 8 — Sprint 5
# ─────────────────────────────────────────────────────────────────────
def build_chap8(doc):
    doc.add_page_break()
    add_h1(doc, "Chapitre 8 : Sprint 5 — Pilotage, tableaux de bord, chatbot et aide à la décision")

    add_h2(doc, "1. Introduction")
    add_para(doc,
        "Ce chapitre présente le cinquième et dernier sprint de la solution. À ce stade, "
        "l'ensemble des fonctionnalités opérationnelles — catalogue, commandes, paiement, "
        "confirmation, livraison, réclamations, transit — est en place. Le Sprint 5 ajoute "
        "la couche de pilotage et d'aide à la décision destinée à l'administrateur. Trois "
        "axes sont traités en parallèle : les tableaux de bord multi-dimensionnels qui "
        "agrègent les données de tous les modules précédents, les exports Excel et PDF qui "
        "permettent d'extraire ces données vers des outils externes, et un chatbot "
        "administrateur capable d'interroger la base en langage naturel grâce à n8n, un "
        "modèle de langage Groq et trois modèles ML.NET."
    )

    add_h2(doc, "2. Objectif et périmètre du Sprint 5")
    add_para(doc,
        "Le Sprint 5 s'est déroulé du 27/04/2026 au 23/05/2026. Son objectif est de "
        "transformer la solution en une plateforme pilotable. L'administrateur doit pouvoir "
        "consulter les indicateurs clés, identifier les anomalies, extraire les données et "
        "obtenir des réponses synthétiques à des questions en français. Le périmètre couvre "
        "onze tableaux de bord, deux services d'export, un chatbot conversationnel multimodal "
        "(texte + vocal sur mobile) et un mécanisme d'insights proactifs qui détecte "
        "automatiquement les tendances et les écarts."
    )

    add_h2(doc, "3. Backlog du Sprint 5")
    add_para(doc,
        "Le tableau ci-après présente les vingt-deux histoires utilisateurs du sprint. Elles "
        "se répartissent entre les tableaux de bord, les exports, le chatbot et les services "
        "d'arrière-plan associés."
    )
    t = doc.add_table(rows=1, cols=5)
    for i, h in enumerate(["ID", "Histoire utilisateur", "Tâches techniques", "Responsable", "Estim."]):
        t.rows[0].cells[i].text = h
    for row in S5_BACKLOG:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = v
    style_table(t)
    add_caption(doc, "Tableau 15 : Backlog du Sprint 5")

    add_h2(doc, "4. Analyse et conception du Sprint 5")
    add_para(doc,
        "L'analyse de ce sprint porte sur des fonctionnalités transverses qui réutilisent les "
        "données produites par les sprints précédents. Les tableaux de bord s'appuient sur un "
        "service d'agrégation unique. Les exports sont centralisés dans un service générique. "
        "Le chatbot est orchestré par un pipeline interne en trois étapes (Router, Exécuteur, "
        "Formatter) et complété par un workflow n8n qui injecte des insights proactifs."
    )

    add_h3(doc, "4.1 Diagramme de cas d'utilisation du Sprint 5")
    add_para(doc,
        "Le diagramme regroupe les fonctionnalités en quatre paquets : tableaux de bord, "
        "exports, chatbot et pilotage avancé. Il fait apparaître les trois acteurs techniques "
        "qui interviennent en arrière-plan : le workflow n8n pour l'orchestration, le modèle "
        "Groq llama-3.3-70b-versatile pour le routage et la mise en forme des réponses, et "
        "ML.NET pour les prédictions. Sage X3 reste la source du tableau de bord de "
        "synchronisation."
    )
    add_figure(doc, DIAG_DIR / "s5_usecase.png",
               "Figure 68 : Diagramme de cas d'utilisation du Sprint 5")

    add_h3(doc, "4.2 Description des cas d'utilisation")
    add_para(doc,
        "L'administrateur accède à la vue d'ensemble, à l'analyse des ventes, des commandes, "
        "des produits, des stocks et des dépôts, de la logistique, des livreurs, des clients, "
        "des réclamations, des confirmatrices et de la synchronisation Sage X3. Il consulte "
        "les insights stratégiques et les insights proactifs générés en arrière-plan, déclenche "
        "les exports Excel et PDF, dialogue avec le chatbot en mode texte ou voix, parcourt "
        "l'historique des conversations et peut rafraîchir manuellement la base de "
        "connaissances utilisée par le router Groq."
    )

    add_h3(doc, "4.3 Diagrammes de séquence du Sprint 5")
    add_para(doc,
        "Deux flux ont été modélisés en détail : la chaîne complète du chatbot administrateur "
        "(React → n8n → Backend → Groq → ML.NET) et la génération des exports Excel et PDF par "
        "le service ExportService."
    )

    add_h4(doc, "4.3.1 Diagramme de séquence du chatbot administrateur")
    add_para(doc,
        "L'administrateur saisit sa question depuis la page ChatbotSandboxPage. Le payload "
        "(question, sessionId, locale) est envoyé au webhook n8n /admin-chat-v3. Le premier "
        "nœud détecte la langue (français, arabe ou « tounsi ») puis transmet la requête au "
        "backend via POST /api/admin/chat/ask. L'AdminChatOrchestratorService applique alors le "
        "pipeline interne : un premier appel à Groq joue le rôle de routeur et choisit l'action "
        "(KB, QUERY, ANALYZE, PREDICT ou CHITCHAT). L'action est exécutée : interrogation SQL "
        "pour QUERY, modèle ML.NET pour PREDICT, lecture de la base de connaissances pour KB. "
        "Un second appel à Groq met en forme la réponse en français et propose éventuellement "
        "une structure de tableau et de graphique. Le message est journalisé dans "
        "F_CHATBOT_MESSAGE. Côté n8n, un nœud complémentaire récupère les insights proactifs en "
        "attente (F_CHATBOT_INSIGHTS) et les fusionne à la réponse avant le retour à React. "
        "Une variante de cet endpoint, POST /api/admin/chat/ask-stream, renvoie un flux SSE en "
        "quatre phases (routing, data, chunks, done) pour un rendu mot-à-mot."
    )
    add_figure(doc, DIAG_DIR / "s5_seq_chatbot.png",
               "Figure 69 : Diagramme de séquence — Chatbot administrateur (React → n8n → Groq)")

    add_h4(doc, "4.3.2 Diagramme de séquence d'export Excel et PDF")
    add_para(doc,
        "L'administrateur déclenche un export depuis l'interface de synthèse. La route "
        "GET /api/admin/orders/export reçoit deux paramètres : format (xlsx ou pdf) et "
        "period. La méthode parse la période, projette F_DOCENTETES dans la limite de "
        "MaxRows = 10 000 lignes et délègue la sérialisation à ExportService. ClosedXML construit "
        "le classeur Excel avec entêtes en gras et colonnes ajustées automatiquement. QuestPDF "
        "construit un document A4 avec en-tête contextualisé (titre, période), un tableau "
        "paginé et un pied de page. Le flux binaire est renvoyé en téléchargement direct via "
        "l'en-tête Content-Disposition. Le même service est réutilisé pour "
        "/api/admin/reclamations/export."
    )
    add_figure(doc, DIAG_DIR / "s5_seq_export.png",
               "Figure 70 : Diagramme de séquence — Export Excel et PDF")

    # 5. Réalisation
    add_h2(doc, "5. Réalisation du Sprint 5")
    add_para(doc,
        "Les onze tableaux de bord partagent une même infrastructure côté React : le composant "
        "DashboardAnalyticsPage est paramétré par une clé pageKey qui détermine l'endpoint à "
        "appeler, la configuration des KPI et la liste des alertes à afficher. Le rendu des "
        "graphiques utilise Recharts. Les exports sont déclenchés depuis les écrans de synthèse "
        "via un appel HTTP qui produit un téléchargement direct. Le chatbot est accessible "
        "depuis une fenêtre flottante (ChatbotFab) sur toutes les pages d'administration."
    )

    real5 = [
        ("5.1 Interface du tableau de bord d'ensemble",
         "L'écran AdminOverviewDashboardPage agrège les indicateurs clés : chiffre d'affaires, "
         "nombre de livraisons, réclamations ouvertes, livreurs actifs et tendance "
         "hebdomadaire. Le panneau d'alertes AlertPanel met en évidence les anomalies par "
         "sévérité (avertissement, critique). Les insights stratégiques et les insights "
         "proactifs alimentés par Hangfire toutes les 30 minutes y sont également visibles.",
         "tableau de bord d'ensemble",
         "Figure 71 : Interface — Tableau de bord d'ensemble"),
        ("5.2 Interface du dashboard ventes et commandes",
         "AdminSalesDashboardPage présente le chiffre d'affaires sur la période, le top "
         "produits et la répartition B2C / B2B. AdminOrdersDashboardPage détaille les statuts "
         "des commandes, les délais moyens (BC → BL → livraison) et la segmentation "
         "géographique. Les graphiques sont produits par Recharts à partir du JSON retourné "
         "par la Web API.",
         "dashboard ventes et commandes",
         "Figure 72 : Interface — Dashboards ventes et commandes"),
        ("5.3 Interface du dashboard logistique et livreurs",
         "AdminLogisticsDashboardPage présente le taux de livraison par zone, la heatmap des "
         "échecs sur 90 jours et le délai moyen de livraison. AdminDriversDashboardPage "
         "expose le classement des livreurs (taux de réussite, volume, caisse COD) et leurs "
         "alertes individuelles. Les huit codes de statut sont restitués séparément pour "
         "permettre une lecture précise.",
         "dashboards logistique et livreurs",
         "Figure 73 : Interface — Dashboards logistique et livreurs"),
        ("5.4 Interface du dashboard réclamations et confirmatrices",
         "AdminReclamationsDashboardPage croise les réclamations par motif (COLIS_ENDOMMAGE, "
         "NON_LIVRE, MAUVAIS_ARTICLE, NUMERO_INCORRECT, ADRESSE_INCORRECTE) et par statut "
         "(ENVOYEE, EN_COURS, CLOTUREE, REFUSEE). Elle suit également les performances des "
         "confirmatrices : cas actifs, taux de résolution, temps moyen et durée des sessions "
         "F_CONFIRMATRICE_SESSION.",
         "dashboards réclamations et confirmatrices",
         "Figure 74 : Interface — Dashboards réclamations et confirmatrices"),
        ("5.5 Interface du dashboard synchronisation Sage X3",
         "AdminSyncDashboardPage présente l'état des données synchronisées : nombre "
         "d'articles, de catalogues, de dépôts et de stocks à jour, alertes d'intégrité "
         "(clients sans adresse, articles sans famille, stocks négatifs) et horodatage de la "
         "dernière synchronisation. Elle permet à l'administrateur de vérifier en permanence "
         "la cohérence entre la plateforme et l'ERP.",
         "dashboard synchronisation Sage X3",
         "Figure 75 : Interface — Dashboard synchronisation Sage X3"),
        ("5.6 Interface des exports Excel et PDF",
         "L'écran de synthèse propose, à côté de chaque liste, deux boutons « Exporter en "
         "Excel » et « Exporter en PDF ». Les fichiers sont générés par ExportService "
         "(ClosedXML 0.104.1 et QuestPDF 2024.7.3) dans la limite de 10 000 lignes par "
         "export. La nomenclature des fichiers inclut la période sélectionnée pour faciliter "
         "le classement.",
         "interface des exports",
         "Figure 76 : Interface — Exports Excel et PDF"),
        ("5.7 Interface du chatbot administrateur",
         "La page ChatbotSandboxPage offre une zone de saisie, l'historique du fil de "
         "conversation et le rendu structuré des réponses (message, tableau, graphique). Le "
         "bouton ChatbotFab permet d'ouvrir le chatbot depuis n'importe quelle page "
         "d'administration. Les exemples de questions traitées sont nombreux : « CA du jour », "
         "« Top 5 produits du mois », « Livreurs à plus de 10 % de retours », « Réclamations "
         "ouvertes depuis plus de 48 heures ». Une bascule active le streaming SSE pour "
         "afficher la réponse mot à mot.",
         "chatbot administrateur",
         "Figure 77 : Interface — Chatbot administrateur"),
        ("5.8 Interface des conversations et insights chatbot",
         "ChatbotConversationsPage parcourt l'historique des sessions, filtrable par date et "
         "par utilisateur. ChatbotInsightsPage présente les insights proactifs sous forme de "
         "cartes : chaque insight expose un titre, un constat chiffré et une action suggérée. "
         "Un bouton de feedback alimente la qualité des futures recommandations. L'endpoint "
         "/api/admin/chat/kb/refresh permet de régénérer la base de connaissances qui sert de "
         "contexte au router Groq.",
         "conversations et insights chatbot",
         "Figure 78 : Interface — Conversations et insights chatbot"),
        ("5.9 Chatbot vocal sur mobile",
         "Sur l'application mobile, l'écran admin_chat_screen expose un bouton micro et un "
         "bouton haut-parleur (voice_buttons.dart). speech_to_text capte la question et "
         "flutter_tts lit la réponse à voix haute. Le mode vocal facilite l'usage du chatbot "
         "en mobilité, par exemple lors d'une réunion ou d'une visite sur le terrain.",
         "chatbot vocal mobile",
         "Figure 79 : Interface — Chatbot vocal sur mobile"),
        ("5.10 Services d'arrière-plan",
         "Trois services tournent en permanence pour alimenter la couche de pilotage. "
         "KbGeneratorService régénère la base de connaissances exploitée par le router Groq. "
         "ProactiveInsightsJob (Hangfire, toutes les 30 minutes) calcule les insights et les "
         "stocke dans F_CHATBOT_INSIGHTS. DepotIncrementJob (quotidien à minuit) met à jour "
         "les compteurs des dépôts. L'ensemble est paramétrable depuis l'écran "
         "AdminWorkflowScreen côté mobile.",
         "supervision des jobs Hangfire",
         "Figure 80 : Interface — Supervision des jobs Hangfire"),
    ]
    for title, body, ph, cap in real5:
        add_h3(doc, title)
        add_para(doc, body)
        add_placeholder(doc, ph)
        add_caption(doc, cap)

    # 6. Tests
    add_h2(doc, "6. Tests et validation du Sprint 5")
    add_para(doc,
        "Les tests vérifient la cohérence des indicateurs par rapport aux données sources, "
        "la conformité des artefacts d'export aux spécifications et la qualité des réponses du "
        "chatbot. Les tests automatisés du projet Web-Api.Tests (xUnit + EF InMemory) couvrent "
        "les modules sensibles (paiements virtuels, devis B2B, favoris, géofencing). Les tests "
        "manuels suivants complètent l'évaluation fonctionnelle du sprint."
    )
    t = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["ID", "Objectif", "Résultat attendu", "Statut"]):
        t.rows[0].cells[i].text = h
    for row in S5_TESTS:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = v
    style_table(t)
    add_caption(doc, "Tableau 16 : Tests fonctionnels du Sprint 5")

    add_h2(doc, "7. Conclusion")
    add_para(doc,
        "Le Sprint 5 referme le projet en plaçant l'administrateur au centre d'une plateforme "
        "instrumentée. Les onze tableaux de bord offrent une vision unifiée et homogène — un "
        "contrat unique ProDashboardPageResponseDto et un composant React unique. Les exports "
        "Excel et PDF construits avec ClosedXML et QuestPDF permettent de matérialiser cette "
        "vision dans des supports partageables. Le chatbot orchestré par n8n, par le pipeline "
        "interne Router-Exécuteur-Formatter et par trois modèles ML.NET, agit comme une "
        "couche d'accès conversationnelle (texte sur le web, voix sur le mobile) qui rapproche "
        "les données métier des questions opérationnelles quotidiennes. Les insights proactifs, "
        "calculés en arrière-plan par Hangfire, anticipent les besoins de l'administrateur. "
        "Avec ce cinquième sprint, la solution atteint un niveau de maturité opérationnelle "
        "compatible avec une démonstration finale et une mise en production progressive."
    )

# ─────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────
def main():
    doc = Document()
    # marges identiques à la copie 06
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(1.94)
    s.bottom_margin = Cm(0.49)
    s.left_margin = Cm(1.76)
    s.right_margin = Cm(0.67)

    # styles
    for sname, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12), ("Heading 4", 12), ("Normal", 12)]:
        try:
            doc.styles[sname].font.color.rgb = BLACK
            doc.styles[sname].font.size = Pt(size)
        except Exception:
            pass

    build_chap7(doc)
    build_chap8(doc)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"OK : {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
